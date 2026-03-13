# -*- coding: utf-8 -*-
import streamlit as st
import requests
import pandas as pd
import plotly.graph_objects as go
import numpy as np
from datetime import datetime, timedelta
import re, io, random, json, os, smtplib, threading
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from collections import Counter
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import streamlit.components.v1 as components

try:
    from apscheduler.schedulers.background import BackgroundScheduler
    from apscheduler.triggers.cron import CronTrigger
    SCHEDULER_OK = True
except:
    SCHEDULER_OK = False

try:
    import yfinance as yf
    YF_OK = True
except:
    YF_OK = False

# ── 보안: API 키는 Streamlit Secrets에서 로드 ──────────
CLIENT_ID     = st.secrets["NAVER_CLIENT_ID"]
CLIENT_SECRET = st.secrets["NAVER_CLIENT_SECRET"]
APP_URL       = st.secrets.get("APP_URL", "https://kepco-news-monitor-gbff2xm5nzatkkmvjsd9tm.streamlit.app")
FONT_KR = "'Noto Sans KR', 'Apple SD Gothic Neo', 'Malgun Gothic', Arial, sans-serif"

MEDIA_GRADE = {
    "조선일보":{"rank":1,"rate":3.73,"grade":"S"},"중앙일보":{"rank":2,"rate":2.45,"grade":"A"},
    "동아일보":{"rank":3,"rate":1.95,"grade":"A"},"매일경제":{"rank":4,"rate":0.97,"grade":"A"},
    "한겨레":{"rank":5,"rate":0.62,"grade":"B"},"한국경제":{"rank":6,"rate":0.43,"grade":"B"},
    "경향신문":{"rank":7,"rate":0.41,"grade":"B"},"연합뉴스":{"rank":8,"rate":0.38,"grade":"B"},
    "YTN":{"rank":9,"rate":0.35,"grade":"B"},"KBS":{"rank":10,"rate":0.30,"grade":"B"},
    "MBC":{"rank":11,"rate":0.28,"grade":"B"},"SBS":{"rank":12,"rate":0.25,"grade":"B"},
    "한국일보":{"rank":13,"rate":0.31,"grade":"B"},"국민일보":{"rank":14,"rate":0.19,"grade":"C"},
    "문화일보":{"rank":15,"rate":0.12,"grade":"C"},"서울신문":{"rank":16,"rate":0.10,"grade":"C"},
    "서울경제":{"rank":17,"rate":0.08,"grade":"C"},"세계일보":{"rank":18,"rate":0.02,"grade":"D"},
    "머니투데이":{"rank":19,"rate":0.05,"grade":"D"},"이데일리":{"rank":20,"rate":0.04,"grade":"D"},
    "헤럴드경제":{"rank":21,"rate":0.03,"grade":"D"},"뉴시스":{"rank":22,"rate":0.02,"grade":"D"},
    "뉴스핌":{"rank":23,"rate":0.02,"grade":"D"},
}
GRADE_COLOR = {"S":"#B71C1C","A":"#E64A19","B":"#1565C0","C":"#2E7D32","D":"#616161"}

POSITIVE_WORDS = ["성과","달성","개선","혁신","성장","우수","협력","기여","선정","수상","기대","친환경","안정","흑자","증가","추진","완료","승인","투자","확대","증대","국익","선도","도약","강화","지원","활성화","성공","최초","호평","기록","돌파","수주","계약","협약","양해각서","출범","개통","준공","신기록"]
NEGATIVE_WORDS = ["사고","부패","비리","손실","적자","위기","파업","갈등","논란","문제","우려","하락","감소","실패","사망","중단","취소","반발","지연","비판","폭탄","부담","폐쇄","위반","처벌","고발","감사원","적발","의혹","과태료","경고","해임","부실","낭비","특혜","불법","소송","압수수색","조사","수사","민원","항의","경찰","패소","기소","고소","피해","피의자","벌금","구속","재판","징계","결함","은폐","허위","과장","오염","누출","폭발","붕괴","침수","정전"]

IRRELEVANT_PATTERNS = [
    r'배구',r'축구',r'야구',r'농구',r'골프',r'올림픽',r'월드컵',r'선수단',r'드래프트',r'챔피언십',
    r'세트스코어',r'\d+세트',r'배구단',r'득점왕',r'홈런',r'퇴장',r'심판',r'리그',r'경기장',
    r'선수(?:\s|가|은|는|이|을|의|도)',r'감독(?:\s|이|은|의)',r'코치(?:\s|이|의)',
]

MEDIA_MAP = {
    "chosun":"조선일보","joongang":"중앙일보","donga":"동아일보","hani":"한겨레","khan":"경향신문",
    "yna":"연합뉴스","ytn":"와이티엔","imnews":"MBC","kbs":"KBS","sbs":"SBS","mt.co":"머니투데이",
    "edaily":"이데일리","heraldcorp":"헤럴드경제","newsis":"뉴시스","newspim":"뉴스핌",
    "etnews":"전자신문","energy-news":"에너지신문","electimes":"일렉트릭타임스",
    "hankyung":"한국경제","mk.co":"매일경제","sedaily":"서울경제","ajunews":"아주경제",
    "businesspost":"비즈니스포스트","fnnews":"파이낸셜뉴스","inews24":"아이뉴스24",
    "dt.co":"디지털타임스","hankookilbo":"한국일보","munhwa":"문화일보","ohmynews":"오마이뉴스",
    "pressian":"프레시안","energydaily":"에너지데일리","naeil":"내일신문","seoul":"서울신문",
    "ekn":"에너지경제","kukminilbo":"국민일보","segyetimes":"세계일보","e2news":"이투뉴스",
    "nocutnews":"노컷뉴스","hani.co":"한겨레","kmib":"국민일보","donga.com":"동아일보",
    "joins":"중앙일보","chosun.com":"조선일보","jtbc":"JTBC","tvchosun":"TV조선",
    "mbn":"MBN","channel.a":"채널A","yonhap":"연합뉴스","asiae":"아시아경제",
    "koreaherald":"코리아헤럴드","koreatimes":"코리아타임스","koreajoongang":"중앙일보",
    "dailian":"데일리안","newdaily":"뉴데일리","pennmike":"펜앤드마이크",
    "wikileaks":"위키리크스코리아","mediatoday":"미디어오늘","pdjournal":"PD저널",
    "journalist":"저널리스트","mindlenews":"마인들뉴스","naver":"네이버뉴스",
    "daum":"다음뉴스","kakao":"카카오","biz.chosun":"조선비즈","news.chosun":"조선일보",
    "news.jtbc":"JTBC","news1":"뉴스1","tf.co":"더팩트","sisajournal":"시사저널",
    "weekly.chosun":"주간조선","monthly.chosun":"월간조선","economist":"이코노미스트",
    "moneys":"머니S","bizwatch":"비즈워치","thebell":"더벨","hankyoreh":"한겨레",
    "kookje":"국제신문","idomin":"경남도민일보","gnmaeil":"경남매일","busan":"부산일보",
    "knnews":"경인일보","kyeonggi":"경기일보","joongdo":"중도일보","daejonilbo":"대전일보",
    "jeonbuk":"전북일보","jnilbo":"전남일보","gwangnam":"광남일보","jemin":"제민일보",
    "jejunews":"제주일보","kado":"강원도민일보","kwnews":"강원일보",
}
# 표시용 한글 매체명 정규화 (영문 약칭 → 한글)
MEDIA_KR_NAME = {
    "YTN":"와이티엔","KBS":"케이비에스","MBC":"MBC","SBS":"SBS",
    "JTBC":"JTBC","MBN":"MBN","TV조선":"TV조선","채널A":"채널A",
}


TOPIC_GROUPS = {
    "전기요금":["전기요금","요금","전력요금","인상","누진제","전기세"],
    "원전·수출":["원전","수출","원자력","UAE","체코","APR","해외수주"],
    "재무·경영":["흑자","적자","부채","재무","비상경영","원가","실적","손실"],
    "전력망·설비":["송전","배전","전력망","변전","선로","전력설비","정전","계통"],
    "탄소중립·에너지전환":["탄소중립","RE100","온실가스","수소","재생에너지","넷제로","태양광","풍력"],
    "노사관계":["노사","노조","파업","임금","단체협약","쟁의"],
    "안전·사고":["안전","사고","재해","산재","폭발","화재","부상","사망"],
    "AI·디지털혁신":["AI","인공지능","디지털","스마트","자동화","AX","빅데이터"],
    "공기업·거버넌스":["공기업","감사","이사회","투명","거버넌스","윤리","비리"],
    "고객·서비스":["서비스","고객","민원","복지","국민","전기복지"],
    "정책·규제":["정책","규제","법안","제도","정부","국회","의원","경찰","조사","소송"],
}
DISAMBIG_MAP = {"김동철":["한전","사장","한국전력","KEPCO"],"김성환":["장관","산업부"]}

# ── INSIGHT DB: 이슈별 PR 전략 (문서4 기반 전면 보강) ──────────────────
INSIGHT_DB = {
    "전기요금": {
        "bg": "요금 인상 이슈는 소비자 민감도 최상위. 물가 자극·부채·성과급 비판이 복합적으로 작동.",
        "sub_issues": [
            ("물가 자극 비판", "인상의 불가피성을 '국가적 비용' 관점에서 설명. 억제 시 채권 발행·금융 왜곡 데이터 제시."),
            ("200조+ 부채 보도", "부채 증가 속도 둔화·자구 대책 달성률을 헤드라인으로 추출."),
            ("연료비 연동제 유보 논란", "유보의 사회적 비용(미수금 증가)을 시각화. 정부 협의 현황 강조."),
            ("OECD 요금 비교", "세계 최저 수준 요금과 에너지 안보 위기를 연계한 팩트시트 배포."),
            ("흑자 전환 시 성과급 비판", "외부 요인(연료비 하락) 겸허 인정, 흑자분 부채 상환 우선 투입 명시."),
        ],
        "action": "요금 억제의 '숨겨진 비용' 시각화 자료 선제 배포. OECD 비교 팩트시트·미수금 누적 데이터 출입기자단 정기 제공.",
        "steps": [
            "원가 회수율·OECD 비교 팩트시트 즉시 배포",
            "취약계층 에너지 복지 수혜자 수 수치화·언론 제공",
            "요금 억제 시 국가 재무 왜곡 시뮬레이션 공개",
        ],
        "msg": "원가를 반영한 합리적 요금이 절실합니다.",
    },
    "재무·경영": {
        "bg": "사상 최대 부채 보도는 공기업 신뢰도를 직격. 흑자 전환 시에도 '성찰 없는 성과급' 비판이 역풍으로 작용.",
        "sub_issues": [
            ("사상 최대 부채 보도", "누적 부채액보다 '부채 증가 속도 둔화'와 '자구 대책 달성률'을 헤드라인으로 추출."),
            ("흑자 전환 시 성과급 비판", "연료비 하락 등 외부 요인을 겸허히 인정하고, 흑자분을 부채 상환에 우선 투입함을 명시."),
            ("비상경영 실효성 논란", "부동산 매각·조직 슬림화 자구 노력 이행률 100% 수치로 제시."),
        ],
        "action": "부채 증가 속도 둔화·자구 이행률 수치를 헤드라인으로 배치. 흑자분 부채 상환 우선 투입 계획 명시 공개.",
        "steps": [
            "부채 감축 로드맵·자구 이행률 정량 보도자료 즉시 배포",
            "부동산 매각·조직 효율화 구체 사례 스토리화",
            "경제지 전담 관계 강화 및 정기 백브리핑 개최",
        ],
        "msg": " 뼈를 깎는 자구 노력과 요금 정상화로 재무건전성을 강화하겠습니다.",
    },
    "전력망·설비": {
        "bg": "송전망 건설 지연은 반도체·산단 공급 차질로 이어지며 국익 이슈로 확산. 지역 반발과 지자체 인허가 거부가 복합 리스크.",
        "sub_issues": [
            ("반도체 산단 공급 차질", "전력망을 '에너지 고속도로'로 브랜딩, 국익 차원 시급성 호소. 전력망 특별법 필요성 강조."),
            ("지역 주민 송전탑 반대", "'보상' 위주에서 '지역 상생·공익 가치' 중심으로 전환. 주민 참여형 신재생 사업 모델 제시."),
            ("지자체 인허가 거부", "해당 지역 전력 자립도·산업 유치 저해 데이터를 정량 분석해 지역 언론 배포."),
            ("노후 설비 대규모 정전", "사고 원인 투명 공개, 설비 현대화 투자 예산 부족(재무 위기 연계) 솔직 설명."),
            ("그리드락(계통 접속 대기) 심화", "계통 여유 지도 공개, FACTS 등 기술적 해소 노력 홍보."),
        ],
        "action": "전력망을 '경제 동맥' 프레임으로 격상. 계통 여유 지도 공개 및 특별법 제정 필요성 공론화.",
        "steps": [
            "전력망 투자 로드맵·진행률 수치 중심 보도자료 배포",
            "계통 여유 지도 공개 및 접속 대기 해소 기술 홍보",
            "전력망 특별법 제정 필요성 의회·언론 동시 공론화",
        ],
        "msg": "에너지고속도로는 대한민국 경제의 미래입니다.",
    },
    "안전·사고": {
        "bg": "현장 사고는 사회적 충격이 크고 협력업체 차별·대형 재난 연루 의혹으로 확산되는 전형적 패턴.",
        "sub_issues": [
            ("현장 감전·추락 사고", "CEO 즉시 현장 방문. 48시간 내 재발 방지책 공식 발표. 감성적 사과와 이성적 대책 병행."),
            ("협력업체 차별 처우 논란", "'상생 협력 생태계' 강화 방안·협력사 안전 지원 예산 집행 현황 공개."),
            ("산불 등 대형 재난 원인 의혹", "정밀 조사 전 '책임 회피 금지' 기조 유지. 드론 순시 등 산불 예방 투자 실적 강조."),
            ("취약계층 에너지 사각지대", "수혜자 미발굴 전수조사 및 '찾아가는 서비스' 실적 홍보."),
        ],
        "action": "CEO 즉시 현장 방문, 48시간 내 재발 방지책 공식 발표. 협력사 안전 지원 예산 공개.",
        "steps": [
            "CEO 현장 방문·48시간 내 재발 방지책 공식 발표",
            "협력사 안전 지원 예산·이행 현황 정량 공개",
            "드론 순시·산불 예방 투자 실적 적극 홍보",
        ],
        "msg": "안전은 타협할 수 없는 가치입니다.",
    },
    "노사관계": {
        "bg": "파업 결의는 공공서비스 안정성 우려로 즉각 확산. MZ세대 이탈·희망퇴직 반발이 내부 리스크로 중첩.",
        "sub_issues": [
            ("파업 결의·단체 행동", "파업 시에도 전력 공급 차질 없음을 최우선 홍보. 노사 대화 채널 상시 브리핑."),
            ("MZ세대 이탈·사기 저하", "CEO 현장 소통 설명회·유연 조직 문화 개선 사례 카드뉴스 제작."),
            ("희망퇴직·인력 감축 반발", "경영 위기 불가피성 설명과 전직 지원 프로그램 보완책 동시 제시."),
            ("채용 비리·불공정 평가 의혹", "AI 채용 시스템·외부 참관인 제도 등 투명성 강화 조치 홍보."),
            ("정치적 외풍 인사 논란", "전문성 중심 인사 원칙 재확인·성과 중심 보직 임명 사례 홍보."),
        ],
        "action": "전력 공급 차질 없음 즉각 선언. 노사 대화 채널 가동 현황 실시간 브리핑으로 루머 선제 차단.",
        "steps": [
            "전력 공급 안정 유지 공식 선언 및 실시간 브리핑",
            "노사 협상 진행 현황 정기 공개(루머 차단)",
            "MZ세대 조직 문화 개선 카드뉴스 SNS 배포",
        ],
        "msg": "세계 최고 수준의 전기품질을 유지하겠습니다.",
    },
    "탄소중립·에너지전환": {
        "bg": "신재생 전환 지연과 계통 불안정 비판이 이행 속도 논란으로 확산. RE100·OECD 비교가 국제 압박 채널로 작동.",
        "sub_issues": [
            ("신재생 전환 지연·RE100 하락", "ESS·계통 안정화 투자 로드맵 및 '에너지 휴게소' 개념 홍보."),
            ("계통 불안정 비판", "변동성 큰 재생에너지를 수용하는 기술(FACTS·ESS) 구체 실적 제시."),
            ("R&D 예산 삭감 우려", "5대 핵심 기술(HVDC·신재생 등) 위주 투자 성과 발표."),
        ],
        "action": "온실가스 감축 실적 정량 공개. ESS·FACTS 등 계통 안정화 기술 투자 로드맵 시각화.",
        "steps": [
            "온실가스 감축 실적·재생에너지 발전량 수치 공개",
            "ESS·FACTS 등 계통 안정화 기술 투자 로드맵 배포",
            "국제 협약 대비 성과 비교 자료 정기 제공",
        ],
        "msg": "재생에너지를 안정적으로 수용하는 것이 기술의 핵심입니다.",
    },
    "원전·수출": {
        "bg": "수출 협상 지연 보도는 국가 신인도와 직결. 저가 수주 덤핑 논란·기술 유출 우려가 복합 리스크로 작동.",
        "sub_issues": [
            ("해외 수출 협상 지연", "원전 수출은 '장기전'임을 강조. 바라카 원전 성공 운영 사례를 레퍼런스로 지속 노출."),
            ("저가 수주 덤핑 논란", "Team Korea 압도적 시공 역량과 경제성 데이터로 정면 반박."),
            ("기술 유출·안전 우려", "국제 인증 현황·안전 기준 구체 자료 즉시 공개."),
        ],
        "action": "바라카 원전 성공 사례 집중 레퍼런스. 계약·협상 현황 정기 업데이트로 불확실성 해소.",
        "steps": [
            "바라카 원전 성공 사례·국제 인증 자료 지속 노출",
            "수출 협상 진행 상황 정기 업데이트 공개",
            "Team Korea 역량·경제성 비교 팩트시트 배포",
        ],
        "msg": "바라카원전의 성공경험을 계속 확대해가겠습니다.",
    },
    "공기업·거버넌스": {
        "bg": "횡령·금품 수수·낙하산 인사·감사원 지적이 복합 발화. 민영화·분할 루머까지 가세하면 신뢰 훼손이 가속.",
        "sub_issues": [
            ("횡령·금품 수수", "'원스트라이크 아웃제' 적용 공표. 전사 윤리 경영 서약 캠페인 전개."),
            ("낙하산 인사 논란", "전문성 중심 인사 원칙 재확인·성과 중심 보직 임명 사례 홍보."),
            ("감사원 지적 보도", "지적 사항 겸허 수용·개선 조치 완료 및 진행 사항 즉시 공개."),
            ("민영화·분할 루머", "공식 입장문으로 '에너지 안보·공익성' 최우선 가치 재확인."),
            ("CEO 발언 왜곡 보도", "발언 전문(Full Script) 신속 공개·오보 대응팀 정정 보도 요청."),
        ],
        "action": "'원스트라이크 아웃제' 공표 및 감사 결과 즉시 자진 공개. 윤리 경영 서약 캠페인 전사 전개.",
        "steps": [
            "원스트라이크 아웃제·윤리 경영 서약 전사 공표",
            "감사 지적 개선 조치 완료 현황 즉시 공개",
            "민영화 루머 공식 입장문·공익성 가치 재확인 배포",
        ],
        "msg": "청렴도와 고객만족도를 함께 끌어올리는 100% 서비스기업으로 거듭나겠습니다.",
    },
    "AI·디지털혁신": {
        "bg": "투자 대비 성과 불명확 시 예산 낭비 비판. AI 일자리 축소 우려가 내부·외부 동시 저항으로 전이.",
        "sub_issues": [
            ("투자 대비 성과 부족", "AI 도입 전후 효율 지표 수치 비교 Before-After 보도자료 배포."),
            ("AI 일자리 축소 우려", "AI는 '대체'가 아닌 '지원' 도구. 고위험 현장 업무 AI 대체로 안전 강화 사례 홍보."),
            ("보안·개인정보 리스크", "보안·개인정보 보호 조치 별도 세부 홍보."),
        ],
        "action": "AI 도입 Before-After 효율 지표 배포. 'AI는 직원 지원 도구' 메시지로 일자리 불안 선제 해소.",
        "steps": [
            "AI 도입 전후 효율 지표 수치 비교 보도자료 배포",
            "고위험 현장 AI 대체 안전 강화 사례 스토리화",
            "보안·개인정보 보호 조치 별도 홍보",
        ],
        "msg": "AI기반 경영혁신으로 업무효율성과 안전성을 높이겠습니다.",
    },
    "고객·서비스": {
        "bg": "AMI 오류·불투명 요금 청구·고객센터 연결 지연이 SNS 확산의 주요 트리거. 소외계층 접근성 문제가 공분으로 연결.",
        "sub_issues": [
            ("AMI 오류·요금 청구 불만", "선제적 오류 보상 체계 구축·맞춤형 서비스 혁신 사례 홍보."),
            ("고객센터 연결 지연", "AI 챗봇 24/7 대응 체계 구축 사실 홍보."),
            ("취약계층 에너지 사각지대", "'찾아가는 서비스' 실적·수혜자 수 수치화 홍보."),
        ],
        "action": "민원 처리 속도·만족도 지표 공개. '찾아가는 에너지 복지 서비스' 실적 수치화·배포.",
        "steps": [
            "민원 처리 속도·만족도 지표 공개",
            "AI 챗봇 24/7 대응 체계 구축 사실 홍보",
            "취약계층 수혜자 수·전수조사 결과 수치화",
        ],
        "msg": "고객이 요청하기 전에 고객의 문제를 해결하겠습니다.",
    },
    "정책·규제": {
        "bg": "경찰 조사·소송은 기관 이미지에 즉각 타격. 정부 정책과의 엇박자 논란은 공기업 존재 가치를 흔드는 이슈로 확산.",
        "sub_issues": [
            ("경찰 조사·소송 리스크", "법무팀 공식 입장 즉시 발표. '적극 협조·투명 소명' 메시지 여론 선점."),
            ("정부 정책과의 엇박자", "정책 실행 기관으로서 충실한 이행 의지 강조. 기업적 효율 대안 조율 중임 시사."),
            ("일방적 정책 추진 비판", "이해관계자 사전 소통 채널 가동 현황 공개."),
        ],
        "action": "법무팀 공식 입장 48시간 내 즉각 발표. '적극 협조·투명 소명' 메시지로 법적 대응보다 여론 우선 선점.",
        "steps": [
            "법무팀 공식 입장 48시간 내 발표",
            "사실 관계 오보 정정 요청 즉각 집행",
            "조사 협조 의지·투명성 강조 메시지 선점",
        ],
        "msg": "국가 에너지 정책 실현을 이끌어가겠습니다.",
    },
}

# ── 위기관리 언론홍보 실행 전략 20선 (문서5 기반) ─────────────────────────
EXECUTION_STRATEGY = {
    "Ⅰ. 커뮤니케이션 거버넌스": [
        ("골든아워(1시간) 대응 워룸 상시화",
         "위기 감지 즉시 홍보처 중심 의사결정 협의체 가동 → 1시간 내 공식 입장(Statement) 발표"),
        ("팩트체크 기반 '원스톱' 승인 프로세스",
         "실무 부서↔홍보처 직통 라인으로 팩트 확인·메시지 승인 단계 최소화, 대응 속도 극대화"),
        ("위기 단계별 다크 사이트(Dark Site) 운영",
         "대형 사고 발생 시 즉각 전환 가능한 위기 전용 홈페이지 사전 설계, 정확한 정보 집중 제공"),
        ("24시간 실시간 여론 모니터링 고도화",
         "AI 기반 뉴스·SNS 키워드 분석으로 위기 징후 선제 포착 및 조기 경보 체계 운영"),
    ],
    "Ⅱ. 전략적 미디어 릴레이션십": [
        ("핵심 현안 '정기 백브리핑' 정례화",
         "요금 현실화·전력망 확충 등 복잡 현안을 출입기자단 대상 심층 설명회로 오보 가능성 차단"),
        ("미디어 파트너십 기반 아웃리치 강화",
         "주요 매체 논설위원·데스크와 정기 간담회 → 에너지 보국 가치 공유"),
        ("가짜뉴스·오보 대응 '팩트 허브' 운영",
         "사실 다른 보도에 증거 데이터와 함께 즉각 반박 자료 게시, 언론 정정 요청 강화"),
        ("미래 기술(AX/Grid) 미디어 데이 개최",
         "'에너지고속도로', 'AI 혁신' 시연·체험 기회 제공으로 긍정 보도 확산 유도"),
    ],
    "Ⅲ. 이해관계자 및 사회적 책임": [
        ("전력망 건설 지역 상생 소통 라운드테이블",
         "송전망 건설 갈등 지역 주민과 정기 소통 채널 구축, 주민 의견 건설 계획에 선제 반영"),
        ("에너지 취약계층 지원 성과 시각화",
         "복지 할인·에너지 바우처 지원 실적을 국민 체감형 인포그래픽으로 제작·배포"),
        ("국민 참여형 '에너지 안보' 캠페인",
         "요금 정상화가 한전 수익이 아닌 '국가 에너지 안보'를 위한 투자임을 알리는 대국민 공감 캠페인"),
        ("지역사회 맞춤형 '에너지보국' 스토리텔링",
         "단순 기부를 넘어 지역 전력 설비 점검·주거 환경 개선 등 업(業) 특성 살린 CSR 홍보"),
    ],
    "Ⅳ. 디지털 및 뉴미디어 소통": [
        ("정책 전문 '숏폼' 콘텐츠 제작",
         "어려운 에너지 정책·요금 체계를 1분 내외 영상으로 제작, MZ세대 소통 접점 확대"),
        ("AI 챗봇 기반 24/7 위기대응 FAQ",
         "위기 시 대량 고객 문의에 AI 활용, 정확하고 일관된 정보 실시간 제공"),
        ("소셜 리스닝 기반 맞춤형 메시지 송출",
         "온라인 논란 키워드 분석 → 직접적 해답을 주는 콘텐츠 배포"),
        ("사내 인플루언서(앰배서더) 육성·활용",
         "현장 직원 목소리로 한전의 노고·진정성 전달하는 친근한 브랜딩 콘텐츠 강화"),
    ],
    "Ⅴ. 조직 내부 및 리더십 소통": [
        ("CEO 주재 전사 위기관리 타운홀 미팅",
         "위기 상황 사장이 직접 직원에게 경영 현황 공유, 주인의식 고취 소통의 장 마련"),
        ("위기 대응 매뉴얼 현장 전파 및 교육",
         "홍보처뿐만 아니라 전 사업소 실무자 정기 모의 훈련 실시"),
        ("글로벌 에너지 솔루션 기업 이미지 제고",
         "원전 수출·해외 사업 성과 집중 홍보 → 내부 구성원 자부심 고취 및 대외 신뢰도 향상"),
        ("사후 평가 시스템(AAR) 도입",
         "위기 종료 후 대응 과정 정밀 분석 백서 발간, 다음 위기 대응 자산으로 축적"),
    ],
}


DEFAULT_INSIGHT = {
    "bg": "커뮤니케이션 공백이 부정 보도의 가장 큰 원인. 침묵은 언론이 대신 채운다.",
    "sub_issues": [
        ("신속 대응 체계 부재", "위기 감지 즉시 홍보처 중심 의사결정 협의체 가동. 1시간 내 공식 입장 발표 체계 구축."),
        ("공식 채널 속도 미흡", "팩트체크 기반 원스톱 승인 프로세스로 대응 속도 극대화."),
    ],
    "action": "해당 이슈 공식 입장 48시간 내 발표 및 담당 부서 창구 일원화.",
    "steps": [
        "해당 이슈 공식 입장 48시간 내 발표",
        "담당 부서 창구 일원화·원스톱 승인 프로세스 가동",
        "미디어 대응 매뉴얼 즉시 실행",
    ],
    "msg": "말하지 않으면 언론이 대신 말한다. 먼저, 빠르게, 구체적으로.",
}

# ── 위기관리 언론홍보 실행전략 20선 (문서5 기반) ──────────────────
CRISIS_EXECUTION_STRATEGY = [
    {
        "group": "Ⅰ. 커뮤니케이션 거버넌스",
        "color": "#B71C1C",
        "icon": "🏛️",
        "items": [
            ("골든아워(1시간) 워룸 상시화", "위기 감지 즉시 홍보처 중심 의사결정 협의체 가동. 1시간 내 공식 Statement 발표 체계 구축."),
            ("팩트체크 기반 원스톱 승인", "실무 부서↔홍보처 직통 라인으로 팩트 확인·메시지 승인 단계 최소화. 대응 속도 극대화."),
            ("위기 단계별 다크 사이트 운영", "대형 사고 발생 시 즉각 전환 가능한 위기 대응 전용 홈페이지 미리 설계."),
            ("24시간 AI 여론 모니터링", "AI 기반 뉴스·SNS 키워드 분석으로 위기 징후 선제 포착. 조기 경보 체계 운영."),
        ],
    },
    {
        "group": "Ⅱ. 전략적 미디어 릴레이션십",
        "color": "#1565C0",
        "icon": "📡",
        "items": [
            ("핵심 현안 정기 백브리핑", "요금 현실화·전력망 확충 등 복잡 현안 출입기자단 심층 설명회. 오보 가능성 선제 차단."),
            ("논설위원·데스크 정기 간담회", "주요 매체 데스크와 정기 간담회로 한전 경영 철학·'에너지 보국' 가치 공유."),
            ("팩트 허브 운영·오보 즉각 반박", "사실과 다른 보도에 증거 데이터와 함께 즉각 반박 자료 게시. 언론 정정 요청 강화."),
            ("AX·Grid 미디어 데이 개최", "에너지고속도로·AI 혁신을 시연·체험하는 행사로 긍정 보도 확산 유도."),
        ],
    },
    {
        "group": "Ⅲ. 이해관계자 및 사회적 책임",
        "color": "#2E7D32",
        "icon": "🤝",
        "items": [
            ("전력망 갈등 지역 소통 라운드테이블", "송전망 건설 갈등 지역 주민들과 정기 소통 채널 구축. 주민 의견 건설 계획에 선제 반영."),
            ("취약계층 지원 성과 인포그래픽 배포", "복지 할인·에너지 바우처 지원 실적을 국민이 체감하는 인포그래픽으로 제작 배포."),
            ("'에너지 안보' 대국민 공감 캠페인", "요금 정상화가 한전 수익이 아닌 국가 에너지 안보를 위한 투자임을 알리는 캠페인 전개."),
            ("지역사회 맞춤형 에너지보국 스토리텔링", "단순 기부를 넘어 지역 전력 설비 점검·주거 환경 개선 등 업 특성 살린 CSR 홍보."),
        ],
    },
    {
        "group": "Ⅳ. 디지털 및 뉴미디어 소통",
        "color": "#6A1B9A",
        "icon": "📱",
        "items": [
            ("정책 전문 숏폼 콘텐츠 제작", "어려운 에너지 정책·요금 체계를 1분 영상으로 제작. MZ세대 소통 접점 확대."),
            ("AI 챗봇 24/7 위기대응 FAQ", "위기 시 대량 고객 문의에 AI를 활용한 정확하고 일관된 정보 실시간 제공."),
            ("소셜 리스닝 기반 맞춤형 메시지 송출", "온라인 논란 키워드 분석 후 직접적 해답을 주는 콘텐츠 배포."),
            ("사내 인플루언서(앰배서더) 육성", "현장 직원의 목소리로 한전 노고와 진정성을 전달하는 친근한 브랜딩 콘텐츠 강화."),
        ],
    },
    {
        "group": "Ⅴ. 조직 내부 및 리더십 소통",
        "color": "#E65100",
        "icon": "🏢",
        "items": [
            ("CEO 전사 위기관리 타운홀 미팅", "위기 상황에서 사장이 직접 직원들에게 경영 현황 공유. 주인의식 고취 소통의 장 마련."),
            ("위기 대응 매뉴얼 현장 전파 교육", "홍보처 외 전 사업소 실무자가 위기 시 행동 요령 숙지. 정기 모의 훈련 실시."),
            ("글로벌 에너지 솔루션 기업 이미지 제고", "원전 수출·해외 사업 성과 집중 홍보. 내부 구성원 자부심 고취 및 대외 신뢰도 향상."),
            ("사후 평가 시스템 도입", "위기 종료 후 대응 과정 정밀 분석·백서 발간. 다음 위기 대응 자산으로 축적."),
        ],
    },
]

def _extract_core_issue(headlines):
    """헤드라인 목록에서 핵심 쟁점 키워드와 이슈 유형을 추출"""
    if not headlines:
        return "", []
    issue_patterns = [
        (r"(수사|조사|압수수색|기소|고발|고소|경찰|검찰)", "수사·법적 리스크"),
        (r"(파업|쟁의|노조|단체협약|갈등|파행)", "노사 갈등"),
        (r"(인상|요금|누진|전기세|가격|부담)", "요금 인상 부담"),
        (r"(사망|부상|사고|재해|화재|폭발|안전)", "안전사고"),
        (r"(적자|손실|부채|재무|적자|결손)", "재무 악화"),
        (r"(지연|취소|중단|차질|연기|무산)", "사업 차질"),
        (r"(비리|부패|횡령|특혜|불법|로비|뇌물)", "비위·도덕성"),
        (r"(오염|누출|탄소|환경|폐기)", "환경 이슈"),
        (r"(불만|민원|항의|반발|피해)", "여론 악화"),
        (r"(낙하산|인사|임명|선임|선발)", "인사 논란"),
    ]
    text = " ".join(headlines)
    found = [label for pat, label in issue_patterns if re.search(pat, text)]
    words = re.findall(r"[가-힣]{2,5}", text)
    freq  = Counter(words)
    stopwords = {"관련","보도","기사","언론","뉴스","대한","위한","으로","에서","에게","부터","까지",
                 "그리고","하지만","그러나","또한","이번","지난","올해","올","이후","이전","등록","발표",
                 "확인","예정","진행","추진","계획","방침","검토","강조","지적","밝혀","통해","따라"}
    keywords = [w for w, c in freq.most_common(12) if c >= 2 and w not in stopwords][:3]
    return ", ".join(keywords) if keywords else "", found


# ── 카테고리별 고정 전략 테이블 (카테고리가 1차 기준) ──────────────
_CAT_STRATEGY = {
    "전기요금": {
        "action_tpl": "{kw} 관련 원가·지원 실적 팩트시트 즉시 배포",
        "msg":    "요금 이슈는 감정이 아닌 숫자로 설득해야 한다. 원가회수율·지원 가구 수 구체 수치가 설득의 무기다.",
        "steps":  ["원가회수율·지원 가구 수 수치화 자료 즉시 배포", "취약계층 지원 성과 스토리 발굴·배포", "핵심 매체 1:1 설명회 개최"],
    },
    "재무·경영": {
        "action_tpl": "{kw} 개선 지표 — 전기 대비 변화폭 중심 선제 공개",
        "msg":    "재무 보도에는 '얼마나 줄었나'를 먼저 말해야 한다. 과거 대비 개선폭이 핵심이다.",
        "steps":  ["부채 감축·효율화 실적 정량 보도자료 발표", "경영 효율화 구체 사례 스토리텔링", "경제지 전담 기자 관계 집중 강화"],
    },
    "노사관계": {
        "action_tpl": "협상 현황 정기 브리핑으로 {kw} 관련 루머 선제 차단",
        "msg":    "파업·갈등 보도는 공공서비스 불안으로 프레임이 전환된다. 침묵은 최악이다.",
        "steps":  ["협상 진행 상황 주 1회 정기 브리핑 실시", "공공서비스 정상 유지 메시지 선점", "노사 공동 성명 또는 협의 시그널 제공"],
    },
    "안전·사고": {
        "action_tpl": "{kw} 관련 사고 원인·재발방지책 48시간 내 공식 발표",
        "msg":    "안전 이슈는 '얼마나 빨리, 얼마나 구체적으로' 대응하느냐가 2차 피해를 막는다.",
        "steps":  ["사고 원인 및 재발방지책 48시간 내 발표", "현장 안전 투자 금액·건수 데이터 동시 제공", "협력업체 포함 안전망 확대 조치 공표"],
    },
    "전력망·설비": {
        "action_tpl": "{kw} 설비 현대화 로드맵·투자 계획 수치 중심 공개",
        "msg":    "'문제 있다'는 언론보다 '우리가 먼저 알고 고치고 있다'는 메시지가 우선이다.",
        "steps":  ["투자 계획·진행 현황 수치 중심 보도자료", "정전 원인 신속 공개 및 복구 타임라인 제시", "스마트그리드·디지털 전환 성과 홍보"],
    },
    "탄소중립·에너지전환": {
        "action_tpl": "{kw} 관련 온실가스 감축 실적·재생에너지 투자 정량 공개",
        "msg":    "추상적 목표보다 '작년 대비 몇% 줄었다'는 구체 숫자가 신뢰를 만든다.",
        "steps":  ["온실가스 감축 실적 정량 공개", "재생에너지 투자·발전량 구체 수치 제시", "국제 협약 대비 성과 비교 자료 제공"],
    },
    "공기업·거버넌스": {
        "action_tpl": "{kw} 관련 감사 결과 자진 공개 및 재발방지 조치 발표",
        "msg":    "비리·특혜 의혹은 숨기면 더 크게 터진다. 먼저 공개하는 것이 신뢰 회복의 시작이다.",
        "steps":  ["감사 결과 자진 공개로 선제 대응", "구체적 윤리경영 조치 언론 제공", "외부 제3자 검증 활용으로 신뢰성 확보"],
    },
    "정책·규제": {
        "action_tpl": "{kw} 관련 법무팀 공식 입장 24시간 내 발표 및 사실관계 선제 공개",
        "msg":    "'법적 대응'보다 '적극 협조·투명하게 소명' 메시지가 여론에 유리하다.",
        "steps":  ["법무팀 공식 입장 즉시 발표", "사실 관계 오보 정정 요청 적극 집행", "조사 협조 의지·투명성 강조 메시지 선점"],
    },
    "원전·수출": {
        "action_tpl": "{kw} 관련 계약·협상 진행 상황 정기 업데이트 공개",
        "msg":    "불확실성이 비판을 낳는다. 원전 수출 관련 알려줄 수 있는 정보는 먼저 알려라.",
        "steps":  ["계약·협상 진행 상황 정기 업데이트", "안전 기준·국제 인증 현황 구체 자료 제공", "기존 수출 성공 사례 집중 레퍼런스화"],
    },
    "AI·디지털혁신": {
        "action_tpl": "{kw} 도입 전후 효율 지표 수치 비교 보도자료 배포",
        "msg":    "'AI 도입'이 아니라 '덕분에 이렇게 달라졌다'는 Before-After 스토리가 효과적이다.",
        "steps":  ["AI 도입 전후 효율 지표 수치 비교 자료", "구체적 서비스 개선 사례(응답시간·오류율) 제시", "보안·개인정보 보호 조치 별도 홍보"],
    },
    "고객·서비스": {
        "action_tpl": "{kw} 관련 민원 처리 현황 공개 및 개선 로드맵 발표",
        "msg":    "민원 통계보다 '실제 해결된 사람의 이야기'가 언론에 더 잘 먹힌다.",
        "steps":  ["민원 처리 속도·만족도 지표 공개", "해결 사례 스토리 발굴 및 배포", "24시간 대응 체계 구축 사실 홍보"],
    },
}

def _build_dynamic_insight(cat, headlines, found_issues, label):
    """카테고리 1차 기준 + 기사 핵심키워드로 문장 보강하는 To-Be 생성"""
    db_base  = INSIGHT_DB.get(cat, DEFAULT_INSIGHT)
    core_kw, _ = _extract_core_issue(headlines)
    kw = core_kw if core_kw else label[:8]   # 핵심 키워드 없으면 이슈 제목 앞부분 사용

    # 카테고리가 테이블에 있으면 무조건 그 전략 사용 (중복 방지의 핵심)
    if cat in _CAT_STRATEGY:
        tpl = _CAT_STRATEGY[cat]
        action = tpl["action_tpl"].replace("{kw}", kw)
        msg    = tpl["msg"]
        steps  = list(tpl["steps"])

        # 헤드라인 패턴으로 추가 보강 (동일 카테고리라도 기사 내용이 다르면 action 앞에 컨텍스트 추가)
        if "수사·법적 리스크" in found_issues and cat not in ("정책·규제", "공기업·거버넌스"):
            action = f"[수사 대응 긴급] {action}"
        elif "안전사고" in found_issues and cat != "안전·사고":
            action = f"[안전 우선] {action}"

    else:
        # 기타 카테고리 — found_issues 로 분기
        if "수사·법적 리스크" in found_issues:
            action = f"'{kw}' 관련 법무팀 공식 입장 24시간 내 발표"
            msg    = f"'{kw}' 보도가 수사·법적 이슈로 확산되기 전, 먼저 입장을 내는 쪽이 여론을 선점한다."
            steps  = ["사실관계 확인 즉시 공식 입장 발표", "언론사별 1:1 팩트 브리핑 실시", "담당 부서 창구 일원화"]
        elif "여론 악화" in found_issues:
            action = f"'{kw}' 관련 민원·불만 처리 현황 공개 및 개선 로드맵 발표"
            msg    = f"민원 통계보다 '실제 해결된 사람의 이야기'가 언론에 더 잘 먹힌다."
            steps  = ["민원 처리 속도·만족도 지표 공개", "해결 사례 스토리 발굴 및 배포", "24시간 대응 체계 구축 사실 홍보"]
        else:
            action = f"'{kw}' 이슈 공식 입장 48시간 내 발표 및 담당 창구 일원화"
            msg    = f"'{kw}' 관련 언론 보도에는 먼저, 빠르게, 구체적으로 대응하는 것이 원칙이다."
            steps  = [f"'{kw}' 관련 공식 입장 즉시 발표", "담당 부서 창구 일원화", "미디어 대응 매뉴얼 사전 준비"]

    return {
        "bg":     db_base.get("bg", ""),
        "action": action,
        "msg":    msg,
        "steps":  steps,
    }


def gen_paired_insights(criticisms):
    """기사 헤드라인 분석 기반 동적 To-Be 인사이트 생성 (카테고리별 중복 없음)"""
    result      = []
    used_cats   = set()   # 동일 카테고리 재사용 방지 (fallback 포함)
    used_actions = set()  # 동일 action 문구 재사용 방지

    for c in criticisms:
        cat       = c.get("category", c["title"])
        headlines = c.get("headlines", [])
        label     = c.get("title", cat)
        _, found_issues = _extract_core_issue(headlines)
        db = _build_dynamic_insight(cat, headlines, found_issues, label)

        # 혹시 같은 action이 나오면 카테고리 DB에서 직접 꺼내 덮어씀
        action_key = db["action"][:20]
        if action_key in used_actions:
            fallback_db = INSIGHT_DB.get(cat, DEFAULT_INSIGHT)
            core_kw, _ = _extract_core_issue(headlines)
            kw = core_kw if core_kw else label[:8]
            db = {
                "bg":     fallback_db.get("bg", ""),
                "action": f"[{cat}] {kw} 관련 {fallback_db.get('action', DEFAULT_INSIGHT['action'])}",
                "msg":    fallback_db.get("msg", DEFAULT_INSIGHT["msg"]),
                "steps":  fallback_db.get("steps", DEFAULT_INSIGHT["steps"]),
            }

        used_cats.add(cat)
        used_actions.add(db["action"][:20])
        result.append({"criticism": c, "db": db})
    return result

def gen_criticisms(arts, kw):
    """실제 부정 기사 헤드라인을 기반으로 이슈 제목·요점을 동적 생성"""
    neg = [a for a in arts if a["감성"] == "부정"]
    cat_c = Counter([a["카테고리"] for a in neg])

    # 카테고리별 고정 제목 (카테고리 분류명만 참조용)
    TITLE_DB = {
        "전기요금": "전기요금 관련 부정 보도",
        "재무·경영": "재무·경영 위기 보도",
        "노사관계": "노사갈등 관련 보도",
        "공기업·거버넌스": "공기업 투명성 논란",
        "안전·사고": "안전사고 관련 보도",
        "전력망·설비": "전력망·설비 문제 보도",
        "탄소중립·에너지전환": "탄소중립 이행 논란",
        "정책·규제": "정책·규제 관련 부정 보도",
        "원전·수출": "원전·수출 신뢰성 논란",
        "AI·디지털혁신": "디지털 혁신 실효성 논란",
        "고객·서비스": "고객서비스 불만 보도",
    }

    result = []
    for cat, cnt2 in cat_c.most_common(8):
        if cat == "기타":
            continue
        # 해당 카테고리 부정 기사 중 최신순 헤드라인 최대 3건 추출
        cat_arts = [a for a in neg if a["카테고리"] == cat]
        cat_arts_sorted = sorted(cat_arts, key=lambda x: x.get("일자", ""), reverse=True)

        # 헤드라인에서 핵심 부정 키워드 등장 문장 우선 선별
        headlines = []
        for a in cat_arts_sorted:
            h = str(a.get("헤드라인", "")).strip()
            if h and h not in headlines:
                headlines.append(h)
            if len(headlines) >= 3:
                break

        # points = 실제 헤드라인 요약 (30자 이내 truncate)
        points = [h[:32] + ("..." if len(h) > 32 else "") for h in headlines[:2]]
        if not points:
            points = ["관련 부정 보도 집중 모니터링 필요"]

        # 이슈 제목: 카테고리 기본 제목 사용 (헤드라인 기반이므로 동일 카테고리라도 매번 다른 기사 반영)
        title = TITLE_DB.get(cat, f"{cat} 비판 보도")

        dots = min(5, max(2, cnt2 // max(1, len(neg) // 10) + 2))
        result.append({"title": title, "points": points, "dots": dots, "category": cat, "headlines": headlines})
        if len(result) == 3:
            break

    # 3개 미만이면 부족분 보충 (실제 기사 기반 기타 이슈)
    if len(result) < 3:
        other_neg = [a for a in neg if a["카테고리"] == "기타"]
        other_sorted = sorted(other_neg, key=lambda x: x.get("일자", ""), reverse=True)
        other_headlines = list({a.get("헤드라인", "") for a in other_sorted if a.get("헤드라인")})[:3]
        while len(result) < 3:
            h_list = other_headlines[:2] if other_headlines else []
            pts = [h[:32] + ("..." if len(h) > 32 else "") for h in h_list] or ["커뮤니케이션 강화 필요"]
            result.append({"title": "기타 부정 보도 동향", "points": pts, "dots": 2, "category": "기타", "headlines": h_list})
            other_headlines = other_headlines[2:]
            if not other_headlines:
                break

    # 그래도 3개 미만이면 placeholder
    fallbacks = [
        {"title": "커뮤니케이션 체계 미흡", "points": ["위기 시 신속 대응 부족", "공식 채널 속도 개선 필요"], "dots": 3, "category": "기타", "headlines": []},
        {"title": "사회적 책임 이행 부족",  "points": ["CSR 기대치 미충족", "이해관계자 소통 강화"], "dots": 2, "category": "기타", "headlines": []},
        {"title": "미디어 관계 강화 필요",  "points": ["전담 기자 관계 구축", "정기 브리핑 채널 부재"], "dots": 2, "category": "기타", "headlines": []},
    ]
    i = 0
    while len(result) < 3 and i < len(fallbacks):
        result.append(fallbacks[i]); i += 1

    return result[:3]

# ── 유틸 ──────────────────────────────────────────────
def clean(t): return re.sub(r'<[^>]+>','',str(t)).strip()

def extract_reporter(title, desc):
    """헤드라인·description에서 기자명 추출. '홍길동 기자' 패턴."""
    text = title + " " + desc
    # 패턴1: [이름] 기자  (2~4글자 한글)
    m = re.search(r'([가-힣]{2,4})\s*기자', text)
    if m:
        name = m.group(1)
        # 오탐 방지: 일반 명사 제거
        noise = {"특별","취재","전문","선임","수석","부장","차장","편집","논설","객원","사진","영상","온라인"}
        if name not in noise:
            return name
    return "—"
def get_media(o,l):
    url=o if o else l
    for d,n in MEDIA_MAP.items():
        if d in url: return n
    # 도메인에서 추출 후 한글 불가능하면 '기타'로
    try:
        domain = url.split("//")[-1].split("/")[0].replace("www.","")
        part = domain.split(".")[0]
        # 숫자나 영문만이면 기타 처리
        if part and not part.isascii():
            return part
        return "기타"
    except: return "기타"

def is_major_media(media):
    """주요매체 여부 (MEDIA_GRADE 등록 = 전국지/방송)"""
    return media in MEDIA_GRADE

def media_sort_key(media):
    """주요매체 우선, 그 다음 rank 순"""
    is_major = 0 if is_major_media(media) else 1
    rank = MEDIA_GRADE.get(media,{}).get("rank", 999)
    return (is_major, rank)

def is_relevant(t): return not any(re.search(p,t) for p in IRRELEVANT_PATTERNS)
def get_sentiment(t):
    p=sum(1 for w in POSITIVE_WORDS if w in t); n=sum(1 for w in NEGATIVE_WORDS if w in t)
    return "긍정" if p>n else "부정" if n>p else "중립"
def summarize(t,n=30):
    t=re.sub(r'\s+',' ',t).strip(); return t[:n]+"..." if len(t)>n else t
def parse_kw(raw):
    raw=raw.replace("(","").replace(")",""); result=[]
    for p in [x.strip() for x in raw.split(",") if x.strip()]:
        if "+" in p:
            sub=[k.strip() for k in p.split("+") if k.strip()]
            result.append({"type":"AND","keywords":sub,"label":" + ".join(sub)})
        else: result.append({"type":"SINGLE","keywords":[p],"label":p})
    return result
def matches_and(t,g): return all(k in t for k in g["keywords"])
def apply_disambig(arts,label):
    for base,req in DISAMBIG_MAP.items():
        if base in label:
            return [a for a in arts if any(r in a["헤드라인"]+" "+a.get("요약","") for r in req)]
    return arts
def get_news(q,mx=1000):
    url="https://openapi.naver.com/v1/search/news.json"
    hdr={"X-Naver-Client-Id":CLIENT_ID,"X-Naver-Client-Secret":CLIENT_SECRET}
    items,s=[],1
    while s<=mx:
        try:
            r=requests.get(url,headers=hdr,params={"query":q,"display":100,"start":s,"sort":"date"},timeout=10)
            batch=r.json().get("items",[])
            if not batch: break
            items.extend(batch)
            if len(batch)<100: break
            s+=100
        except: break
    return items
def auto_cat(arts):
    for a in arts:
        t=a["헤드라인"]+" "+a.get("요약","")
        sc={c:sum(1 for w in ws if w in t) for c,ws in TOPIC_GROUPS.items()}
        sc={k:v for k,v in sc.items() if v>0}
        a["카테고리"]=max(sc,key=sc.get) if sc else "기타"
    return arts
def extract_kws(arts,sent,n=3):
    ft=[a for a in arts if a["감성"]==sent]
    txt=" ".join([a["헤드라인"]+" "+a.get("요약","") for a in ft])
    pool=(NEGATIVE_WORDS if sent=="부정" else POSITIVE_WORDS)
    cnt={w:txt.count(w) for w in pool if txt.count(w)>0}
    return sorted(cnt.items(),key=lambda x:x[1],reverse=True)[:n]
def get_media_rank(media): return MEDIA_GRADE.get(media,{}).get("rank",999)
def sentiment_light(s): return {"긍정":"🟢","부정":"🔴","중립":"🟡"}.get(s,"⚪")
def calc_pr_risk(neg_n,total,neg_kws,crisis_found,top_neg_media):
    s=0; neg_r=neg_n/total*100 if total>0 else 0
    s+=min(40,neg_r*0.8)
    if crisis_found: s+=20
    s+=min(20,len(neg_kws)*4)
    sa=[m for m in top_neg_media if MEDIA_GRADE.get(m,{}).get("grade","") in ["S","A"]]
    s+=min(20,len(sa)*7)
    s=min(100,round(s,1))
    if s>=70: return s,"HIGH","#C62828"
    elif s>=40: return s,"MEDIUM","#E65100"
    return s,"LOW","#2E7D32"

# ══ 구독 알리미 시스템 ══════════════════════════════════

SUBSCRIPTION_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "subscription.json")

def load_sub():
    try:
        with open(SUBSCRIPTION_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except:
        return {
            "enabled": False,
            "sender_email": "",
            "sender_pw": "",
            "recipients": "",
            "send_hour": 6,
            "send_minute": 30,
            "keyword": "한국전력",
            "days": 1,
            "last_sent": "",
        }

def save_sub(cfg):
    try:
        with open(SUBSCRIPTION_FILE, "w", encoding="utf-8") as f:
            json.dump(cfg, f, ensure_ascii=False, indent=2)
        return True
    except:
        return False

def build_email_html(arts, df, label, period_str):
    """분석 결과 전체 레포트를 HTML 이메일로 변환"""
    total = len(df)
    if total == 0:
        return "<p>수집된 기사가 없습니다.</p>"

    cv = df["감성"].value_counts()
    neg_n = int(cv.get("부정", 0)); pos_n = int(cv.get("긍정", 0)); neu_n = int(cv.get("중립", 0))
    neg_rate = neg_n / total * 100; pos_rate = pos_n / total * 100
    tone_color = "#C62828" if neg_n > pos_n*1.5 else "#1565C0" if pos_n > neg_n*1.5 else "#E65100"
    tone_txt   = "부정 우세" if neg_n > pos_n*1.5 else "긍정 우세" if pos_n > neg_n*1.5 else "균형"

    nk = extract_kws(arts, "부정", n=5)
    pk = extract_kws(arts, "긍정", n=5)
    neg_med = [m for m,_ in df[df["감성"]=="부정"]["매체"].value_counts().head(5).items()]
    pr_s, pr_l, pr_c = calc_pr_risk(neg_n, total, nk, False, neg_med)
    pr_l_kr = {"HIGH":"높음","MEDIUM":"보통","LOW":"낮음"}.get(pr_l, pr_l)
    pr_bg   = "#FFEBEE" if pr_s>=70 else "#FFF3E0" if pr_s>=40 else "#E8F5E9"
    now_str = (datetime.utcnow()+timedelta(hours=9)).strftime("%Y년 %m월 %d일 %H:%M")

    # ── 00. 요약 서술 ──
    neg_top1_kw = nk[0][0] if nk else "해당없음"
    tnc = df[df["감성"]=="부정"]["카테고리"].value_counts().index[0] if neg_n>0 else "없음"
    tpc = df[df["감성"]=="긍정"]["카테고리"].value_counts().index[0] if pos_n>0 else "없음"
    summary_txt = (
        f"<b>{period_str}</b> 네이버 기사 전체 <b>{total}건</b>을 전수 분석했습니다. "
        f"기간 내 <b>부정 보도 {neg_rate:.0f}%</b>, 긍정 보도 {pos_rate:.0f}%로 "
        f"{'부정 보도가 많은 위기 국면입니다' if tone_txt=='부정 우세' else '긍정 보도가 우세한 호의적 환경입니다' if tone_txt=='긍정 우세' else '균형 있는 언론 환경이 유지되고 있습니다'}. "
        f"'{neg_top1_kw}' 키워드가 부정 보도의 핵심이며 언론 리스크는 <b>{pr_s}점({pr_l_kr})</b>입니다. "
        f"위기관리 차원에서 선제적 대응과 '{tnc}' 이슈에 대한 공식 입장 발표가 권고됩니다. "
        f"반면 '{tpc}'와 관련한 보도는 긍정적입니다."
    )

    # ── 01. 매체별 논조 테이블 ──
    media_list = sorted(df["매체"].value_counts().head(15).index.tolist(), key=get_media_rank)
    media_rows = ""
    for mname in media_list:
        gi = MEDIA_GRADE.get(mname,{}); grade=gi.get("grade",""); rate=gi.get("rate","")
        gc = GRADE_COLOR.get(grade,"#999")
        n_tot = int(df[df["매체"]==mname].shape[0])
        n_neg = int(df[(df["매체"]==mname)&(df["감성"]=="부정")].shape[0])
        n_pos = int(df[(df["매체"]==mname)&(df["감성"]=="긍정")].shape[0])
        neg_pct = round(n_neg/n_tot*100) if n_tot>0 else 0
        bar_neg = "■"*int(neg_pct/10) + "□"*(10-int(neg_pct/10))
        media_rows += f"""<tr style='border-bottom:1px solid #f0f0f0;'>
          <td style='padding:5px 8px;font-size:12px;'>
            <span style='background:{gc};color:white;padding:1px 4px;border-radius:2px;font-size:9px;font-weight:700;margin-right:4px;'>{grade}</span>{mname}
          </td>
          <td style='padding:5px 8px;font-size:11px;color:#888;text-align:center;'>{rate}%</td>
          <td style='padding:5px 8px;font-size:12px;font-weight:700;color:#C62828;text-align:center;'>{n_neg}</td>
          <td style='padding:5px 8px;font-size:12px;font-weight:700;color:#1565C0;text-align:center;'>{n_pos}</td>
          <td style='padding:5px 8px;font-size:12px;font-weight:700;color:#555;text-align:center;'>{n_tot}</td>
          <td style='padding:5px 8px;font-size:9px;color:#C62828;'>{bar_neg}</td>
        </tr>"""

    # ── 02. 키워드 워드클라우드 (순수 HTML — 이메일 호환) ──
    all_kws = []
    if nk:
        max_neg = nk[0][1] if nk else 1
        for k, v in nk[:12]:
            size = max(13, min(30, int(13 + (v / max_neg) * 18)))
            all_kws.append(f"<span style='font-size:{size}px;font-weight:700;color:#C62828;background:#FFEBEE;padding:4px 10px;border-radius:20px;margin:4px;display:inline-block;line-height:1.4;'>{k}<span style='font-size:9px;opacity:.7;margin-left:3px;'>({v})</span></span>")
    if pk:
        max_pos = pk[0][1] if pk else 1
        for k, v in pk[:8]:
            size = max(12, min(26, int(12 + (v / max_pos) * 15)))
            all_kws.append(f"<span style='font-size:{size}px;font-weight:700;color:#1565C0;background:#E3F2FD;padding:4px 10px;border-radius:20px;margin:4px;display:inline-block;line-height:1.4;'>{k}<span style='font-size:9px;opacity:.7;margin-left:3px;'>({v})</span></span>")
    # 중심어 (검색 키워드)
    center_kw_html = f"<span style='font-size:28px;font-weight:900;color:#003366;background:#EEF2FF;padding:6px 16px;border-radius:24px;margin:4px;display:inline-block;line-height:1.4;'>⚡ {label}</span>"
    wordcloud_html = center_kw_html + " " + " ".join(all_kws)

    neg_kw_html = " ".join([f"<span style='background:#FFEBEE;color:#C62828;padding:3px 8px;border-radius:12px;font-size:11px;font-weight:700;margin:2px;display:inline-block;'>{k}({v})</span>" for k,v in nk[:5]])
    pos_kw_html = " ".join([f"<span style='background:#E3F2FD;color:#1565C0;padding:3px 8px;border-radius:12px;font-size:11px;font-weight:700;margin:2px;display:inline-block;'>{k}({v})</span>" for k,v in pk[:5]])

    # ── 03. 매체×이슈 부정 보도율 ──
    top_m = sorted(df["매체"].value_counts().head(8).index.tolist(), key=get_media_rank)
    top_c = [c for c in TOPIC_GROUPS if c in df["카테고리"].values][:7]
    heat_header = "<th style='padding:5px 6px;background:#003366;color:white;font-size:9px;font-weight:700;'>매체</th>" + \
        "".join([f"<th style='padding:5px 4px;background:#003366;color:white;font-size:9px;font-weight:700;text-align:center;'>{c[:4]}</th>" for c in top_c])
    heat_rows = ""
    for mname in top_m:
        cells = f"<td style='padding:5px 8px;font-size:11px;font-weight:700;'>{mname}</td>"
        for cat in top_c:
            nm = len(df[(df["매체"]==mname)&(df["카테고리"]==cat)])
            nn = len(df[(df["매체"]==mname)&(df["카테고리"]==cat)&(df["감성"]=="부정")])
            pct = round(nn/nm*100) if nm>0 else 0
            bg = "#B71C1C" if pct>=70 else "#FFB74D" if pct>=40 else "#FFF9C4" if pct>0 else "#F9F9F9"
            fg = "white" if pct>=70 else "#333"
            cells += f"<td style='padding:5px 4px;font-size:10px;font-weight:700;text-align:center;background:{bg};color:{fg};'>{pct}%</td>" if nm>0 else "<td style='padding:5px 4px;font-size:10px;color:#ccc;text-align:center;'>—</td>"
        heat_rows += f"<tr style='border-bottom:1px solid #eee;'>{cells}</tr>"

    # ── 04. 비판 포인트 & 대응 전략 ──
    criticisms = gen_criticisms(arts, label)
    paired = gen_paired_insights(criticisms)
    strategy_rows = ""
    for i, item in enumerate(paired, 1):
        c = item["criticism"]; db = item["db"]
        dots_str = "●"*c["dots"] + "○"*(5-c["dots"])
        pts = " / ".join(c["points"][:2])
        strategy_rows += f"""<tr style='border-bottom:1px solid #f0f0f0;'>
          <td style='padding:8px;font-size:12px;font-weight:700;color:#C62828;vertical-align:top;'>이슈{i}. {c['title']}<br><span style='font-size:10px;color:#C62828;letter-spacing:1px;'>{dots_str}</span><br><span style='font-size:10px;color:#888;font-weight:400;'>{pts}</span></td>
          <td style='padding:8px;font-size:11px;color:#333;vertical-align:top;background:#F0F8FF;'>{db['action']}<br><br><span style='font-size:10px;color:#003366;font-weight:700;'>📌 {db['msg']}</span></td>
        </tr>"""

    # ── 05. 전체 기사 목록 ──
    df_sorted = df.copy()
    df_sorted["_major"] = df_sorted["매체"].apply(lambda m: 0 if is_major_media(m) else 1)
    df_sorted["_r"]     = df_sorted["매체"].apply(get_media_rank)
    df_sorted = df_sorted.sort_values(["일자","_major","_r"], ascending=[False,True,True]).reset_index(drop=True)

    sent_icon = {"부정":"🔴","긍정":"🔵","중립":"🟡"}
    article_rows = ""
    show_n = 10
    for idx, row in df_sorted.head(show_n).iterrows():
        gi2 = MEDIA_GRADE.get(row["매체"],{}); grade2=gi2.get("grade","")
        gc2 = GRADE_COLOR.get(grade2,"#999")
        icon = sent_icon.get(row["감성"],"⚪")
        article_rows += f"""<tr style='border-bottom:1px solid #f5f5f5;{"background:#FFF8F8;" if row["감성"]=="부정" else "background:#F8FBFF;" if row["감성"]=="긍정" else ""}'>
          <td style='padding:5px 6px;font-size:10px;color:#aaa;text-align:center;'>{idx+1}</td>
          <td style='padding:5px 6px;font-size:10px;color:#888;white-space:nowrap;'>{row['일자']}</td>
          <td style='padding:5px 6px;font-size:11px;white-space:nowrap;'>
            <span style='background:{gc2};color:white;padding:1px 3px;border-radius:2px;font-size:8px;font-weight:700;margin-right:2px;'>{grade2}</span>{row['매체']}
          </td>
          <td style='padding:5px 6px;font-size:12px;'>
            <a href='{row["링크"]}' style='color:#003366;text-decoration:none;'>{row['헤드라인']}</a>
          </td>
          <td style='padding:5px 6px;font-size:11px;color:#888;'>{row.get('카테고리','—')}</td>
          <td style='padding:5px 6px;font-size:14px;text-align:center;'>{icon}</td>
        </tr>"""

    remain = total - show_n
    more_row = ""
    if remain > 0:
        more_row = f"""<tr>
          <td colspan='6' style='text-align:center;padding:14px;background:#F4F6F9;'>
            <span style='font-size:12px;color:#888;'>10건만 표시 중 &nbsp;|&nbsp; 나머지 <b style='color:#003366;'>{remain}건</b>은 앱에서 확인하세요</span><br>
            <a href='{APP_URL}?kw={requests.utils.quote(label)}&days={days}' target='_blank' style='display:inline-block;margin-top:8px;background:#003366;color:white;padding:7px 20px;border-radius:20px;font-size:12px;font-weight:700;letter-spacing:.3px;text-decoration:none;'>⚡ 전체 기사 {total}건 앱에서 보기 →</a>
          </td>
        </tr>"""

    html = f"""<!DOCTYPE html>
<html><head><meta charset='utf-8'>
<style>
  body {{ font-family:'Malgun Gothic','Apple SD Gothic Neo',Arial,sans-serif; margin:0; padding:0; background:#f0f2f5; }}
  .wrap {{ max-width:800px; margin:20px auto; background:white; border-radius:8px; overflow:hidden; box-shadow:0 2px 12px rgba(0,0,0,.12); }}
  table {{ width:100%; border-collapse:collapse; }}
  .sec {{ margin-bottom:0; padding:18px 24px; border-bottom:1px solid #eee; }}
  .sec-title {{ font-size:14px; font-weight:800; color:#003366; border-left:4px solid #003366; padding-left:10px; margin-bottom:12px; }}
</style></head><body>
<div class='wrap'>

  <!-- 헤더 -->
  <div style='background:#003366;color:white;padding:20px 24px;'>
    <div style='font-size:20px;font-weight:800;'>⚡ {label} 언론보도 유형분석 리포트</div>
    <div style='font-size:11px;opacity:.75;margin-top:5px;'>{period_str} | {now_str} 자동 발송 | 홍보실에 꼭 필요한 뉴스 분석시스템 <span style="font-size:9px;opacity:.7;">by 글쓰는 여행자</span></div>
  </div>

  <!-- KPI 카드 -->
  <div style='padding:18px 24px;border-bottom:1px solid #eee;'>
    <table style='table-layout:fixed;'>
      <tr>
        <td style='text-align:center;padding:10px;background:#F4F6F9;border-radius:6px;border-top:3px solid #003366;'>
          <div style='font-size:24px;font-weight:800;color:#003366;'>{total}</div>
          <div style='font-size:10px;color:#888;margin-top:2px;'>총 기사</div>
        </td>
        <td style='width:8px;'></td>
        <td style='text-align:center;padding:10px;background:#FFF8F8;border-radius:6px;border-top:3px solid #C62828;'>
          <div style='font-size:24px;font-weight:800;color:#C62828;'>{neg_n}</div>
          <div style='font-size:10px;color:#888;margin-top:2px;'>부정 ({neg_rate:.0f}%)</div>
        </td>
        <td style='width:8px;'></td>
        <td style='text-align:center;padding:10px;background:#F0F8FF;border-radius:6px;border-top:3px solid #1565C0;'>
          <div style='font-size:24px;font-weight:800;color:#1565C0;'>{pos_n}</div>
          <div style='font-size:10px;color:#888;margin-top:2px;'>긍정 ({pos_rate:.0f}%)</div>
        </td>
        <td style='width:8px;'></td>
        <td style='text-align:center;padding:10px;background:#F5F5F5;border-radius:6px;border-top:3px solid #888;'>
          <div style='font-size:24px;font-weight:800;color:#555;'>{neu_n}</div>
          <div style='font-size:10px;color:#888;margin-top:2px;'>중립</div>
        </td>
        <td style='width:8px;'></td>
        <td style='text-align:center;padding:10px;background:{pr_bg};border-radius:6px;border-top:3px solid {pr_c};'>
          <div style='font-size:24px;font-weight:800;color:{pr_c};'>{pr_s}점</div>
          <div style='font-size:10px;color:#888;margin-top:2px;'>PR리스크({pr_l_kr})</div>
        </td>
        <td style='width:8px;'></td>
        <td style='text-align:center;padding:10px;background:#F4F6F9;border-radius:6px;border-top:3px solid {tone_color};'>
          <div style='font-size:20px;font-weight:800;color:{tone_color};'>{"🔴" if tone_txt=="부정 우세" else "🔵" if tone_txt=="긍정 우세" else "🟡"}</div>
          <div style='font-size:10px;color:#888;margin-top:2px;'>{tone_txt}</div>
        </td>
      </tr>
    </table>
  </div>

  <!-- 00. 종합 결론 -->
  <div class='sec'>
    <div class='sec-title'>00 · 종합 결론 및 제언</div>
    <div style='font-size:13px;line-height:1.9;color:#333;background:#F8F9FA;padding:14px 16px;border-left:3px solid #003366;border-radius:0 4px 4px 0;'>{summary_txt}</div>
  </div>

  <!-- 01. 키워드 -->
  <div class='sec'>
    <div class='sec-title'>01 · 키워드 워드클라우드</div>
    <div style='text-align:center;padding:16px 8px;background:#FAFBFC;border-radius:6px;line-height:2.2;'>
      {wordcloud_html}
    </div>
    <div style='margin-top:10px;font-size:10px;color:#aaa;'>
      <span style='color:#C62828;font-weight:700;'>■ 빨강</span> 부정 키워드 &nbsp;
      <span style='color:#1565C0;font-weight:700;'>■ 파랑</span> 긍정 키워드 &nbsp;
      <span style='color:#003366;font-weight:700;'>■ 남색</span> 검색 키워드 &nbsp;
      글자 크기 = 언급 빈도
    </div>
  </div>

  <!-- 02. 매체별 논조 -->\n  <div class='sec'>
    <div class='sec-title'>02 · 매체별 논조 분석</div>
    <table>
      <thead><tr style='background:#003366;color:white;font-size:10px;'>
        <th style='padding:6px 8px;text-align:left;'>매체</th>
        <th style='padding:6px 8px;text-align:center;'>열독률</th>
        <th style='padding:6px 8px;text-align:center;'>부정</th>
        <th style='padding:6px 8px;text-align:center;'>긍정</th>
        <th style='padding:6px 8px;text-align:center;'>전체</th>
        <th style='padding:6px 8px;text-align:left;'>부정비중</th>
      </tr></thead>
      <tbody>{media_rows}</tbody>
    </table>
  </div>

  <!-- 03. 매체×이슈 부정 보도율 -->
  <div class='sec'>
    <div class='sec-title'>03 · 매체×이슈 부정 보도율</div>
    <table>
      <thead><tr>{heat_header}</tr></thead>
      <tbody>{heat_rows}</tbody>
    </table>
    <div style='font-size:10px;color:#aaa;margin-top:6px;'>■ 진한 빨강(70%↑) ■ 주황(40~70%) ■ 노랑(1~40%) □ 없음</div>
  </div>

  <!-- 04. 비판 포인트 & 대응 전략 -->
  <div class='sec'>
    <div class='sec-title'>04 · 비판 포인트 & 대응 전략</div>
    <table>
      <thead><tr style='background:#003366;color:white;font-size:10px;'>
        <th style='padding:6px 8px;text-align:left;width:40%;'>🔴 현재 문제점 (As-Is)</th>
        <th style='padding:6px 8px;text-align:left;width:60%;'>✅ 개선 방향 (To-Be)</th>
      </tr></thead>
      <tbody>{strategy_rows}</tbody>
    </table>
  </div>

  <!-- 05. 전체 기사 목록 -->
  <div class='sec'>
    <div class='sec-title'>05 · 기사 목록 (상위 {min(show_n, total)}건 / 전체 {total}건)</div>
    <table>
      <thead><tr style='background:#003366;color:white;font-size:10px;'>
        <th style='padding:5px 6px;text-align:center;'>No.</th>
        <th style='padding:5px 6px;'>일자</th>
        <th style='padding:5px 6px;'>매체</th>
        <th style='padding:5px 6px;'>헤드라인</th>
        <th style='padding:5px 6px;'>카테고리</th>
        <th style='padding:5px 6px;text-align:center;'>논조</th>
      </tr></thead>
      <tbody>{article_rows}{more_row}</tbody>
    </table>
  </div>

  <!-- 푸터 -->
  <div style='background:#f8f8f8;padding:14px 24px;font-size:10px;color:#aaa;text-align:center;border-top:1px solid #eee;'>
    ⚡ 홍보실에 꼭 필요한 뉴스 분석시스템 &nbsp;<span style="font-size:9px;opacity:.7;">by 글쓰는 여행자</span> &nbsp;|&nbsp; 네이버 뉴스 API 기반 자동 분석 &nbsp;|&nbsp; 열독률: 언론진흥재단('23)<br>
    본 메일은 구독 설정에 따라 자동 발송되었습니다. 수신을 원하지 않으면 앱에서 구독을 해제해 주세요.
  </div>
</div>
</body></html>"""
    return html



    total = len(df)
    if total == 0:
        return "<p>수집된 기사가 없습니다.</p>"
    cv = df["감성"].value_counts()
    neg_n = int(cv.get("부정", 0))
    pos_n = int(cv.get("긍정", 0))
    neu_n = int(cv.get("중립", 0))
    neg_rate = neg_n / total * 100
    pos_rate = pos_n / total * 100

    # PR 리스크
    nk = extract_kws(df.to_dict("records"), "부정", n=5)
    neg_med = [m for m, _ in df[df["감성"] == "부정"]["매체"].value_counts().head(5).items()]
    pr_s, pr_l, pr_c = calc_pr_risk(neg_n, total, nk, False, neg_med)
    pr_l_kr = {"HIGH": "높음", "MEDIUM": "보통", "LOW": "낮음"}.get(pr_l, pr_l)

    # TOP 부정 기사 5건
    neg_top = df[df["감성"] == "부정"].sort_values("일자", ascending=False).head(5)
    neg_rows = ""
    for _, row in neg_top.iterrows():
        gi = MEDIA_GRADE.get(row["매체"], {}); grade = gi.get("grade", "")
        gc = GRADE_COLOR.get(grade, "#999")
        neg_rows += f"""<tr>
          <td style='padding:6px 8px;font-size:12px;color:#333;border-bottom:1px solid #f0f0f0;'>
            <span style='background:{gc};color:white;padding:1px 5px;border-radius:3px;font-size:10px;font-weight:700;margin-right:4px;'>{grade}</span>
            <b style='color:#777;'>{row['매체']}</b>
          </td>
          <td style='padding:6px 8px;font-size:12px;color:#333;border-bottom:1px solid #f0f0f0;'>
            <a href='{row["링크"]}' style='color:#003366;text-decoration:none;'>{row['헤드라인']}</a>
          </td>
          <td style='padding:6px 8px;font-size:11px;color:#999;border-bottom:1px solid #f0f0f0;white-space:nowrap;'>{row['일자']}</td>
        </tr>"""

    # TOP 긍정 기사 3건
    pos_top = df[df["감성"] == "긍정"].sort_values("일자", ascending=False).head(3)
    pos_rows = ""
    for _, row in pos_top.iterrows():
        pos_rows += f"""<tr>
          <td style='padding:6px 8px;font-size:12px;border-bottom:1px solid #f0f0f0;'>
            <b style='color:#777;'>{row['매체']}</b>
          </td>
          <td style='padding:6px 8px;font-size:12px;border-bottom:1px solid #f0f0f0;'>
            <a href='{row["링크"]}' style='color:#1565C0;text-decoration:none;'>{row['헤드라인']}</a>
          </td>
          <td style='padding:6px 8px;font-size:11px;color:#999;border-bottom:1px solid #f0f0f0;white-space:nowrap;'>{row['일자']}</td>
        </tr>"""

    # 키워드
    neg_kw_html = " ".join([f"<span style='background:#FFEBEE;color:#C62828;padding:3px 8px;border-radius:12px;font-size:11px;font-weight:700;margin:2px;display:inline-block;'>{k}({v})</span>" for k, v in nk[:5]])
    pos_kws = extract_kws(df.to_dict("records"), "긍정", n=3)
    pos_kw_html = " ".join([f"<span style='background:#E3F2FD;color:#1565C0;padding:3px 8px;border-radius:12px;font-size:11px;font-weight:700;margin:2px;display:inline-block;'>{k}({v})</span>" for k, v in pos_kws[:3]])

    # 카테고리 TOP 이슈
    cat_neg = df[df["감성"] == "부정"]["카테고리"].value_counts().head(3)
    cat_rows = "".join([f"<li style='margin-bottom:4px;font-size:12px;'><b>{c}</b> — 부정 {n}건</li>" for c, n in cat_neg.items()])

    tone_color = "#C62828" if neg_n > pos_n * 1.5 else "#1565C0" if pos_n > neg_n * 1.5 else "#E65100"
    tone_txt = "부정 우세" if neg_n > pos_n * 1.5 else "긍정 우세" if pos_n > neg_n * 1.5 else "균형"
    pr_bg = "#FFEBEE" if pr_s >= 70 else "#FFF3E0" if pr_s >= 40 else "#E8F5E9"

    now_str = (datetime.utcnow()+timedelta(hours=9)).strftime("%Y년 %m월 %d일 %H:%M")

    html = f"""<!DOCTYPE html>
<html><head><meta charset='utf-8'>
<style>
  body {{ font-family: 'Malgun Gothic', 'Apple SD Gothic Neo', Arial, sans-serif; margin:0; padding:0; background:#f5f5f5; }}
  .wrap {{ max-width:680px; margin:20px auto; background:white; border-radius:8px; overflow:hidden; box-shadow:0 2px 8px rgba(0,0,0,.1); }}
  .hdr  {{ background:#003366; color:white; padding:18px 24px; }}
  .body {{ padding:20px 24px; }}
  .kpi  {{ display:flex; gap:10px; margin-bottom:16px; flex-wrap:wrap; }}
  .kpi-box {{ flex:1; min-width:90px; background:#F4F6F9; border-radius:6px; padding:10px 12px; text-align:center; border-top:3px solid {tone_color}; }}
  table {{ width:100%; border-collapse:collapse; }}
  .sec  {{ margin-bottom:20px; }}
  .sec-title {{ font-size:13px; font-weight:800; color:#003366; border-bottom:2px solid #003366; padding-bottom:5px; margin-bottom:10px; }}
  .ftr  {{ background:#f8f8f8; padding:12px 24px; font-size:10px; color:#aaa; text-align:center; border-top:1px solid #eee; }}
</style></head><body>
<div class='wrap'>
  <div class='hdr'>
    <div style='font-size:18px;font-weight:800;'>⚡ {label} 뉴스 모니터링 리포트</div>
    <div style='font-size:11px;opacity:.75;margin-top:4px;'>{period_str} | {now_str} 발송 | 홍보실에 꼭 필요한 뉴스 분석시스템 <span style="font-size:9px;opacity:.7;">by 글쓰는 여행자</span></div>
  </div>
  <div class='body'>

    <!-- KPI -->
    <div class='kpi'>
      <div class='kpi-box'>
        <div style='font-size:22px;font-weight:800;color:#003366;'>{total}</div>
        <div style='font-size:10px;color:#888;margin-top:2px;'>총 기사</div>
      </div>
      <div class='kpi-box' style='border-top-color:#C62828;'>
        <div style='font-size:22px;font-weight:800;color:#C62828;'>{neg_n}</div>
        <div style='font-size:10px;color:#888;margin-top:2px;'>부정 ({neg_rate:.0f}%)</div>
      </div>
      <div class='kpi-box' style='border-top-color:#1565C0;'>
        <div style='font-size:22px;font-weight:800;color:#1565C0;'>{pos_n}</div>
        <div style='font-size:10px;color:#888;margin-top:2px;'>긍정 ({pos_rate:.0f}%)</div>
      </div>
      <div class='kpi-box' style='border-top-color:#888;'>
        <div style='font-size:22px;font-weight:800;color:#555;'>{neu_n}</div>
        <div style='font-size:10px;color:#888;margin-top:2px;'>중립</div>
      </div>
      <div class='kpi-box' style='border-top-color:{pr_c};background:{pr_bg};'>
        <div style='font-size:22px;font-weight:800;color:{pr_c};'>{pr_s}점</div>
        <div style='font-size:10px;color:#888;margin-top:2px;'>PR리스크 ({pr_l_kr})</div>
      </div>
    </div>

    <!-- 키워드 -->
    <div class='sec'>
      <div class='sec-title'>🔑 주요 키워드</div>
      <div style='margin-bottom:6px;'><b style='font-size:11px;color:#C62828;'>부정 키워드</b><br>{neg_kw_html if neg_kw_html else '<span style="color:#aaa;font-size:11px;">없음</span>'}</div>
      <div><b style='font-size:11px;color:#1565C0;'>긍정 키워드</b><br>{pos_kw_html if pos_kw_html else '<span style="color:#aaa;font-size:11px;">없음</span>'}</div>
    </div>

    <!-- 주요 이슈 -->
    <div class='sec'>
      <div class='sec-title'>📌 주요 비판 이슈 TOP3</div>
      <ul style='margin:0;padding-left:18px;color:#333;'>{cat_rows if cat_rows else '<li>없음</li>'}</ul>
    </div>

    <!-- 부정 기사 -->
    <div class='sec'>
      <div class='sec-title'>🔴 주요 부정 기사 TOP5</div>
      {'<table>' + neg_rows + '</table>' if neg_rows else '<p style="color:#aaa;font-size:12px;">부정 기사 없음</p>'}
    </div>

    <!-- 긍정 기사 -->
    <div class='sec'>
      <div class='sec-title'>🔵 주요 긍정 기사 TOP3</div>
      {'<table>' + pos_rows + '</table>' if pos_rows else '<p style="color:#aaa;font-size:12px;">긍정 기사 없음</p>'}
    </div>

  </div>
  <div class='ftr'>
    ⚡ 홍보실에 꼭 필요한 뉴스 분석시스템 &nbsp;<span style="font-size:9px;opacity:.7;">by 글쓰는 여행자</span> &nbsp;|&nbsp; 네이버 뉴스 API 기반 자동 분석 &nbsp;|&nbsp; 열독률: 언론진흥재단('23)<br>
    본 메일은 구독 설정에 따라 자동 발송되었습니다. 수신을 원하지 않으면 앱에서 구독을 해제해 주세요.
  </div>
</div>
</body></html>"""
    return html


def _collect_news_for(label, days):
    """특정 키워드·기간으로 기사 수집 후 DataFrame 반환"""
    end_dt   = datetime.now().date()
    start_dt = end_dt - timedelta(days=max(1, int(days)))
    raw = get_news(label, 1000)
    arts = []
    for a in raw:
        pub = a.get("pubDate", "")
        try:
            ad = datetime.strptime(pub[:16], "%a, %d %b %Y").date()
            if not (start_dt <= ad <= end_dt): continue
            ds = ad.strftime("%Y-%m-%d"); hs = pub[17:19] if len(pub) > 18 else "00"
        except:
            ds = pub[:10]; hs = "00"
        title = clean(a.get("title", "")); desc = clean(a.get("description", ""))
        text  = title + " " + desc
        orig  = a.get("originallink", ""); link = a.get("link", "")
        if not is_relevant(text): continue
        media = get_media(orig, link); gi = MEDIA_GRADE.get(media, {})
        reporter = extract_reporter(title, desc)
        arts.append({"키워드그룹": label, "일자": ds, "월": ds[:7], "시간": hs,
                     "매체": media, "등급": gi.get("grade","—"), "열독률": gi.get("rate", 0.05),
                     "헤드라인": title, "요약": summarize(desc, 30),
                     "감성": get_sentiment(text), "카테고리": "",
                     "기자": reporter, "링크": orig if orig else link})
    if not arts:
        return None, None, None
    arts = auto_cat(arts)
    df   = pd.DataFrame(arts)
    period_str = f"{start_dt.strftime('%Y.%m.%d')} ~ {end_dt.strftime('%m.%d')}"
    return arts, df, period_str


def send_email_report(cfg, test_addr=None, is_broadcast_test=False, custom_message=None):
    """구독자별 개인 키워드로 리포트 발송. test_addr 지정 시 단일 테스트 발송.
    is_broadcast_test=True 이면 제목 앞에 (테스트) 말머리 추가.
    custom_message 있으면 이메일 HTML 상단에 삽입."""
    try:
        subs = cfg.get("subscribers", [])
        if not subs and not test_addr:
            return False, "구독자 없음"

        targets = subs if not test_addr else [
            next((s for s in subs if s["email"] == test_addr),
                 {"email": test_addr, "keyword": "한국전력",
                  "send_hour": 6, "send_minute": 30})
        ]

        fail_list = []
        with smtplib.SMTP_SSL("smtp.naver.com", 465) as server:
            server.login(cfg["sender_email"], cfg["sender_pw"])
            for sub in targets:
                label = sub.get("keyword", "한국전력")
                days  = sub.get("days", cfg.get("days", 1))  # 구독자별 설정 우선
                arts, df, period_str = _collect_news_for(label, days)
                if arts is None:
                    fail_list.append(f"{sub['email']}(기사 없음)")
                    continue
                try:
                    html_body = build_email_html(arts, df, label, period_str)
                    if custom_message and custom_message.strip():
                        msg_banner = (
                            f"<div style='background:#FFF8E1;border-left:4px solid #F9A825;"
                            f"padding:12px 16px;margin-bottom:16px;font-family:\"Noto Sans KR\",sans-serif;"
                            f"font-size:13px;color:#333;border-radius:0 6px 6px 0;'>"
                            f"📝 <b>발신자 메시지:</b><br>{custom_message.strip().replace(chr(10), '<br>')}"
                            f"</div>"
                        )
                        html_body = msg_banner + html_body
                    addr      = sub["email"]
                    prefix    = "(테스트) " if is_broadcast_test else ""
                    today_str = (datetime.utcnow()+timedelta(hours=9)).strftime('%Y.%m.%d')
                    subject   = (f"{prefix}[({today_str}) 글쓰는 여행자의 뉴스 모니터링 레포트] - {label}")
                    msg = MIMEMultipart("alternative")
                    msg["Subject"] = subject
                    msg["From"]    = cfg["sender_email"]
                    msg["To"]      = addr
                    msg.attach(MIMEText(html_body, "html", "utf-8"))
                    server.sendmail(cfg["sender_email"], [addr], msg.as_string())
                except Exception as e:
                    fail_list.append(f"{sub['email']}({e})")

        cfg["last_sent"] = (datetime.utcnow()+timedelta(hours=9)).strftime("%Y-%m-%d %H:%M")
        save_sub(cfg)

        if fail_list:
            return False, f"일부 실패: {', '.join(fail_list)}"
        return True, f"{len(targets)}명에게 발송 완료"

    except Exception as e:
        return False, str(e)


def init_scheduler():
    """APScheduler 싱글톤 초기화"""
    if not SCHEDULER_OK:
        return None
    key = "_kepco_scheduler"
    if not hasattr(st, key):
        sched = BackgroundScheduler(timezone="Asia/Seoul")
        sched.start()
        setattr(st, key, sched)
    return getattr(st, key)


def apply_scheduler(cfg):
    """구독자별 발송 시각에 맞게 스케줄 잡 등록/갱신"""
    sched = init_scheduler()
    if sched is None:
        return
    # 기존 kepco 잡 전부 제거
    for job in sched.get_jobs():
        if job.id.startswith("kepco_sub_"):
            sched.remove_job(job.id)
    if not cfg.get("enabled"):
        return
    # 발송 시각별로 그룹핑
    from collections import defaultdict
    time_groups = defaultdict(list)
    for sub in cfg.get("subscribers", []):
        h = int(sub.get("send_hour", 6))
        m = int(sub.get("send_minute", 30))
        time_groups[(h, m)].append(sub)
    for (h, m), group in time_groups.items():
        job_id = f"kepco_sub_{h:02d}{m:02d}"
        def make_fn(g):
            def fn():
                c = load_sub()
                try:
                    with smtplib.SMTP_SSL("smtp.naver.com", 465) as srv:
                        srv.login(c["sender_email"], c["sender_pw"])
                        for sub in g:
                            label = sub.get("keyword","한국전력")
                            arts, df, period_str = _collect_news_for(label, c.get("days",1))
                            if arts is None: continue
                            html_body = build_email_html(arts, df, label, period_str)
                            addr = sub["email"]
                            subject = f"[KEPCO 뉴스] {label} 모니터링 리포트 — {(datetime.utcnow()+timedelta(hours=9)).strftime('%Y.%m.%d')}"
                            msg = MIMEMultipart("alternative")
                            msg["Subject"]=subject; msg["From"]=c["sender_email"]; msg["To"]=addr
                            msg.attach(MIMEText(html_body,"html","utf-8"))
                            srv.sendmail(c["sender_email"],[addr],msg.as_string())
                    c["last_sent"] = (datetime.utcnow()+timedelta(hours=9)).strftime("%Y-%m-%d %H:%M")
                    save_sub(c)
                except Exception: pass
            return fn
        sched.add_job(
            make_fn(group),
            CronTrigger(hour=h, minute=m, timezone="Asia/Seoul"),
            id=job_id, replace_existing=True, misfire_grace_time=300
        )


# ── 시장 데이터 ───────────────────────────────────────
# ── 잘 알려진 상장사 → 티커 빠른 매핑 (네트워크 없이도 동작) ──
_KNOWN_TICKERS = {
    "한국전력": ("한국전력", "015760.KS"),
    "한전": ("한국전력", "015760.KS"),
    "삼성전자": ("삼성전자", "005930.KS"),
    "SK하이닉스": ("SK하이닉스", "000660.KS"),
    "현대차": ("현대자동차", "005380.KS"),
    "현대자동차": ("현대자동차", "005380.KS"),
    "LG에너지솔루션": ("LG에너지솔루션", "373220.KS"),
    "삼성바이오로직스": ("삼성바이오로직스", "207940.KS"),
    "셀트리온": ("셀트리온", "068270.KS"),
    "POSCO홀딩스": ("POSCO홀딩스", "005490.KS"),
    "포스코": ("POSCO홀딩스", "005490.KS"),
    "LG화학": ("LG화학", "051910.KS"),
    "KB금융": ("KB금융", "105560.KS"),
    "신한지주": ("신한지주", "055550.KS"),
    "하나금융지주": ("하나금융지주", "086790.KS"),
    "기아": ("기아", "000270.KS"),
    "카카오": ("카카오", "035720.KS"),
    "네이버": ("NAVER", "035420.KS"),
    "NAVER": ("NAVER", "035420.KS"),
    "현대모비스": ("현대모비스", "012330.KS"),
    "삼성SDI": ("삼성SDI", "006400.KS"),
    "LG전자": ("LG전자", "066570.KS"),
    "고려아연": ("고려아연", "010130.KS"),
    "삼성물산": ("삼성물산", "028260.KS"),
    "두산에너빌리티": ("두산에너빌리티", "034020.KS"),
    "한국가스공사": ("한국가스공사", "036460.KS"),
    "한국수력원자력": ("한국수력원자력", ""),  # 비상장
    "카카오뱅크": ("카카오뱅크", "323410.KS"),
    "크래프톤": ("크래프톤", "259960.KS"),
    "한국항공우주": ("한국항공우주", "047810.KS"),
    "에코프로비엠": ("에코프로비엠", "247540.KQ"),
    "에코프로": ("에코프로", "086520.KQ"),
    "엔씨소프트": ("엔씨소프트", "036570.KS"),
    "넷마블": ("넷마블", "251270.KS"),
    "카카오게임즈": ("카카오게임즈", "293490.KQ"),
    "HMM": ("HMM", "011200.KS"),
    "한진칼": ("한진칼", "180640.KS"),
    "대한항공": ("대한항공", "003490.KS"),
    "롯데케미칼": ("롯데케미칼", "011170.KS"),
    "SK이노베이션": ("SK이노베이션", "096770.KS"),
    "SKT": ("SK텔레콤", "017670.KS"),
    "SK텔레콤": ("SK텔레콤", "017670.KS"),
    "KT": ("KT", "030200.KS"),
    "LGU+": ("LG유플러스", "032640.KS"),
    "LG유플러스": ("LG유플러스", "032640.KS"),
    "NH투자증권": ("NH투자증권", "005940.KS"),
    "미래에셋증권": ("미래에셋증권", "006800.KS"),
    "삼성증권": ("삼성증권", "016360.KS"),
    "하이브": ("하이브", "352820.KS"),
    "CJ ENM": ("CJ ENM", "035760.KQ"),
    "한화에어로스페이스": ("한화에어로스페이스", "012450.KS"),
    "한화솔루션": ("한화솔루션", "009830.KS"),
}

@st.cache_data(ttl=86400)
def lookup_krx_ticker(company_name):
    """회사명으로 KRX 티커 자동검색.
    1순위: 내장 매핑 테이블 (네트워크 불필요)
    2순위: 네이버 증권 자동완성 API
    3순위: KRX KIND API"""
    if not company_name:
        return "", ""
    q = company_name.strip()

    # ── 1순위: 내장 매핑 ──
    if q in _KNOWN_TICKERS:
        name, ticker = _KNOWN_TICKERS[q]
        return (name, ticker) if ticker else (name, "")

    # ── 2순위: 네이버 증권 자동완성 ──
    try:
        r = requests.get(
            "https://ac.finance.naver.com/ac",
            params={"q": q, "q_enc": "UTF-8", "st": "111", "frm": "nasdaq",
                    "le": "100", "r_format": "json", "r_enc": "UTF-8"},
            headers={"User-Agent": "Mozilla/5.0", "Referer": "https://finance.naver.com/"},
            timeout=4
        )
        data = r.json()
        items = data.get("items", [[]])[0]   # 첫 번째 그룹
        for item in items:
            # item: [종목코드, 종목명, 시장구분, ...]
            if len(item) >= 2:
                code = str(item[0]).strip().zfill(6)
                name = str(item[1]).strip()
                mkt  = str(item[2]).strip() if len(item) > 2 else ""
                if q in name or name in q:
                    suffix = ".KQ" if "KOSDAQ" in mkt.upper() or "코스닥" in mkt else ".KS"
                    return name, code + suffix
    except: pass

    # ── 3순위: KRX KIND API ──
    try:
        r = requests.get(
            "https://kind.krx.co.kr/common/searchcorpname.do",
            params={"method": "searchCorpNameJson", "searchCorpName": q,
                    "copyPageSize": "10", "currentPageSize": "10"},
            headers={"User-Agent": "Mozilla/5.0", "Referer": "https://kind.krx.co.kr/"},
            timeout=5
        )
        data = r.json()
        rows = data.get("result", data.get("list", []))
        if not rows and isinstance(data, list):
            rows = data
        if rows:
            exact = [x for x in rows if str(x.get("corpNm","")).strip() == q]
            row   = exact[0] if exact else rows[0]
            code  = str(row.get("stockCode", row.get("isu_cd",""))).strip().zfill(6)
            name  = str(row.get("corpNm", q)).strip()
            mkt   = str(row.get("marketName", row.get("mkt",""))).strip()
            suffix = ".KQ" if "코스닥" in mkt or "KOSDAQ" in mkt.upper() else ".KS"
            return name, code + suffix if code else ""
    except: pass

    return q, ""


@st.cache_data(ttl=3600)
def get_weekly_pres_schedule():
    """대통령실 캘린더(https://www.president.go.kr/president/calendar)에서
    오늘 이후 향후 일정 2~3건을 수집해 1줄 문자열로 반환.
    파싱 실패 시 빈 문자열 반환."""
    try:
        from datetime import date as _date
        kst_now  = datetime.utcnow() + timedelta(hours=9)
        today    = kst_now.date()

        # ── 시도 1: 대통령실 캘린더 페이지 직접 파싱 ──
        # 캘린더 페이지가 FullCalendar 기반 SPA인 경우 API 엔드포인트로 요청
        entries = []
        api_urls = [
            # 공식 JSON API 후보 (실제 배포 환경에서 동작)
            f"https://www.president.go.kr/api/schedule/list?startDate={today.strftime('%Y-%m-01')}&endDate={(today + timedelta(days=60)).strftime('%Y-%m-%d')}",
            "https://www.president.go.kr/president/calendar/getScheduleList.do",
        ]
        for api_url in api_urls:
            try:
                resp = requests.get(api_url,
                    headers={"User-Agent":"Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
                             "Referer":"https://www.president.go.kr/president/calendar",
                             "Accept":"application/json, text/html, */*"},
                    timeout=5)
                if resp.status_code != 200:
                    continue
                ct = resp.headers.get("Content-Type","")
                if "json" in ct:
                    data = resp.json()
                    # 일반적 응답 구조 탐색
                    rows = (data.get("data") or data.get("list") or
                            data.get("result") or data.get("schedules") or
                            (data if isinstance(data, list) else []))
                    for row in rows:
                        title   = str(row.get("title") or row.get("subjectNm") or row.get("scheduleNm","")).strip()
                        date_s  = str(row.get("startDate") or row.get("scheduleDate") or row.get("date",""))[:10]
                        if not title or not date_s:
                            continue
                        try:
                            ev_date = datetime.strptime(date_s, "%Y-%m-%d").date()
                        except:
                            continue
                        if ev_date >= today:
                            entries.append((ev_date, title))
                elif "html" in ct:
                    # HTML 내 script 태그에서 JSON 추출 시도
                    json_blobs = re.findall(r'scheduleData\s*=\s*(\[.*?\]);', resp.text, re.DOTALL)
                    for blob in json_blobs:
                        try:
                            rows = __import__('json').loads(blob)
                            for row in rows:
                                title  = str(row.get("title","")).strip()
                                date_s = str(row.get("start") or row.get("date",""))[:10]
                                if title and date_s:
                                    try:
                                        ev_date = datetime.strptime(date_s, "%Y-%m-%d").date()
                                        if ev_date >= today:
                                            entries.append((ev_date, title))
                                    except: pass
                        except: pass
                if entries:
                    break
            except: continue

        # ── 시도 2: 캘린더 HTML 직접 파싱 ──
        if not entries:
            try:
                resp = requests.get("https://www.president.go.kr/president/calendar",
                    headers={"User-Agent":"Mozilla/5.0 (Windows NT 10.0; Win64; x64)",
                             "Accept":"text/html,application/xhtml+xml"},
                    timeout=6)
                if resp.status_code == 200:
                    html = resp.text
                    # data-date 속성에서 날짜, 인접 텍스트에서 제목 추출
                    date_pat = re.compile(r'data-date=["\']([0-9]{4}-[0-9]{2}-[0-9]{2})["\']')
                    tag_pat  = re.compile(r'<[^>]+>')
                    chunks = re.split(r'(?=data-date=)', html)
                    for chunk in chunks[:30]:
                        dm = date_pat.search(chunk)
                        if not dm: continue
                        date_s = dm.group(1)
                        try:
                            ev_date = datetime.strptime(date_s, "%Y-%m-%d").date()
                        except: continue
                        if ev_date < today: continue
                        # 이후 100자에서 태그 제거 후 텍스트 추출
                        raw = tag_pat.sub("", chunk[:200]).strip()
                        title = " ".join(raw.split())[:30]
                        if title and len(title) > 2:
                            entries.append((ev_date, title))
                    # 월일 패턴 (예: 3월 7일 국무회의)
                    for m2 in re.finditer(r'([0-9]{1,2})월\s*([0-9]{1,2})일[^<]{0,3}<[^>]+>([^<]{3,30})', html):
                        try:
                            ev_date = datetime(today.year, int(m2.group(1)), int(m2.group(2))).date()
                            if ev_date >= today:
                                title = m2.group(3).strip()[:25]
                                if title: entries.append((ev_date, title))
                        except: pass
            except: pass

        # ── 시도 3: 네이버 뉴스 API fallback ──
        if not entries:
            try:
                resp = requests.get("https://openapi.naver.com/v1/search/news.json",
                    headers={"X-Naver-Client-Id": CLIENT_ID, "X-Naver-Client-Secret": CLIENT_SECRET},
                    params={"query": "대통령 일정", "display": 6, "sort": "date"},
                    timeout=4)
                for item in resp.json().get("items", []):
                    pub = item.get("pubDate","")
                    try:
                        from email.utils import parsedate_to_datetime
                        pub_dt = parsedate_to_datetime(pub).date()
                    except:
                        pub_dt = today
                    if pub_dt >= today:
                        title = re.sub(r"<[^>]+>","", item.get("title","")).strip()
                        if title:
                            entries.append((pub_dt, title[:22]))
            except: pass

        # ── 정렬 + 1줄 포맷 ──
        if not entries:
            return ""
        entries.sort(key=lambda x: x[0])
        # 중복 제거
        seen, unique = set(), []
        for d, t in entries:
            key = t[:10]
            if key not in seen:
                seen.add(key); unique.append((d, t))
        parts = [
            f"{ev.month}.{ev.day:02d} ▶ {title[:18]}{'…' if len(title)>18 else ''}"
            for ev, title in unique[:3]
        ]
        return " / ".join(parts)

    except:
        return ""


@st.cache_data(ttl=1800)
def get_market_data(custom_ticker=""):
    d = {
        "kospi":"—","kospi_c":"","kospi_p":"","kospi_up":True,
        "kosdaq":"—","kosdaq_c":"","kosdaq_p":"","kosdaq_up":True,
        "nasdaq":"—","nasdaq_c":"","nasdaq_p":"","nasdaq_up":True,
        "sp500":"—","sp500_c":"","sp500_p":"","sp500_up":True,
        "usd_krw":"—","usd_c":"","usd_up":True,
        "oil":"—","oil_c":"","oil_up":True,
        "custom_name": custom_ticker,
        "custom_price":"—","custom_c":"","custom_up":True,
        "updated": (datetime.utcnow() + timedelta(hours=9)).strftime("%Y.%m.%d %H:%M"),
    }
    if YF_OK:
        tickers = {"^KS11":"kospi","^KQ11":"kosdaq","^IXIC":"nasdaq","^GSPC":"sp500",
                   "USDKRW=X":"usd","BZ=F":"oil"}
        if custom_ticker:
            tickers[custom_ticker] = "custom"
        for sym, key in tickers.items():
            try:
                h = yf.Ticker(sym).history(period="2d")
                if h.empty: continue
                cur  = float(h["Close"].iloc[-1])
                prev = float(h["Close"].iloc[-2]) if len(h) >= 2 else cur
                chg  = cur - prev; pct = chg/prev*100 if prev else 0
                arr  = "▲" if chg >= 0 else "▼"; up = (chg >= 0)
                if key == "kospi":
                    d.update({"kospi":f"{cur:,.2f}","kospi_c":f"{arr}{abs(chg):,.2f}","kospi_p":f"{pct:+.2f}%","kospi_up":up})
                elif key == "kosdaq":
                    d.update({"kosdaq":f"{cur:,.2f}","kosdaq_c":f"{arr}{abs(chg):,.2f}","kosdaq_p":f"{pct:+.2f}%","kosdaq_up":up})
                elif key == "nasdaq":
                    d.update({"nasdaq":f"{cur:,.2f}","nasdaq_c":f"{arr}{abs(chg):,.2f}","nasdaq_p":f"{pct:+.2f}%","nasdaq_up":up})
                elif key == "sp500":
                    d.update({"sp500":f"{cur:,.2f}","sp500_c":f"{arr}{abs(chg):,.2f}","sp500_p":f"{pct:+.2f}%","sp500_up":up})
                elif key == "usd":
                    d.update({"usd_krw":f"{cur:,.2f}","usd_c":f"{arr}{abs(chg):,.2f}","usd_up":up})
                elif key == "oil":
                    d.update({"oil":f"{cur:.2f}","oil_c":f"{arr}{abs(chg):.2f}","oil_up":up})
                elif key == "custom":
                    # 국내주(KS/KQ) → 원화, 해외 → USD
                    if sym.endswith(".KS") or sym.endswith(".KQ"):
                        price_str = f"{cur:,.0f}원"
                        chg_str   = f"{arr}{abs(chg):,.0f}"
                    else:
                        price_str = f"{cur:,.2f}"
                        chg_str   = f"{arr}{abs(chg):,.2f}"
                    d.update({"custom_price": price_str,"custom_c": chg_str,"custom_up": up})
            except: pass

    return d

def mhdr(d):
    def cs(v, up):
        c = "#C62828" if up else "#1565C0"
        return f"<span style='color:{c};font-size:10px;font-weight:600;'>{v}</span>"

    def cell(label, val, chg_html, border_left=True):
        bl = "border-left:1px solid #eee;padding-left:10px;" if border_left else ""
        return (f"<div style='margin-right:10px;{bl}'>"
                f"<div style='font-size:8px;color:#888;font-weight:700;'>{label}</div>"
                f"<div style='font-size:12px;font-weight:700;color:#003366;'>{val}</div>"
                f"<div>{chg_html}</div></div>")

    # 코스피 / 코스닥
    kospi_row  = cell("코스피",  d["kospi"],  cs(d["kospi_c"]+" "+d["kospi_p"],  d["kospi_up"]),  border_left=False)
    kosdaq_row = cell("코스닥",  d["kosdaq"], cs(d["kosdaq_c"]+" "+d["kosdaq_p"], d["kosdaq_up"]))
    sep1 = "<div style='border-left:2px solid #003366;height:30px;margin:0 10px;'></div>"

    # 나스닥 / S&P500
    nasdaq_row = cell("나스닥",  d["nasdaq"], cs(d["nasdaq_c"]+" "+d["nasdaq_p"], d["nasdaq_up"]))
    sp500_row  = cell("S&P500", d["sp500"],  cs(d["sp500_c"]+" "+d["sp500_p"],  d["sp500_up"]))
    sep2 = "<div style='border-left:2px solid #ddd;height:30px;margin:0 10px;'></div>"

    # USD/KRW + 두바이유
    usd_row = cell("USD/KRW",   d["usd_krw"], cs(d["usd_c"], d["usd_up"]))
    oil_row = cell("두바이유($/bbl)", d["oil"], cs(d["oil_c"], d["oil_up"]))

    # 구독자 지정 회사 주가 — 있으면 별도 섹션, 없으면 공간 자체 제거
    cn = d.get("custom_name","").strip()
    has_custom = bool(cn and d.get("custom_price","—") != "—")
    if has_custom:
        custom_section = (
            sep2 +
            "<div style='background:#F3E5F5;border-radius:4px;padding:4px 10px;margin-right:6px;'>"
            f"<div style='font-size:8px;color:#7B1FA2;font-weight:700;'>📌 {cn}</div>"
            f"<div style='font-size:13px;font-weight:800;color:#6A1B9A;'>{d['custom_price']}</div>"
            f"<div>{cs(d['custom_c'], d['custom_up'])}</div>"
            "</div>"
        )
    else:
        custom_section = ""

    updated = f"<div style='margin-left:auto;font-size:8px;color:#aaa;white-space:nowrap;'>{d['updated']}</div>"

    # 레이아웃: 코스피/코스닥 | 나스닥/S&P | [회사주가(있을때)] | USD/두바이유 | [대통령일정] | 업데이트
    return (
        f"<div style='background:white;border:1px solid #ddd;border-radius:5px;"
        f"padding:7px 14px;margin-bottom:8px;display:flex;align-items:center;"
        f"flex-wrap:wrap;gap:4px;font-family:{FONT_KR};'>"
        + kospi_row + kosdaq_row
        + sep1
        + nasdaq_row + sp500_row
        + custom_section      # 구독자 전용: 회사 주가
        + sep2
        + usd_row + oil_row
        + updated
        + "</div>"
    )

# ── 차트 함수 ──────────────────────────────────────────
def cfg(): return {'displayModeBar':False}

def extract_article_keywords(df, center_word='', top_n=28):
    """실제 기사 헤드라인에서 핵심 명사 추출 (Okt 형태소 분석)"""
    try:
        from konlpy.tag import Okt
        okt = Okt()
    except:
        return {}

    # 불용어 목록 (조사, 의존명사, 일반 단어 제외)
    STOPWORDS = {
        '것','수','등','및','또','이','그','저','이번','지난','올해','올','해','내','이날','당일',
        '관련','통해','위해','따라','대한','대해','따른','한','된','될','하는','있는','없는',
        '하며','하여','되는','되어','이에','오는','지','년','월','일','시','분','원','명',
        '건','개','건수','기사','뉴스','보도','기자','취재','언론','매체',
        '한국','전력','한전','KEPCO','kepco', center_word,
        '가운데','속에','앞서','이후','이전','이미','더욱','또한','특히','아울러',
        '계획','예정','방침','입장','밝혀','강조','주장','설명','언급',
    }

    word_sent_map = {}  # word -> {sent: count}
    word_headline_map = {}  # word -> [(date, media, headline)]

    for _, row in df.iterrows():
        text = str(row.get('헤드라인', ''))
        sent = row.get('감성', '중립')
        date = str(row.get('일자', ''))
        media = str(row.get('매체', ''))

        try:
            nouns = okt.nouns(text)
        except:
            continue

        seen = set()
        for noun in nouns:
            if len(noun) < 2: continue
            if noun in STOPWORDS: continue
            if noun not in word_sent_map:
                word_sent_map[noun] = {'부정': 0, '긍정': 0, '중립': 0}
                word_headline_map[noun] = []
            if noun not in seen:
                word_sent_map[noun][sent] += 1
                seen.add(noun)
            if len(word_headline_map[noun]) < 3:
                word_headline_map[noun].append((date, media, text))

    # 총 빈도 기준 정렬, 최소 2회 이상
    results = {}
    for word, sc in word_sent_map.items():
        total = sum(sc.values())
        if total < 2: continue
        dom_sent = max(sc, key=sc.get)
        results[word] = (dom_sent, total, word_headline_map.get(word, []))

    # top_n개 반환
    sorted_r = sorted(results.items(), key=lambda x: -x[1][1])[:top_n]
    return dict(sorted_r)


def plot_wordcloud(df, center_word='한국전력'):
    random.seed(42)

    # 실제 기사 기반 키워드 추출
    art_kws = extract_article_keywords(df, center_word=center_word, top_n=30)

    if art_kws:
        # 감성 색상: 부정>중립 → 빨강, 긍정>중립 → 파랑, 균형 → 회색
        items = list(art_kws.items())  # [(word, (sent, cnt, headlines))]
        max_cnt = items[0][1][1] if items else 1
    else:
        # fallback: 기존 POSITIVE/NEGATIVE 방식
        word_data = {}
        for sent, words in [('부정', NEGATIVE_WORDS), ('긍정', POSITIVE_WORDS)]:
            sub = df[df['감성']==sent]
            txt = " ".join(sub['헤드라인'].tolist())
            for w in words:
                cnt = txt.count(w)
                if cnt >= 1:
                    if w not in word_data or word_data[w][1] < cnt:
                        word_data[w] = (sent, cnt, [])
        items = sorted(word_data.items(), key=lambda x: -x[1][1])[:28]
        max_cnt = items[0][1][1] if items else 1

    # 중심 단어
    center_cnt = df['헤드라인'].str.contains(center_word, na=False, regex=False).sum() if center_word else len(df)
    xs   = [0]
    ys   = [0]
    texts= [center_word]
    sizes= [52]
    cols = ['#003366']
    hover= [f'<b>{center_word}</b> | 검색어 | {center_cnt}건']

    angle_step = 2.399; r_step = 0.15; base_r = 0.45
    for i, (word, info) in enumerate(items):
        if word == center_word:
            continue
        sent, cnt = info[0], info[1]
        headlines_raw = info[2] if len(info) > 2 else []

        angle = i * angle_step
        r = base_r + r_step * (i // 6)
        x = r * np.cos(angle) * 2.2 + random.uniform(-0.1, 0.1)
        y = r * np.sin(angle) + random.uniform(-0.07, 0.07)
        size = max(13, min(34, int(13 + (cnt / max_cnt) * 22)))
        color = '#C62828' if sent == '부정' else '#1565C0' if sent == '긍정' else '#777777'
        xs.append(x); ys.append(y); texts.append(word)
        sizes.append(size); cols.append(color)

        if headlines_raw:
            hl_lines = "<br>".join([f"· {h[2][:26]}  <span style='color:#aaa;font-size:9px;'>({h[0]} {h[1]})</span>" for h in headlines_raw[:3]])
        else:
            mask = df['헤드라인'].str.contains(word, na=False, regex=False)
            sample = df[mask][['일자','매체','헤드라인']].head(3)
            hl_lines = "<br>".join([f"· {r2['헤드라인'][:26]}  <span style='color:#aaa;font-size:9px;'>({r2['일자']} {r2['매체']})</span>" for _,r2 in sample.iterrows()]) if not sample.empty else "헤드라인 없음"

        hover.append(f'<b>{word}</b> | {sent} | {cnt}회<br>──────────<br>{hl_lines}')


    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=xs, y=ys, mode='text',
        text=texts,
        textfont=dict(size=sizes, color=cols, family=FONT_KR),
        hovertext=hover, hoverinfo='text',
        customdata=hover,
    ))
    fig.add_annotation(x=0, y=0, text='', showarrow=False)
    fig.update_layout(
        xaxis=dict(showgrid=False, showticklabels=False, zeroline=False, range=[-3.5, 3.5]),
        yaxis=dict(showgrid=False, showticklabels=False, zeroline=False, range=[-1.6, 1.6]),
        paper_bgcolor='white', plot_bgcolor='#FAFBFC',
        margin=dict(l=5, r=5, t=5, b=5), height=280,
        font=dict(family=FONT_KR),
    )
    return fig

def plot_donut(pos_n, neg_n, neu_n, total):
    fig = go.Figure(data=[go.Pie(
        labels=['긍정','중립','부정'], values=[pos_n, neu_n, neg_n],
        hole=0.55, marker=dict(colors=['#1565C0','#9E9E9E','#C62828'], line=dict(color='white', width=2)),
        textinfo='percent+label', textfont=dict(size=11, family=FONT_KR),
        hovertemplate='%{label}: %{value}건 (%{percent})<extra></extra>',
        direction='clockwise', sort=False, rotation=90,
    )])
    fig.update_layout(
        showlegend=False, margin=dict(l=5, r=5, t=5, b=5), height=230,
        paper_bgcolor='white', font=dict(family=FONT_KR),
        annotations=[dict(text=f"<b>{total}</b><br>건", x=0.5, y=0.5, font_size=16, showarrow=False, font=dict(family=FONT_KR, color='#003366'))]
    )
    return fig

def plot_buzz(df):
    daily = df.groupby('일자').size().reset_index(name='건수')
    daily['dt'] = pd.to_datetime(daily['일자'])
    by_sent = df.groupby(['일자','감성']).size().unstack(fill_value=0)
    fig = go.Figure()
    for sent, color in [('부정','#FFCDD2'),('중립','#E0E0E0'),('긍정','#BBDEFB')]:
        if sent in by_sent.columns:
            y = by_sent[sent].reindex(daily['일자'], fill_value=0).values
            fig.add_trace(go.Bar(x=daily['dt'], y=y, name=sent, marker_color=color, hovertemplate=f'{sent}: %{{y}}건<extra></extra>'))
    fig.add_trace(go.Scatter(x=daily['dt'], y=daily['건수'], mode='lines+markers', name='전체', line=dict(color='#003366', width=2), marker=dict(size=5, color='white', line=dict(width=2, color='#003366')), hovertemplate='%{x|%Y-%m-%d}<br>전체: <b>%{y}건</b><extra></extra>'))
    fig.update_layout(barmode='stack', plot_bgcolor='white', paper_bgcolor='white', font=dict(family=FONT_KR, size=11), margin=dict(l=40, r=10, t=10, b=35), height=230, hovermode='x unified', showlegend=True, legend=dict(orientation='h', y=1.08, x=1, xanchor='right', font=dict(size=10)), xaxis=dict(tickformat='%m/%d', showgrid=False, tickangle=-30), yaxis=dict(showgrid=True, gridcolor='#f5f5f5', rangemode='tozero'))
    return fig

def plot_kw_trend(df, kw, mode='daily', date_from=None, date_to=None):
    mask = df['헤드라인'].str.contains(kw, na=False, regex=False) | df['요약'].str.contains(kw, na=False, regex=False)
    kdf = df[mask].copy()
    if kdf.empty: return None
    if date_from: kdf = kdf[kdf['일자'] >= str(date_from)]
    if date_to:   kdf = kdf[kdf['일자'] <= str(date_to)]
    if kdf.empty: return None
    color_map = {'부정':'#C62828','긍정':'#1565C0','중립':'#777777'}
    if mode == 'daily':
        try:
            d_from = pd.to_datetime(str(date_from)) if date_from else pd.to_datetime(kdf['일자'].min())
            d_to   = pd.to_datetime(str(date_to))   if date_to   else pd.to_datetime(kdf['일자'].max())
            all_dates_full = pd.date_range(d_from, d_to, freq='D').strftime('%Y-%m-%d').tolist()
        except:
            all_dates_full = sorted(kdf['일자'].unique())
        grouped = kdf.groupby(['일자','감성']).size().unstack(fill_value=0).reindex(all_dates_full, fill_value=0)
        x = [pd.to_datetime(d) for d in grouped.index]; tick_fmt = '%m/%d'
        n_days = len(all_dates_full)
        dtick_ms = max(1, n_days // 20) * 86400000
    elif mode == 'monthly':
        kdf = kdf.copy(); kdf['월'] = kdf['일자'].str[:7]
        all_months = sorted(kdf['월'].unique())
        grouped = kdf.groupby(['월','감성']).size().unstack(fill_value=0).reindex(all_months, fill_value=0)
        x = grouped.index.tolist(); tick_fmt = None; dtick_ms = None
    else:
        kdf2 = kdf.copy()
        kdf2['시간_int'] = pd.to_numeric(kdf2['시간'], errors='coerce').fillna(0).astype(int)
        grouped = kdf2.groupby(['시간_int','감성']).size().unstack(fill_value=0).reindex(range(24), fill_value=0)
        x = list(range(24)); tick_fmt = None; dtick_ms = 2
    fig = go.Figure()
    for sent, color in color_map.items():
        y = grouped[sent].tolist() if sent in grouped.columns else [0]*len(x)
        m_mode = 'lines' if (mode == 'daily' and len(x) > 14) else 'lines+markers'
        fig.add_trace(go.Scatter(x=x, y=y, mode=m_mode, name=sent,
            line=dict(color=color, width=1.5), marker=dict(size=3),
            hovertemplate=f'<b>{sent}</b> %{{y}}건<extra></extra>'))
    mode_lbl = {'daily':'일자별','monthly':'월별','hourly':'시간대별'}
    xaxis_cfg = dict(showgrid=True, gridcolor='#f0f0f0', tickangle=-45, tickformat=tick_fmt)
    if mode == 'daily' and dtick_ms: xaxis_cfg.update({'dtick': dtick_ms, 'tickmode': 'linear'})
    elif mode == 'hourly': xaxis_cfg.update({'dtick': 2, 'tickmode': 'linear'})
    fig.update_layout(
        title=dict(text=f"<b>「{kw}」 {mode_lbl.get(mode,'')} 추이</b>", font=dict(size=13, color='#003366', family=FONT_KR)),
        plot_bgcolor='white', paper_bgcolor='white', font=dict(family=FONT_KR, size=11),
        margin=dict(l=40, r=10, t=40, b=55), height=270, hovermode='x unified',
        legend=dict(orientation='h', y=1.15, x=1, xanchor='right', font=dict(size=10)),
        xaxis=xaxis_cfg, yaxis=dict(showgrid=True, gridcolor='#f5f5f5', rangemode='tozero')
    )
    return fig

def plot_heatmap_with_hover(df):
    media_counts = df["매체"].value_counts().head(10)
    top_m = sorted(media_counts.index.tolist(), key=get_media_rank)[:8]
    top_c = [c for c in TOPIC_GROUPS if c in df["카테고리"].values][:8]
    if not top_m or not top_c: return None

    z = np.array([[
        round(len(df[(df["매체"]==m)&(df["카테고리"]==cat)&(df["감성"]=="부정")])/
              max(1, len(df[(df["매체"]==m)&(df["카테고리"]==cat)]))*100, 0)
        for cat in top_c] for m in top_m])

    # Build hover text with article headlines per cell
    hover_text = []
    for i, m in enumerate(top_m):
        row_hover = []
        for j, cat in enumerate(top_c):
            arts_cell = df[(df["매체"]==m)&(df["카테고리"]==cat)&(df["감성"]=="부정")]
            arts_cell = arts_cell.sort_values("일자", ascending=False).head(3)
            rate = z[i][j]
            if len(arts_cell) > 0:
                lines = [f"<b>{m} × {cat} ({rate:.0f}%)</b><br>──────────────"]
                for _, r2 in arts_cell.iterrows():
                    lines.append(f"· {r2['일자']}  {r2['헤드라인'][:22]}")
                cell_str = "<br>".join(lines)
            else:
                cell_str = f"<b>{m} × {cat}</b><br>부정 기사 없음"
            row_hover.append(cell_str)
        hover_text.append(row_hover)

    # Annotations (값만, 순위 없음)
    annotations = []
    for i in range(len(top_m)):
        for j in range(len(top_c)):
            annotations.append(dict(x=j, y=i, text=f"{z[i][j]:.0f}%", showarrow=False,
                font=dict(size=10, color='white' if z[i][j]>55 else '#333', family=FONT_KR)))

    fig = go.Figure(data=go.Heatmap(
        z=z, x=top_c, y=top_m,
        colorscale=[[0,'#F1F8E9'],[0.4,'#FFF9C4'],[0.7,'#FFB74D'],[1,'#B71C1C']],
        zmin=0, zmax=100,
        text=hover_text,
        hovertemplate='%{text}<extra></extra>',
        showscale=True,
    ))
    fig.update_layout(
        xaxis=dict(tickangle=-30, side='bottom', tickfont=dict(family=FONT_KR, size=10)),
        yaxis=dict(autorange='reversed', tickfont=dict(family=FONT_KR, size=10)),
        plot_bgcolor='white', paper_bgcolor='white',
        font=dict(family=FONT_KR, size=10),
        margin=dict(l=90, r=10, t=10, b=70), height=310,
        annotations=annotations
    )
    return fig

def plot_pr_gauge(pr_s, pr_c):
    fig = go.Figure(go.Indicator(
        mode="gauge+number", value=pr_s,
        number={"suffix":"점","font":{"size":22,"color":pr_c,"family":FONT_KR}},
        gauge={"axis":{"range":[0,100],"tickwidth":1,"tickcolor":"#ccc","tickfont":{"size":9}},
               "bar":{"color":pr_c,"thickness":0.25},"bgcolor":"#f5f5f5","borderwidth":0,
               "steps":[{"range":[0,40],"color":"#E8F5E9"},{"range":[40,70],"color":"#FFF8E1"},{"range":[70,100],"color":"#FFEBEE"}],
               "threshold":{"line":{"color":pr_c,"width":3},"thickness":0.75,"value":pr_s}}))
    fig.update_layout(height=140, margin=dict(l=20,r=20,t=10,b=10), paper_bgcolor='white', font=dict(family=FONT_KR))
    return fig

# ── Word 보고서 ────────────────────────────────────────
def set_table_header(table, headers, bg="003366", fg="FFFFFF"):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]; cell.text = h
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0] if p.runs else p.add_run(h)
        run.bold = True; run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(int(fg[:2],16), int(fg[2:4],16), int(fg[4:],16))
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),bg)
        tcPr.append(shd)

def make_full_word(cd):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)
    label=cd['label']; period_str=cd['period_str']; df=cd['df']; total=cd['total']
    pos_n=cd['pos_n']; neg_n=cd['neg_n']; neu_n=cd['neu_n']
    neg_rate=neg_n/total*100; pos_rate=pos_n/total*100
    pr_s=cd.get('pr_score',0); pr_l=cd.get('pr_lvl','—')
    neg_kws=cd['neg_kws']; pos_kws=cd['pos_kws']
    criticisms=cd['criticisms']; top_neg_cat=cd['top_neg_cat']; top_pos_cat=cd['top_pos_cat']

    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("한국전력 언론보도 유형분석 보고서"); r.bold=True; r.font.size=Pt(18); r.font.color.rgb=RGBColor(0,51,102)
    p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(f"{label}  |  {period_str}  |  {(datetime.utcnow()+timedelta(hours=9)).strftime('%Y년 %m월 %d일')}")
    doc.add_paragraph()
    def hd(txt, lv=1):
        h=doc.add_heading(txt,level=lv); h.runs[0].font.color.rgb=RGBColor(0,51,102); return h

    hd("01. 요약 및 제언")
    tone="부정 우세" if neg_n>pos_n*1.5 else "긍정 우세" if pos_n>neg_n*1.5 else "균형"
    doc.add_paragraph(cd['insights_text'])
    doc.add_paragraph(f"PR 리스크 스코어: {pr_s}점 ({pr_l}) | 논조: {tone}")
    doc.add_paragraph()

    hd("01. 논조 분석")
    tbl0=doc.add_table(rows=2,cols=4); tbl0.style='Table Grid'
    set_table_header(tbl0, ["구분","부정","긍정","중립"])
    row=tbl0.rows[1].cells
    row[0].text="건수 (비율)"; row[1].text=f"{neg_n}건 ({neg_rate:.1f}%)"
    row[2].text=f"{pos_n}건 ({pos_rate:.1f}%)"; row[3].text=f"{neu_n}건 ({neu_n/total*100:.1f}%)"
    doc.add_paragraph()

    hd("02. 매체별 논조 분석")
    media_list=df["매체"].value_counts().head(15).index.tolist()
    media_sorted=sorted(media_list,key=lambda m:(-df[(df["매체"]==m)&(df["감성"]=="부정")].shape[0]/max(1,df[df["매체"]==m].shape[0]), get_media_rank(m)))
    tbl1=doc.add_table(rows=1+len(media_sorted),cols=6); tbl1.style='Table Grid'
    set_table_header(tbl1, ["매체","등급","열독률(%)","총기사","부정","긍정"])
    for i,mname in enumerate(media_sorted,1):
        gi=MEDIA_GRADE.get(mname,{}); cells=tbl1.rows[i].cells
        cells[0].text=mname; cells[1].text=gi.get("grade","—"); cells[2].text=f"{gi.get('rate','')}"
        cells[3].text=str(df[df["매체"]==mname].shape[0])
        cells[4].text=str(df[(df["매체"]==mname)&(df["감성"]=="부정")].shape[0])
        cells[5].text=str(df[(df["매체"]==mname)&(df["감성"]=="긍정")].shape[0])
    doc.add_paragraph()

    hd("03. 논조별 키워드 TOP5")
    tbl2=doc.add_table(rows=1,cols=2); tbl2.style='Table Grid'
    set_table_header(tbl2, ["부정 키워드","긍정 키워드"])
    mx=max(len(neg_kws),len(pos_kws),1)
    for i in range(mx):
        r=tbl2.add_row().cells
        r[0].text=f"{neg_kws[i][0]}({neg_kws[i][1]}회)" if i<len(neg_kws) else ""
        r[1].text=f"{pos_kws[i][0]}({pos_kws[i][1]}회)" if i<len(pos_kws) else ""
    doc.add_paragraph()

    hd("04. 매체×이슈 부정 보도율")
    top_m_l=sorted(df["매체"].value_counts().head(8).index.tolist(),key=get_media_rank)
    top_c_l=[c for c in TOPIC_GROUPS if c in df["카테고리"].values][:7]
    if top_m_l and top_c_l:
        tbl3=doc.add_table(rows=1+len(top_m_l),cols=1+len(top_c_l)); tbl3.style='Table Grid'
        set_table_header(tbl3, ["매체"]+top_c_l)
        for i,mname in enumerate(top_m_l,1):
            cells=tbl3.rows[i].cells; cells[0].text=mname
            for j,cat in enumerate(top_c_l):
                nm=len(df[(df["매체"]==mname)&(df["카테고리"]==cat)])
                nn=len(df[(df["매체"]==mname)&(df["카테고리"]==cat)&(df["감성"]=="부정")])
                cells[j+1].text=f"{round(nn/max(1,nm)*100)}%" if nm>0 else "—"
    doc.add_paragraph()

    hd("05. 비판 포인트 & 대응 전략")
    paired=gen_paired_insights(criticisms)
    tbl4=doc.add_table(rows=1+len(paired),cols=4); tbl4.style='Table Grid'
    set_table_header(tbl4, ["비판 이슈","심각도","대응 전략","핵심 메시지"])
    for i,item in enumerate(paired,1):
        c=item["criticism"]; db=item["db"]; cells=tbl4.rows[i].cells
        cells[0].text=c["title"]; cells[1].text="●"*c["dots"]+"○"*(5-c["dots"])
        cells[2].text=db["action"]; cells[3].text=db["msg"]
    doc.add_paragraph()

    hd("06. 기사 전체 목록 (주요매체)")
    df_s = df.copy()
    df_s['_major'] = df_s['매체'].apply(lambda m: 0 if is_major_media(m) else 1)
    df_s['_r']     = df_s['매체'].apply(get_media_rank)
    # 주요매체만 필터
    df_s = df_s[df_s['_major'] == 0]
    df_s = df_s.sort_values(['일자','_r'], ascending=[False, True]).reset_index(drop=True)

    tbl5 = doc.add_table(rows=1, cols=5); tbl5.style='Table Grid'
    set_table_header(tbl5, ["No.","일자","매체","헤드라인 (링크)","논조"])
    for idx, row in enumerate(df_s.to_dict("records"), 1):
        cells = tbl5.add_row().cells
        cells[0].text = str(idx)
        cells[1].text = str(row["일자"])
        cells[2].text = str(row["매체"])
        cells[4].text = str(row["감성"])
        # 헤드라인 셀에 하이퍼링크 삽입
        headline_cell = cells[3]
        headline_cell.text = ""
        para = headline_cell.paragraphs[0]
        run  = para.add_run(str(row["헤드라인"]))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 51, 153)
        run.font.underline = True
        link_url = str(row.get("링크",""))
        if link_url:
            # docx 하이퍼링크 XML 삽입
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            part = doc.part
            r_id = part.relate_to(link_url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("r:id"), r_id)
            hyperlink.append(run._r)
            para._p.append(hyperlink)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0); return buf

# ── 섹션 헤더 ──────────────────────────────────────────
def divider(n, count_html=""):
    st.markdown(f"<div style='font-size:15px;font-weight:800;color:#003366;letter-spacing:.5px;border-bottom:2px solid #003366;padding-bottom:6px;margin:20px 0 10px;font-family:{FONT_KR};'>{n}{count_html}</div>", unsafe_allow_html=True)

# ── 위기관리 권고 박스 ──────────────────────────────────
def show_crisis_recommendation(pr_s, pr_l, label, cd=None):
    """PR 리스크 70점 이상 시 강력 권고 박스 + 실행전략 상위 3선 표출"""
    if pr_s < 70:
        return

    if pr_s >= 85:
        tier = "A등급 (긴급·전사 위기)"
        bg = "#FFEBEE"; border = "#B71C1C"; icon = "🚨"
        badge_color = "#B71C1C"
        trigger = "중앙부처 문의·조사 개시 / 중앙지·방송 집중 보도 / 전국 규모 단체 항의 가능성"
        immediate = [
            "① 골든아워(1시간) 내 홍보처 War Room 가동, 공식 Statement 즉각 발표",
            "② CEO 직접 담화 또는 현장 방문으로 책임 있는 리더십 메시지 발신",
            "③ 다크 사이트(Dark Site) 즉시 전환, 단일 창구로 정보 집중 제공",
            "④ 외부 전문가 공조 체제 구축 및 그룹 전체 위기관리 프로토콜 가동",
        ]
    elif pr_s >= 75:
        tier = "B등급 (경계·집중 대응)"
        bg = "#FFF3E0"; border = "#E65100"; icon = "⚠️"
        badge_color = "#E65100"
        trigger = "시·도 관공서 문의 / 지방 신문·지역방송 집중 보도 / 지역 단체 항의 가능성"
        immediate = [
            "① 48시간 내 공식 입장 발표 및 담당 부서 창구 일원화",
            "② 출입기자단 대상 백브리핑 즉시 개최, 오보 확산 차단",
            "③ 지역 언론 대응 강화 및 사업소 단위 위기관리 가동",
            "④ 소셜 리스닝 강화, 키워드 확산 경로 실시간 모니터링",
        ]
    else:  # 70~74
        tier = "C등급 (주의·예방 관리)"
        bg = "#FFFDE7"; border = "#F9A825"; icon = "⚡"
        badge_color = "#F57F17"
        trigger = "온라인 키워드 확산 / 전문지·업계 언론 집중 보도 / 이해관계자 문의 증가"
        immediate = [
            "① 주요 이슈 공식 입장 선제 정리 및 FAQ 문서 사전 준비",
            "② 정기 백브리핑 일정 앞당겨 관련 기자 선제 설명",
            "③ 팩트 허브(FAQ 페이지) 업데이트로 오보 예방 기반 구축",
            "④ 내부 실무진 위기 대응 매뉴얼 공유 및 모의 훈련 실시",
        ]

    # 해당 이슈의 상위 실행전략 3선 자동 매핑
    top_neg_cat = cd.get('top_neg_cat', '') if cd else ''
    strategy_hints = []
    if top_neg_cat in ("전기요금", "재무·경영"):
        strategy_hints = [
            ("Ⅱ-① 정기 백브리핑 정례화", "요금 현실화 복잡 현안을 심층 설명회로 오보 가능성 차단"),
            ("Ⅲ-③ 에너지 안보 캠페인", "요금 정상화가 국민 이익임을 알리는 대국민 공감 캠페인 즉시 기획"),
            ("Ⅳ-① 숏폼 콘텐츠 제작", "1분 내외 요금 체계 해설 영상으로 MZ세대 공감 확보"),
        ]
    elif top_neg_cat in ("전력망·설비",):
        strategy_hints = [
            ("Ⅰ-④ 실시간 여론 모니터링", "전력망 이슈 키워드 확산 경로 AI 기반 조기 포착"),
            ("Ⅱ-④ 미디어 데이 개최", "'에너지고속도로' 현장 시연으로 긍정 프레임 전환"),
            ("Ⅲ-① 상생 소통 라운드테이블", "건설 갈등 지역 주민과 정기 소통 채널 즉시 구축"),
        ]
    elif top_neg_cat in ("안전·사고",):
        strategy_hints = [
            ("Ⅰ-① 골든아워 War Room", "사고 감지 1시간 내 공식 입장 발표 체계 가동"),
            ("Ⅱ-③ 팩트 허브 운영", "사고 원인·재발 방지책 데이터 즉각 게시, 루머 차단"),
            ("Ⅴ-① CEO 타운홀 미팅", "CEO 직접 현장 방문·전 직원 대상 안전 의지 선언"),
        ]
    elif top_neg_cat in ("노사관계",):
        strategy_hints = [
            ("Ⅰ-① 골든아워 War Room", "파업 선언 즉시 전력 공급 안정 유지 Statement 발표"),
            ("Ⅴ-① CEO 타운홀 미팅", "CEO가 직접 직원에게 경영 현황·대화 의지 공유"),
            ("Ⅳ-② AI 챗봇 FAQ", "파업 관련 국민 문의에 24/7 정확한 정보 즉각 제공"),
        ]
    elif top_neg_cat in ("공기업·거버넌스",):
        strategy_hints = [
            ("Ⅰ-② 원스톱 승인 프로세스", "사건 감지 즉시 사실 확인·공식 입장 최단 시간 발표"),
            ("Ⅱ-③ 팩트 허브 운영", "의혹 제기 보도에 증거 데이터 즉각 반박 자료 게시"),
            ("Ⅴ-④ AAR 시스템 도입", "사건 처리 후 전 과정 정밀 분석, 재발 방지 백서 발간"),
        ]
    else:
        strategy_hints = [
            ("Ⅰ-① 골든아워 War Room", "위기 감지 1시간 내 공식 Statement 발표 체계 즉시 가동"),
            ("Ⅱ-① 정기 백브리핑 정례화", "현안 관련 출입기자단 심층 설명회로 오보 확산 차단"),
            ("Ⅳ-③ 소셜 리스닝 맞춤 대응", "온라인 논란 키워드 분석 후 직접 해답 콘텐츠 배포"),
        ]

    immediate_html = "".join([
        f"<div style='font-size:11px;color:#333;padding:4px 0;border-bottom:1px dashed #f0f0f0;'>{step}</div>"
        for step in immediate
    ])
    strategy_html = "".join([
        f"<div style='margin-bottom:6px;'>"
        f"<span style='font-size:11px;font-weight:700;color:{badge_color};'>{s[0]}</span>"
        f"<div style='font-size:10px;color:#666;margin-top:2px;'>{s[1]}</div>"
        f"</div>"
        for s in strategy_hints
    ])

    st.markdown(f"""<div style='background:{bg};border:2px solid {border};border-radius:8px;padding:16px 18px;margin-bottom:14px;font-family:{FONT_KR};'>
  <div style='display:flex;align-items:center;justify-content:space-between;margin-bottom:10px;'>
    <div>
      <span style='font-size:16px;font-weight:900;color:{border};'>{icon} PR 리스크 {pr_s}점</span>
      <span style='display:inline-block;background:{border};color:white;font-size:11px;font-weight:800;padding:2px 10px;border-radius:12px;margin-left:8px;'>{tier}</span>
    </div>
    <div style='font-size:10px;color:#999;text-align:right;'>위기관리 절차 즉시 시행 권고</div>
  </div>
  <div style='font-size:11px;color:#555;margin-bottom:10px;'><b>판단 기준:</b> {trigger}</div>
  <div style='display:flex;gap:12px;flex-wrap:wrap;'>
    <div style='flex:3;min-width:280px;background:rgba(255,255,255,0.7);border-radius:6px;padding:10px 14px;'>
      <div style='font-size:12px;font-weight:800;color:{border};margin-bottom:6px;'>🔥 즉각 실행 4대 조치</div>
      {immediate_html}
    </div>
    <div style='flex:2;min-width:200px;background:rgba(255,255,255,0.7);border-radius:6px;padding:10px 14px;'>
      <div style='font-size:12px;font-weight:800;color:#003366;margin-bottom:8px;'>📋 이슈별 실행전략 TOP 3</div>
      {strategy_html}
    </div>
  </div>
  <div style='font-size:10px;color:#aaa;margin-top:8px;border-top:1px solid #e0e0e0;padding-top:6px;'>
    ※ 세부 실행전략 20선 전문은 홍보처 위기관리 매뉴얼(Ⅰ~Ⅴ) 참조
  </div>
</div>""", unsafe_allow_html=True)


# ══ 보고서 렌더링 ══════════════════════════════════════
def render_report(cd):
    label=cd['label']; period_str=cd['period_str']; df=cd['df']
    total=cd['total']; pos_n=cd['pos_n']; neg_n=cd['neg_n']; neu_n=cd['neu_n']
    neg_kws=cd['neg_kws']; pos_kws=cd['pos_kws']
    top_neg_kw=cd['top_neg_kw']; criticisms=cd['criticisms']
    insights_text=cd['insights_text']; top_neg_cat=cd['top_neg_cat']; top_pos_cat=cd['top_pos_cat']
    top3_media=cd['top3_media']; trend_txt=cd['trend_txt']
    pr_s=cd.get('pr_score',0); pr_l=cd.get('pr_lvl','—'); pr_c=cd.get('pr_color','#888')
    neg_rate=neg_n/total*100; pos_rate=pos_n/total*100
    tone_sym="🔴" if neg_n>pos_n*1.5 else "🟢" if pos_n>neg_n*1.5 else "🟡"
    tone_txt="부정 우세" if neg_n>pos_n*1.5 else "긍정 우세" if pos_n>neg_n*1.5 else "균형"
    neg_kw_str=", ".join([f'{k}({v}회)' for k,v in neg_kws[:3]]) if neg_kws else "없음"
    neg_media_top=df[df['감성']=='부정']['매체'].value_counts().head(3)
    top_neg_m=", ".join([f"{m}({n}건)" for m,n in neg_media_top.items()]) if not neg_media_top.empty else "해당없음"

    # ─ 위기관리 권고 (조건부) ─

    # ═══ 00. KPI + 결론 (서술형 내러티브) ═══
    divider("01 · 요약 및 제언")

    # 내러티브 생성
    neg_top1_kw = neg_kws[0][0] if neg_kws else "해당없음"
    tone_desc = ("균형 있는 언론 환경이 유지되고 있습니다" if tone_txt=="균형"
                 else ("부정 보도가 많은 위기 국면입니다" if tone_txt=="부정 우세"
                       else "긍정 보도가 우세한 호의적 환경입니다"))
    # PR 리스크 한글 레벨
    pr_l_kr = {"HIGH":"높음","MEDIUM":"보통","LOW":"낮음"}.get(pr_l, pr_l)
    narrative_html = (
        f"<b>{period_str}</b> 네이버 기사 전체 <b>{total}건</b>을 전수 분석했습니다. "
        f"기간 내 <b>부정 보도 {neg_rate:.0f}%</b>, 긍정 보도 {pos_rate:.0f}%로 {tone_desc}. "
        f"'{neg_top1_kw}' 키워드가 부정 보도의 핵심이며 언론 리스크는 <b>{pr_s}점(보통)</b>입니다. "
        if pr_l == "MEDIUM" else
        f"<b>{period_str}</b> 네이버 기사 전체 <b>{total}건</b>을 전수 분석했습니다. "
        f"기간 내 <b>부정 보도 {neg_rate:.0f}%</b>, 긍정 보도 {pos_rate:.0f}%로 {tone_desc}. "
        f"'{neg_top1_kw}' 키워드가 부정 보도의 핵심이며 언론 리스크는 <b>{pr_s}점({pr_l_kr})</b>입니다. "
    )
    narrative_html += (
        f"위기관리 차원에서 선제적 대응과 '{top_neg_cat}' 이슈에 대한 공식 입장 발표가 권고됩니다. "
        f"반면 '{top_pos_cat}'와 관련한 보도는 긍정적입니다."
    )
    # KPI 서브텍스트용 한글 레벨
    pr_l_disp = pr_l_kr

    k1,k2,k3,k4,k5,k6=st.columns(6)
    for col,val,lbl,color,sub in [
        (k1,str(total),"총 기사","#003366",trend_txt[:18]),
        (k2,f"{neg_n}건","부정","#C62828",f"{neg_rate:.0f}%  {top_neg_cat[:6]}"),
        (k3,f"{pos_n}건","긍정","#1565C0",f"{pos_rate:.0f}%  {top_pos_cat[:6]}"),
        (k4,f"{neu_n}건","중립","#555",f"{neu_n/total*100:.0f}%"),
        (k5,tone_sym,"논조","#333",tone_txt),
        (k6,f"{pr_s}","PR리스크",pr_c,f"{pr_l_disp}  /100점"),
    ]:
        col.markdown(f"""<div style='background:white;border:1px solid #e8e8e8;border-top:3px solid {color};border-radius:4px;padding:8px 6px;text-align:center;font-family:{FONT_KR};'>
        <div style='font-size:19px;font-weight:700;color:{color};line-height:1.1;'>{val}</div>
        <div style='font-size:9px;font-weight:700;color:#999;letter-spacing:.4px;margin-top:2px;'>{lbl}</div>
        <div style='font-size:9px;color:#bbb;margin-top:1px;white-space:nowrap;overflow:hidden;'>{sub}</div></div>""", unsafe_allow_html=True)

    g1,g2=st.columns([1,3])
    with g1:
        st.plotly_chart(plot_pr_gauge(pr_s, pr_c), use_container_width=True, config=cfg())
    with g2:
        st.markdown(f"""<div style='background:#F8F9FA;border-left:3px solid #003366;border-radius:0 4px 4px 0;padding:14px 18px;font-size:12px;line-height:2.0;font-family:{FONT_KR};height:130px;overflow:auto;'>
        {narrative_html}
        </div>""", unsafe_allow_html=True)

    # ═══ 01. 워드 클라우드 ═══
    divider("02 · 워드 클라우드")
    wc1, wc2 = st.columns([3,1])
    with wc1:
        fig_wc = plot_wordcloud(df, center_word=label)
        st.plotly_chart(fig_wc, use_container_width=True, config=cfg())
    with wc2:
        st.markdown(f"""<div style='background:#F8F9FA;border-radius:6px;padding:10px;font-family:{FONT_KR};font-size:11px;'>
        <div style='font-weight:700;color:#003366;margin-bottom:6px;'>범례</div>
        <div style='margin-bottom:4px;'><span style='color:#C62828;font-weight:700;'>■</span> 부정 키워드</div>
        <div style='margin-bottom:4px;'><span style='color:#1565C0;font-weight:700;'>■</span> 긍정 키워드</div>
        <div style='margin-bottom:8px;'><span style='color:#888;font-weight:700;'>■</span> 중립 키워드</div>
        <div style='font-size:10px;color:#aaa;'>글자 크기 = 언급 빈도<br>커서를 단어에 올리면<br>상세 정보 표시</div>
        </div>""", unsafe_allow_html=True)

    # ═══ 02. 언론노출 추이 및 논조 분석 + 키워드 TOP3 ═══
    divider("03 · 언론노출 추이 및 논조 분석")
    b1, b2, b3 = st.columns([1, 1, 1])
    with b1:
        st.plotly_chart(plot_donut(pos_n, neg_n, neu_n, total), use_container_width=True, config=cfg())
    with b2:
        # 부정 키워드 TOP3 (추세버튼 없음)
        st.markdown(f"<div style='background:#FFEBEE;border:2px solid #C62828;border-radius:8px 8px 0 0;padding:6px 12px;font-size:12px;font-weight:800;color:#C62828;font-family:{FONT_KR};'>🔴 부정 키워드 TOP3</div>", unsafe_allow_html=True)
        neg3 = neg_kws[:3]
        for kw, cnt in neg3:
            bar_w = int(cnt / max(neg3[0][1], 1) * 100)
            st.markdown(f"""<div style='border:1px solid #FFCDD2;border-top:none;background:white;padding:8px 12px 5px;font-family:{FONT_KR};'>
  <div style='font-size:14px;font-weight:800;color:#C62828;'>{kw}</div>
  <div style='font-size:10px;color:#999;margin:2px 0 4px;'>{cnt}회</div>
  <div style='background:#f5f5f5;border-radius:3px;height:3px;'><div style='background:#C62828;width:{bar_w}%;height:3px;border-radius:3px;'></div></div>
</div>""", unsafe_allow_html=True)
    with b3:
        # 긍정 키워드 TOP3 (추세버튼 없음)
        st.markdown(f"<div style='background:#E3F2FD;border:2px solid #1565C0;border-radius:8px 8px 0 0;padding:6px 12px;font-size:12px;font-weight:800;color:#1565C0;font-family:{FONT_KR};'>🔵 긍정 키워드 TOP3</div>", unsafe_allow_html=True)
        pos3 = pos_kws[:3]
        for kw, cnt in pos3:
            bar_w = int(cnt / max(pos3[0][1], 1) * 100)
            st.markdown(f"""<div style='border:1px solid #BBDEFB;border-top:none;background:white;padding:8px 12px 5px;font-family:{FONT_KR};'>
  <div style='font-size:14px;font-weight:800;color:#1565C0;'>{kw}</div>
  <div style='font-size:10px;color:#999;margin:2px 0 4px;'>{cnt}회</div>
  <div style='background:#f5f5f5;border-radius:3px;height:3px;'><div style='background:#1565C0;width:{bar_w}%;height:3px;border-radius:3px;'></div></div>
</div>""", unsafe_allow_html=True)

    # ═══ 03. 요주의/우호 매체 ═══
    divider("04 · 요주의/우호 매체")

    # 등급 범례
    grade_legend_html = "".join([
        f"<span style='background:{c};color:white;padding:2px 7px;border-radius:3px;"
        f"font-size:10px;font-weight:700;margin-right:6px;'>{g}</span>"
        f"<span style='font-size:10px;color:#555;margin-right:14px;'>{desc}</span>"
        for g, c, desc in [
            ("S","#B71C1C","열독률 상위 10%"),
            ("A","#E64A19","상위 20%"),
            ("B","#1565C0","상위 40%"),
            ("C","#2E7D32","상위 60%"),
            ("D","#616161","하위 40%"),
        ]
    ])
    st.markdown(
        f"<div style='background:#F8F9FA;border:1px solid #eee;border-radius:6px;"
        f"padding:7px 14px;margin-bottom:10px;font-family:{FONT_KR};display:flex;"
        f"align-items:center;flex-wrap:wrap;gap:4px;'>"
        f"<span style='font-size:10px;font-weight:700;color:#888;margin-right:10px;'>등급 범례</span>"
        f"{grade_legend_html}</div>",
        unsafe_allow_html=True
    )

    # ═══ 03. 요주의/우호 매체 (좌우 나란히) ═══
    media_all_r = df["매체"].value_counts()
    media_with_neg = []
    for m, tot in media_all_r.items():
        if tot < 3: continue
        n_n = int(df[(df["매체"]==m)&(df["감성"]=="부정")].shape[0])
        n_p = int(df[(df["매체"]==m)&(df["감성"]=="긍정")].shape[0])
        neg_pct = n_n/tot*100
        pos_pct = n_p/tot*100
        gi = MEDIA_GRADE.get(m, {})
        media_with_neg.append((m, gi, neg_pct, n_n, tot, pos_pct, n_p))

    top_neg_media = sorted(media_with_neg, key=lambda x: -x[2])[:5]
    top_pos_media = sorted(media_with_neg, key=lambda x: -x[5])[:5]

    def grade_badge(gi):
        g = gi.get("grade",""); c = GRADE_COLOR.get(g,"#aaa")
        return f"<span style='background:{c};color:white;padding:1px 5px;border-radius:3px;font-size:9px;font-weight:700;'>{g}</span>" if g else "<span style='background:#ddd;color:#999;padding:1px 5px;border-radius:3px;font-size:9px;'>—</span>"

    mc1, mc2 = st.columns(2)

    # CSS 툴팁 스타일 (한 번만 정의)
    tooltip_css = f"""<style>
.tip-wrap {{ position:relative; display:inline-block; cursor:pointer; }}
.tip-wrap .tip-box {{
  visibility:hidden; opacity:0;
  background:#1a1a2e; color:#fff;
  font-size:10px; line-height:1.6;
  border-radius:6px; padding:8px 10px;
  position:absolute; z-index:999;
  bottom:125%; left:50%; transform:translateX(-50%);
  width:260px; white-space:pre-line;
  box-shadow:0 4px 14px rgba(0,0,0,.35);
  transition:opacity .15s;
  font-family:{FONT_KR};
  pointer-events:none;
}}
.tip-wrap:hover .tip-box {{ visibility:visible; opacity:1; }}
</style>"""

    def make_tip(arts_tip, color):
        """기사 목록을 툴팁 HTML로 변환"""
        if arts_tip is None or len(arts_tip) == 0:
            return "<span style='color:{};font-size:11px;'>0</span>".format(color)
        lines = "\n".join([f"· {r['일자']}  {r['헤드라인'][:22]}" for _, r in arts_tip.head(5).iterrows()])
        cnt = len(arts_tip)
        cnt_display = f"{cnt}건" if cnt <= 5 else f"{cnt}건 (상위 5건)"
        tip_content = f"{cnt_display}\n──────────────\n{lines}"
        return (f"<span class='tip-wrap' style='color:{color};font-size:11px;font-weight:600;'>"
                f"{cnt}"
                f"<span class='tip-box'>{tip_content}</span>"
                f"</span>")

    with mc1:
        neg_rows = ""
        for m, gi, neg_pct, n_n, tot, _, _ in top_neg_media:
            arts_neg = df[(df["매체"]==m)&(df["감성"]=="부정")].sort_values("일자", ascending=False)
            arts_tot = df[df["매체"]==m].sort_values("일자", ascending=False)
            tip_neg = make_tip(arts_neg, "#C62828")
            tip_tot = make_tip(arts_tot, "#888")
            neg_rows += (f"<tr>"
                f"<td style='padding:5px 8px;font-size:11px;font-weight:600;'>{m}</td>"
                f"<td style='padding:5px 8px;text-align:center;'>{grade_badge(gi)}</td>"
                f"<td style='padding:5px 8px;text-align:center;font-size:11px;font-weight:700;color:#C62828;'>{neg_pct:.0f}%</td>"
                f"<td style='padding:5px 8px;text-align:center;'>{tip_neg}</td>"
                f"<td style='padding:5px 8px;text-align:center;'>{tip_tot}</td>"
                f"</tr>")
        st.markdown(f"""{tooltip_css}<div style='background:#FFF5F5;border:1.5px solid #FFCDD2;border-radius:8px;padding:10px;font-family:{FONT_KR};'>
  <div style='font-size:12px;font-weight:800;color:#C62828;margin-bottom:6px;'>🚨 요주의 매체 (부정 보도 집중)</div>
  <table style='width:100%;border-collapse:collapse;'>
    <tr style='background:#FFEBEE;font-size:10px;color:#888;'><th style='padding:4px 8px;text-align:left;'>매체</th><th style='padding:4px 8px;'>등급</th><th style='padding:4px 8px;'>부정%</th><th style='padding:4px 8px;'>부정</th><th style='padding:4px 8px;'>전체</th></tr>
    {neg_rows}
  </table>
</div>""", unsafe_allow_html=True)

    with mc2:
        pos_rows = ""
        for m, gi, _, _, tot, pos_pct, n_p in top_pos_media:
            arts_pos = df[(df["매체"]==m)&(df["감성"]=="긍정")].sort_values("일자", ascending=False)
            arts_tot = df[df["매체"]==m].sort_values("일자", ascending=False)
            tip_pos = make_tip(arts_pos, "#1565C0")
            tip_tot = make_tip(arts_tot, "#888")
            pos_rows += (f"<tr>"
                f"<td style='padding:5px 8px;font-size:11px;font-weight:600;'>{m}</td>"
                f"<td style='padding:5px 8px;text-align:center;'>{grade_badge(gi)}</td>"
                f"<td style='padding:5px 8px;text-align:center;font-size:11px;font-weight:700;color:#1565C0;'>{pos_pct:.0f}%</td>"
                f"<td style='padding:5px 8px;text-align:center;'>{tip_pos}</td>"
                f"<td style='padding:5px 8px;text-align:center;'>{tip_tot}</td>"
                f"</tr>")
        st.markdown(f"""{tooltip_css}<div style='background:#F0F8FF;border:1.5px solid #BBDEFB;border-radius:8px;padding:10px;font-family:{FONT_KR};'>
  <div style='font-size:12px;font-weight:800;color:#1565C0;margin-bottom:6px;'>✅ 우호 매체 (긍정 보도 집중)</div>
  <table style='width:100%;border-collapse:collapse;'>
    <tr style='background:#E3F2FD;font-size:10px;color:#888;'><th style='padding:4px 8px;text-align:left;'>매체</th><th style='padding:4px 8px;'>등급</th><th style='padding:4px 8px;'>긍정%</th><th style='padding:4px 8px;'>긍정</th><th style='padding:4px 8px;'>전체</th></tr>
    {pos_rows}
  </table>
</div>""", unsafe_allow_html=True)

    # ═══ 05. 매체×이슈 부정 보도 매트릭스 ═══
    st.markdown("<div style='margin-top:24px;'></div>", unsafe_allow_html=True)
    divider("05 · 매체×이슈 부정 보도 매트릭스 — 커서를 셀에 올리면 기사 확인")
    fig_hm = plot_heatmap_with_hover(df)
    if fig_hm:
        st.plotly_chart(fig_hm, use_container_width=True, config=cfg())
    else:
        st.caption("데이터 부족으로 히트맵 생성 불가")

    # ═══ 06. 위기관리 키워드 추세 (부정 Top1, 최근 3개월 일자별) ═══
    # ═══ 05-B. 블랙리스트 기자 ═══
    st.markdown("<div style='margin-top:20px;'></div>", unsafe_allow_html=True)
    divider("06 · 기자 블랙리스트  *기간 내 3건 이상 부정보도")

    df_rep = df[df["기자"] != "—"].copy()
    if df_rep.empty:
        st.caption("기간 내 3건 이상의 부정 보도를 작성한 기자가 없습니다")
    else:
        rep_stats = []
        for rep_name, grp in df_rep.groupby("기자"):
            tot_r   = len(grp)
            neg_r   = int((grp["감성"]=="부정").sum())
            pos_r   = int((grp["감성"]=="긍정").sum())
            neg_pct = neg_r / tot_r * 100
            medias  = grp["매체"].value_counts().index[:2].tolist()
            rep_stats.append((rep_name, tot_r, neg_r, pos_r, neg_pct, medias, grp))
        # 부정 건수 3건 이상만, 부정 비율 내림차순
        rep_stats = [r for r in rep_stats if r[2] >= 3]
        rep_stats.sort(key=lambda x: (-x[2], -x[4]))
        show_n_bl = st.session_state.get(f'bl_show_{label}', 3)
        black_list = rep_stats[:show_n_bl]

        if not black_list:
            st.caption("기간 내 3건 이상의 부정 보도를 작성한 기자가 없습니다")
        else:
            bl1, bl2 = st.columns(2)
            with bl1:
                st.markdown(
                    f"<div style='background:#FFF5F5;border:1.5px solid #FFCDD2;border-radius:8px;"
                    f"padding:10px;font-family:{FONT_KR};'>",
                    unsafe_allow_html=True
                )
                st.markdown(
                    f"<div style='font-size:12px;font-weight:800;color:#C62828;margin-bottom:6px;'>"
                    f"🚨 부정 보도 집중 기자 TOP {len(black_list)}</div>",
                    unsafe_allow_html=True
                )

                # 툴팁 CSS (재사용)
                bl_css = f"""<style>
.bl-tip {{ position:relative;display:inline-block;cursor:pointer; }}
.bl-tip .bl-box {{
  visibility:hidden;opacity:0;
  background:#1a1a2e;color:#fff;font-size:10px;line-height:1.7;
  border-radius:6px;padding:8px 10px;
  position:absolute;z-index:999;bottom:120%;left:0;
  width:300px;white-space:pre-line;
  box-shadow:0 4px 14px rgba(0,0,0,.35);
  transition:opacity .15s;font-family:{FONT_KR};pointer-events:none;
}}
.bl-tip:hover .bl-box {{ visibility:visible;opacity:1; }}
</style>"""

                bl_rows = ""
                for rank_i, (rep_name, tot_r, neg_r, pos_r, neg_pct, medias, grp) in enumerate(black_list, 1):
                    media_str = " · ".join(medias)
                    gi_r = MEDIA_GRADE.get(medias[0] if medias else "", {})
                    grade_r = gi_r.get("grade","")
                    gc_r = GRADE_COLOR.get(grade_r,"#aaa")
                    badge_r = (f"<span style='background:{gc_r};color:white;padding:0 3px;"
                               f"border-radius:2px;font-size:8px;font-weight:700;'>{grade_r}</span> ") if grade_r else ""

                    # 툴팁 기사 목록
                    neg_arts = grp[grp["감성"]=="부정"].sort_values("일자", ascending=False)
                    tip_lines = "\n".join([
                        f"· {r['일자']}  [{r['매체']}]  {r['헤드라인'][:20]}"
                        for _, r in neg_arts.head(5).iterrows()
                    ])
                    tip_content = f"부정 기사 {neg_r}건\n──────────────\n{tip_lines}"

                    bar_w = int(neg_pct)
                    bl_rows += f"""<tr style='border-bottom:1px solid #FFE0E0;'>
  <td style='padding:6px 6px;font-size:11px;color:#888;text-align:center;font-weight:700;'>{rank_i}</td>
  <td style='padding:6px 8px;'>
    <div class='bl-tip'>
      <span style='font-size:12px;font-weight:800;color:#C62828;'>{rep_name} 기자</span>
      <span class='bl-box'>{tip_content}</span>
    </div>
    <div style='font-size:9px;color:#888;margin-top:1px;'>{badge_r}{media_str}</div>
  </td>
  <td style='padding:6px 8px;text-align:center;'>
    <span style='font-size:13px;font-weight:800;color:#C62828;'>{neg_r}</span>
    <span style='font-size:9px;color:#aaa;'>/{tot_r}건</span>
  </td>
  <td style='padding:6px 8px;min-width:80px;'>
    <div style='background:#f5f5f5;border-radius:4px;height:7px;'>
      <div style='background:#C62828;width:{bar_w}%;height:7px;border-radius:4px;'></div>
    </div>
    <div style='font-size:9px;color:#C62828;font-weight:700;margin-top:1px;'>{neg_pct:.0f}%</div>
  </td>
</tr>"""

                st.markdown(
                    f"""{bl_css}<table style='width:100%;border-collapse:collapse;font-family:{FONT_KR};'>
  <tr style='background:#FFEBEE;font-size:10px;color:#888;'>
    <th style='padding:4px 6px;'>#</th>
    <th style='padding:4px 8px;text-align:left;'>기자 / 매체</th>
    <th style='padding:4px 8px;'>부정건수</th>
    <th style='padding:4px 8px;text-align:left;'>부정비율</th>
  </tr>
  {bl_rows}
</table>""",
                    unsafe_allow_html=True
                )
                # 더보기 버튼
                total_bl = len(rep_stats)
                cur_show = st.session_state.get(f'bl_show_{label}', 3)
                if cur_show < total_bl:
                    if st.button(f"▼ 더보기 ({total_bl - cur_show}명 더)", key=f"bl_more_{label}"):
                        st.session_state[f'bl_show_{label}'] = cur_show + 3
                        st.rerun()
                st.markdown("</div>", unsafe_allow_html=True)

            with bl2:
                st.markdown(
                    f"""<div style='background:#F8F9FA;border:1px solid #ddd;border-radius:8px;
padding:14px 16px;font-family:{FONT_KR};font-size:11px;line-height:1.9;color:#444;'>
<div style='font-size:12px;font-weight:800;color:#003366;margin-bottom:8px;'>📋 블랙리스트 기자 활용 가이드</div>
<div style='margin-bottom:6px;'><b style='color:#C62828;'>①</b> 커서를 기자명에 올리면 해당 기사 목록 확인</div>
<div style='margin-bottom:6px;'><b style='color:#C62828;'>②</b> 해당 기자 취재 시 사실관계 자료 선제 제공</div>
<div style='margin-bottom:6px;'><b style='color:#C62828;'>③</b> 보도자료 배포 시 우선 배제 또는 별도 관리</div>
<div style='margin-bottom:6px;'><b style='color:#C62828;'>④</b> 동일 매체 내 우호 기자와 관계 집중 강화</div>
<div style='margin-top:10px;background:#FFF8E1;border-left:3px solid #F9A825;
padding:7px 10px;border-radius:0 4px 4px 0;font-size:10px;color:#777;'>
⚠️ 기자명은 description의 '홍길동 기자' 패턴으로 자동 추출합니다.
매체에 따라 미표기되는 경우가 있어 일부 누락될 수 있습니다.</div>
</div>""",
                    unsafe_allow_html=True
                )
    top_crisis_kw = neg_kws[0][0] if neg_kws else None
    if top_crisis_kw:
        divider(f"07 · 위기관리 키워드 추세 — 「{top_crisis_kw}」 최근 3개월 일별 노출")
        crisis_end = datetime.now()
        crisis_start = crisis_end - timedelta(days=90)
        with st.spinner(f"'{top_crisis_kw}' 최근 3개월 데이터 수집 중..."):
            crisis_raw = get_news(top_crisis_kw, 3000)
        crisis_records = []
        for a in crisis_raw:
            pub = a.get("pubDate","")
            try:
                ad = datetime.strptime(pub[:16], "%a, %d %b %Y").date()
                if not (crisis_start.date() <= ad <= crisis_end.date()): continue
                ds = ad.strftime("%Y-%m-%d")
            except:
                ds = pub[:10]
            title = clean(a.get("title",""))
            if not is_relevant(title): continue
            crisis_records.append({"일자": ds})
        if crisis_records:
            cdf = pd.DataFrame(crisis_records)
            daily_crisis = cdf.groupby("일자").size().reset_index(name="건수")
            daily_crisis["dt"] = pd.to_datetime(daily_crisis["일자"])
            all_days = pd.date_range(crisis_start, crisis_end, freq='D')
            daily_full = pd.DataFrame({"dt": all_days})
            daily_full["일자"] = daily_full["dt"].dt.strftime("%Y-%m-%d")
            daily_merged = daily_full.merge(daily_crisis[["일자","건수"]], on="일자", how="left").fillna(0)
            daily_merged["건수"] = daily_merged["건수"].astype(int)

            fig_crisis = go.Figure()
            fig_crisis.add_trace(go.Scatter(
                x=daily_merged["dt"],
                y=daily_merged["건수"],
                mode="lines+markers",
                line=dict(color="#C62828", width=2),
                marker=dict(size=4, color="#C62828"),
                fill="tozeroy",
                fillcolor="rgba(198,40,40,0.08)",
                hovertemplate="%{x|%m월 %d일} · %{y}회 노출<extra></extra>",
            ))
            fig_crisis.update_layout(
                plot_bgcolor="white", paper_bgcolor="white",
                font=dict(family=FONT_KR, size=11),
                margin=dict(l=40, r=10, t=10, b=40), height=240,
                hovermode="x unified",
                xaxis=dict(tickformat="%m/%d", showgrid=False, tickangle=-30,
                           dtick=7*86400000, tickmode="linear"),
                yaxis=dict(showgrid=True, gridcolor="#f5f5f5", rangemode="tozero"),
            )
            st.plotly_chart(fig_crisis, use_container_width=True, config=cfg())
        else:
            st.caption("최근 3개월 해당 키워드 데이터 없음")

    # ═══ 07. 비판 포인트 레이더 + As-Is/To-Be ═══
    divider("08 · 비판 포인트 & 대응 전략")
    paired = gen_paired_insights(criticisms)

    # 레이더차트: 카테고리별 부정 건수 기반 6각형
    cat_labels_all = list(TOPIC_GROUPS.keys())
    cat_neg_counts = {cat: int(df[(df['카테고리']==cat)&(df['감성']=='부정')].shape[0]) for cat in cat_labels_all}
    cat_neg_sorted = sorted(cat_neg_counts.items(), key=lambda x: -x[1])
    radar_cats = [c for c,v in cat_neg_sorted if v > 0][:6]
    if len(radar_cats) < 3:
        radar_cats = [c for c,v in cat_neg_sorted][:6]
    radar_vals = [cat_neg_counts.get(c, 0) for c in radar_cats]
    max_v = max(radar_vals) if radar_vals else 1
    max_v = max_v if max_v > 0 else 1
    radar_norm = [round(v/max_v*5, 1) for v in radar_vals]

    # 레이더 Plotly
    radar_labels_short = [c[:7] for c in radar_cats]
    fig_radar = go.Figure()
    fig_radar.add_trace(go.Scatterpolar(
        r=radar_norm + [radar_norm[0]],
        theta=radar_labels_short + [radar_labels_short[0]],
        fill='toself',
        fillcolor='rgba(198,40,40,0.15)',
        line=dict(color='#C62828', width=2),
        marker=dict(size=7, color='#C62828'),
        name='비판 강도',
        hovertemplate='<b>%{theta}</b><br>강도: %{r:.1f}/5<br>' +
            '<br>'.join([f"{radar_cats[i]}: {radar_vals[i]}건" for i in range(len(radar_cats))]) + '<extra></extra>',
    ))
    # 평균 라인
    avg_v = sum(radar_norm)/len(radar_norm) if radar_norm else 2.5
    fig_radar.add_trace(go.Scatterpolar(
        r=[avg_v]*len(radar_labels_short) + [avg_v],
        theta=radar_labels_short + [radar_labels_short[0]],
        mode='lines',
        line=dict(color='#003366', width=1, dash='dot'),
        name='평균',
        hoverinfo='skip',
    ))
    fig_radar.update_layout(
        polar=dict(
            radialaxis=dict(visible=True, range=[0, 5], tickfont=dict(size=9), showticklabels=True, gridcolor='#eee'),
            angularaxis=dict(tickfont=dict(size=11, family=FONT_KR, color='#333'), gridcolor='#eee'),
            bgcolor='white',
        ),
        showlegend=False,
        paper_bgcolor='white',
        font=dict(family=FONT_KR),
        margin=dict(l=40, r=40, t=30, b=30),
        height=360,
        title=dict(text=f"<b>주요 비판 포인트 우선순위</b><br><sub style='font-size:9px;color:#888;'>비판 강도(5점 만점)</sub>",
                   font=dict(size=12, color='#003366', family=FONT_KR), x=0.5, xanchor='center'),
    )

    r_col, detail_col = st.columns([1, 1])
    with r_col:
        st.plotly_chart(fig_radar, use_container_width=True, config=cfg())

    with detail_col:
        visible = [(cat, val) for cat, val in cat_neg_sorted[:6] if val > 0]
        # 레이더 차트와 동일한 높이(360px) flexbox로 항목 균등 배분
        rows_html = ""
        for rank_i, (cat, val) in enumerate(visible, 1):
            score = round(val/max_v*5, 1)
            bar_w = int(score/5*100)
            num_circle = ["①","②","③","④","⑤","⑥"][rank_i-1]
            rows_html += f"""<div style='display:flex;align-items:center;gap:10px;font-family:{FONT_KR};'>
  <span style='font-size:16px;font-weight:800;color:#C62828;min-width:26px;flex-shrink:0;'>{num_circle}</span>
  <div style='flex:1;min-width:0;'>
    <div style='display:flex;justify-content:space-between;margin-bottom:5px;'>
      <span style='font-size:12px;font-weight:700;color:#333;'>{cat}</span>
      <span style='font-size:12px;font-weight:700;color:#C62828;white-space:nowrap;margin-left:8px;'>{score}/5점 ({val}건)</span>
    </div>
    <div style='background:#f5f5f5;border-radius:4px;height:8px;'><div style='background:#C62828;width:{bar_w}%;height:8px;border-radius:4px;'></div></div>
  </div>
</div>"""

        st.markdown(f"""<div style='height:360px;display:flex;flex-direction:column;justify-content:space-between;padding:10px 0;'>
  <div style='font-size:12px;font-weight:800;color:#003366;font-family:{FONT_KR};padding-bottom:4px;border-bottom:1px solid #eee;'>카테고리별 비판 강도 순위</div>
  <div style='display:flex;flex-direction:column;justify-content:space-evenly;flex:1;padding-top:6px;'>
    {rows_html.replace(chr(10), "")}
  </div>
</div>""", unsafe_allow_html=True)

    # As-Is / To-Be
    col_asis, col_tobe = st.columns(2)
    with col_asis:
        st.markdown(f"<div style='background:#C62828;color:white;padding:7px 14px;border-radius:6px 6px 0 0;font-size:13px;font-weight:800;font-family:{FONT_KR};'>🔴 현재 문제점 (As-Is)</div>", unsafe_allow_html=True)
        for i, item in enumerate(paired, 1):
            c = item["criticism"]
            dots_str = "●"*c["dots"]+"○"*(5-c["dots"])
            cat_val = c.get("category", c["title"])
            # 해당 카테고리 부정 기사 헤드라인 최대 2건
            cat_arts = df[(df["카테고리"]==cat_val)&(df["감성"]=="부정")].sort_values("일자", ascending=False).head(2)
            hl_items = ""
            for _, row in cat_arts.iterrows():
                hl = str(row["헤드라인"])[:30]
                tooltip = f"{row['일자']} · {row['매체']}"
                hl_items += f"<div title='{tooltip}' style='font-size:10px;color:#888;padding:2px 0;border-left:2px solid #FFCDD2;padding-left:6px;margin-top:3px;cursor:default;overflow:hidden;white-space:nowrap;text-overflow:ellipsis;' title='{tooltip}'>📰 {hl}</div>"
            if not hl_items:
                pts_html = "  /  ".join(c["points"])
                hl_items = f"<div style='font-size:10px;color:#aaa;'>{pts_html}</div>"
            st.markdown(f"""<div style='border:1px solid #FFCDD2;border-top:none;background:#FFF8F8;padding:10px 14px;margin-bottom:3px;font-family:{FONT_KR};height:90px;overflow:hidden;'>
  <div style='display:flex;justify-content:space-between;align-items:center;margin-bottom:4px;'>
    <span style='font-size:12px;font-weight:700;color:#C62828;'>이슈 {i}. {c["title"]}</span>
    <span style='font-size:10px;color:#C62828;letter-spacing:1px;'>{dots_str}</span>
  </div>
  {hl_items}
</div>""", unsafe_allow_html=True)
    with col_tobe:
        st.markdown(f"<div style='background:#1565C0;color:white;padding:7px 14px;border-radius:6px 6px 0 0;font-size:13px;font-weight:800;font-family:{FONT_KR};'>✅ 개선 방향 (To-Be)</div>", unsafe_allow_html=True)
        for i, item in enumerate(paired, 1):
            db = item["db"]
            step1 = db["steps"][0] if db["steps"] else db["action"]
            st.markdown(f"""<div style='border:1px solid #BBDEFB;border-top:none;background:#F0F8FF;padding:10px 14px;margin-bottom:3px;font-family:{FONT_KR};height:90px;overflow:hidden;'>
  <div style='font-size:12px;font-weight:700;color:#1565C0;margin-bottom:5px;'>전략 {i}. {db["action"]}</div>
  <div style='font-size:11px;background:white;border-left:3px solid #003366;padding:4px 8px;color:#003366;font-weight:700;'>📌 {db["msg"]}</div>
</div>""", unsafe_allow_html=True)

    # ═══ 07. 기사 목록 ═══
    neg_cnt = int(df['감성'].value_counts().get('부정',0))
    neu_cnt = int(df['감성'].value_counts().get('중립',0))
    pos_cnt = int(df['감성'].value_counts().get('긍정',0))
    count_html = f" <span style='font-size:12px;font-weight:400;color:#888;'>🔴 부정 {neg_cnt} · 🟡 중립 {neu_cnt} · 🟢 긍정 {pos_cnt} · 총 {total}건</span>"
    divider("09 · 기사 목록", count_html)

    fdf = df.copy()
    fdf['_major'] = fdf['매체'].apply(lambda m: 0 if is_major_media(m) else 1)
    fdf['_rank']  = fdf['매체'].apply(get_media_rank)
    fdf = fdf.sort_values(['일자','_major','_rank'], ascending=[False, True, True]).reset_index(drop=True)

    # 컬럼별 필터 — 1줄 인라인
    all_dates = ["전체"]+sorted(fdf["일자"].unique().tolist(), reverse=True)
    all_media = ["전체"]+sorted(fdf["매체"].unique().tolist(), key=get_media_rank)
    all_sent  = ["전체","부정","중립","긍정"]
    all_cat   = ["전체"]+sorted(fdf["카테고리"].unique().tolist())

    fl1, fl2, fl3, fl4 = st.columns(4)
    with fl1: f_date  = st.selectbox("📅 일자별",    all_dates, key=f"fd_{label}", label_visibility="visible")
    with fl2: f_media = st.selectbox("📰 언론사",    all_media, key=f"fm_{label}", label_visibility="visible")
    with fl3: f_sent  = st.selectbox("🎨 논조",      all_sent,  key=f"fs_{label}", label_visibility="visible")
    with fl4: f_cat   = st.selectbox("🏷️ 카테고리", all_cat,   key=f"fc_{label}", label_visibility="visible")


    if f_date!="전체": fdf=fdf[fdf["일자"]==f_date]
    if f_media!="전체": fdf=fdf[fdf["매체"]==f_media]
    if f_sent!="전체": fdf=fdf[fdf["감성"]==f_sent]
    if f_cat!="전체": fdf=fdf[fdf["카테고리"]==f_cat]
    fdf = fdf.reset_index(drop=True)

    sk = f"s_{label}"
    if sk not in st.session_state: st.session_state[sk] = 30
    ddf = fdf.iloc[:st.session_state[sk]]

    rh = ""
    prev_date   = None
    shown_divider = False   # 날짜 그룹 내 지방/전문지 구분선 삽입 여부
    for i, row in enumerate(ddf.to_dict("records"), 1):
        cur_date  = row["일자"]
        is_major  = is_major_media(row["매체"])

        # 날짜가 바뀌면 구분선 플래그 리셋
        if cur_date != prev_date:
            prev_date     = cur_date
            shown_divider = False

        # 같은 날짜 안에서 주요→기타 첫 전환 시 구분선 삽입
        if not is_major and not shown_divider:
            shown_divider = True
            rh += (
                "<tr><td colspan='7' style='padding:2px 8px;background:#f5f5f5;"
                "font-size:9px;color:#999;font-family:" + FONT_KR + ";'>"
                "▼ 지방지·전문지·인터넷매체</td></tr>"
            )

        light = sentiment_light(row["감성"])
        gi2   = MEDIA_GRADE.get(row["매체"],{}); grade=gi2.get("grade",""); gc_=GRADE_COLOR.get(grade,"#ccc")
        gs    = f"<span style='background:{gc_};color:white;padding:0 3px;border-radius:2px;font-size:8px;font-weight:700;'>{grade}</span>" if grade else ""
        # 비주요매체 행 배경 살짝 회색
        row_bg = "background:#fafafa;" if not is_major else ""
        summ   = str(row.get('요약',''))[:30]
        rh += (
            f"<tr style='{row_bg}'>"
            f"<td style='text-align:center;color:#aaa;font-size:10px;padding:4px 6px;'>{i}</td>"
            f"<td style='font-size:10px;padding:4px 6px;'>{row['일자']}</td>"
            f"<td style='font-size:10px;padding:4px 6px;'>{row['매체']} {gs}</td>"
            f"<td style='padding:4px 6px;'><a href='{row['링크']}' target='_blank' "
            f"style='color:{'#003366' if is_major else '#555'};text-decoration:none;font-size:11px;'>"
            f"{row['헤드라인']}</a></td>"
            f"<td style='color:#666;font-size:10px;padding:4px 6px;'>{summ}</td>"
            f"<td style='text-align:center;font-size:16px;padding:4px 6px;'>{light}</td>"
            f"<td style='color:#999;font-size:9px;padding:4px 6px;'>{row.get('카테고리','—')}</td>"
            f"</tr>"
        )

    st.markdown(f"""<div style='overflow-x:auto;margin-top:6px;'><table style='width:100%;border-collapse:collapse;font-family:{FONT_KR};'>
    <thead><tr style='background:#003366;color:white;font-size:11px;'>
      <th style='padding:6px 8px;'>No.</th><th style='padding:6px 8px;'>일자</th><th style='padding:6px 8px;'>언론사</th><th style='padding:6px 8px;'>헤드라인</th><th style='padding:6px 8px;'>30자 요약</th><th style='padding:6px 8px;'>논조</th><th style='padding:6px 8px;'>카테고리</th>
    </tr></thead>
    <tbody>{rh}</tbody></table></div>""", unsafe_allow_html=True)

    if st.session_state[sk] < len(fdf):
        if st.button("▼ 더보기", key=f"more_{label}"): st.session_state[sk]+=30; st.rerun()

    # 다운로드 바
    dl1,dl2,dl3 = st.columns(3)
    with dl1:
        out=io.BytesIO()
        with pd.ExcelWriter(out,engine='openpyxl') as w: df.to_excel(w,index=False,sheet_name="데이터")
        out.seek(0)
        st.download_button("📥 엑셀", data=out, file_name=f"한전뉴스_{label}_{(datetime.utcnow()+timedelta(hours=9)).strftime('%Y%m%d')}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True, key=f"xl_{label}")
    with dl2:
        wb2=make_full_word(cd)
        st.download_button("📄 전체 보고서 워드", data=wb2, file_name=f"KEPCO_{label}_{(datetime.utcnow()+timedelta(hours=9)).strftime('%Y%m%d')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, key=f"wd2_{label}")
    with dl3:
        components.html("""<button id="cpbtn" onclick="(function(){var u=window.parent.location.href;navigator.clipboard.writeText(u).then(function(){document.getElementById('cpbtn').innerText='✅ 복사됨!';document.getElementById('cpbtn').style.background='#2E7D32';setTimeout(function(){document.getElementById('cpbtn').innerText='🔗 링크 복사';document.getElementById('cpbtn').style.background='#003366';},2000);});})();" style="background:#003366;color:white;border:none;padding:8px 16px;border-radius:5px;cursor:pointer;font-size:12px;font-weight:600;width:100%;">🔗 링크 복사</button>""", height=40)

    st.markdown(f"<div style='background:#003366;color:white;text-align:center;padding:7px;border-radius:4px;margin-top:10px;font-size:10px;opacity:.8;font-family:{FONT_KR};'>⚡ 홍보실에 꼭 필요한 뉴스 분석시스템 <span style='font-size:8px;opacity:.75;'>by 글쓰는 여행자</span> | {(datetime.utcnow()+timedelta(hours=9)).strftime('%Y.%m.%d')} | 열독률: 언론진흥재단('23)</div>", unsafe_allow_html=True)
    st.markdown("---")

# ══ APP ═══════════════════════════════════════════════
st.set_page_config(page_title="홍보실에 꼭 필요한 뉴스 분석시스템_by 글쓰는 여행자", layout="wide", page_icon="⚡", initial_sidebar_state="expanded")
st.markdown(f"""<style>
@import url('https://fonts.googleapis.com/css2?family=Noto+Sans+KR:wght@300;400;500;700;800&display=swap');
.main .block-container{{padding-top:.5rem;padding-bottom:.5rem;max-width:1400px;}}
[data-testid="stSidebar"]{{background:#F4F6F9;}}
.stTabs [data-baseweb="tab"]{{font-size:12px;padding:5px 14px;font-family:{FONT_KR};}}
div[data-testid="stVerticalBlock"]>div{{gap:0.3rem;}}
.main p, .main div, .main span, .main td, .main th, .main label {{font-family:{FONT_KR};}}
</style>""", unsafe_allow_html=True)

for k,v in [("history",[]),("analysis_cache",{}),("active_key",None)]:
    if k not in st.session_state: st.session_state[k]=v

# 구독 설정 로드 & 스케줄러 초기화 (앱 시작 시 1회)
if "_sub_loaded" not in st.session_state:
    _sub_cfg = load_sub()
    if _sub_cfg.get("enabled"):
        apply_scheduler(_sub_cfg)
    st.session_state._sub_loaded = True

if not YF_OK: st.warning("📦 주가: pip install yfinance 실행 필요", icon="⚠️")
# 구독자 지정 회사 주가 — 세션에 저장된 ticker 사용 (없으면 빈값)
_custom_ticker = st.session_state.get("header_ticker", "")
_custom_name   = st.session_state.get("header_company", "")
md = get_market_data(custom_ticker=_custom_ticker)
md["custom_name"] = _custom_name  # 표시명 덮어쓰기
st.markdown(f"""<div style='background:#003366;color:white;padding:8px 16px;border-radius:5px;display:flex;justify-content:space-between;align-items:center;margin-bottom:6px;font-family:{FONT_KR};'><span style='font-size:15px;font-weight:700;'>⚡ 홍보실에 꼭 필요한 뉴스 분석시스템</span><span style='font-size:10px;opacity:.7;margin-left:6px;'>by 글쓰는 여행자</span><span style='font-size:8px;opacity:.65;'>{(datetime.utcnow()+timedelta(hours=9)).strftime('%Y.%m.%d')} | 열독률 등급 기반 | 네이버 뉴스 API</span></div>""", unsafe_allow_html=True)
st.markdown(mhdr(md), unsafe_allow_html=True)

with st.sidebar:
    st.markdown(f"<h3 style='font-family:{FONT_KR};'>분석 설정</h3>", unsafe_allow_html=True)

    # ── URL 쿼리 파라미터 자동 분석 ──
    _qp = st.query_params
    _auto_kw   = _qp.get("kw", "")
    _auto_days = int(_qp.get("days", 1))
    if _auto_kw and "auto_run_done" not in st.session_state:
        st.session_state["auto_run_kw"]   = _auto_kw
        st.session_state["auto_run_days"] = _auto_days
        st.session_state["auto_run_done"] = True

    with st.form("mf", clear_on_submit=False):
        kc1s,kc2s = st.columns([5,1])
        with kc1s: keywords_input = st.text_input("🔍 키워드 (Enter=분석)", "", placeholder="키워드 입력 후 Enter")
        with kc2s:
            st.markdown("<div style='padding-top:24px;'>", unsafe_allow_html=True)
            run = st.form_submit_button("Go", use_container_width=True)
            st.markdown("</div>", unsafe_allow_html=True)
        st.caption("쉼표(,)=개별  |  플러스(+)=동시포함")
        cs1,cs2 = st.columns(2)
        with cs1: start_date = st.date_input("시작일", datetime.now()-timedelta(days=7))
        with cs2: end_date = st.date_input("종료일", datetime.now())
        max_articles = st.select_slider("수집 기사 수", [500,1000,2000,3000,5000], value=1000)

    st.markdown("---")
    if st.session_state.history:
        st.markdown("**📋 분석 이력**")
        for i,h in enumerate(st.session_state.history[:10],1):
            nr=h['neg']/h['total']*100 if h['total']>0 else 0
            active=(st.session_state.active_key==h['cache_key'])
            if st.button(f"{'▶ ' if active else ''}#{i} {h['keyword']}\n{h['period']} | 부정{nr:.0f}%", key=f"hb_{i}", use_container_width=True):
                st.session_state.active_key=h['cache_key']; st.rerun()
    else: st.caption("분석 후 이력이 쌓입니다")

    # ══ 구독 알리미 UI ══
    st.markdown("---")
    sub_cfg     = load_sub()
    subs        = sub_cfg.get("subscribers", [])
    status_icon = "🟢" if sub_cfg.get("enabled") else "⚫"
    last_sent   = sub_cfg.get("last_sent", "")
    last_txt    = f"마지막 발송: {last_sent}" if last_sent else "아직 발송 없음"
    ADMIN_PW    = "kepco2025"

    with st.expander(f"{status_icon} 뉴스 알리미 구독", expanded=False):
        st.markdown(
            f"<div style='font-size:10px;color:#888;margin-bottom:10px;font-family:{FONT_KR};'>"
            f"{last_txt} | 현재 구독자 {len(subs)}명</div>",
            unsafe_allow_html=True
        )

        # ── [일반] 구독 신청 / 해제 ──
        st.markdown(
            f"<div style='font-size:12px;font-weight:800;color:#003366;margin-bottom:6px;"
            f"font-family:{FONT_KR};'>📬 구독 신청 / 해제</div>",
            unsafe_allow_html=True
        )
        with st.form("user_sub_form", clear_on_submit=True):
            user_email  = st.text_input("내 이메일", placeholder="my@email.com", label_visibility="collapsed")
            user_kw     = st.text_input("받고 싶은 키워드", value="", placeholder="예: 한국전력, 원전, 전기요금")
            user_days   = st.selectbox(
                "📅 수집 기간",
                [1, 2, 3, 7, 30],
                index=0,
                format_func=lambda x: {1:"최근 1일",2:"최근 2일",3:"최근 3일",7:"최근 1주일",30:"최근 1개월 (트렌드 분석)"}.get(x, f"최근 {x}일"),
                help="리포트에 포함할 기사 수집 기간. 길수록 트렌드 분석에 유리합니다."
            )
            # 회사명 자동검색 (KRX 티커 자동 매핑)
            user_co_name = st.text_input(
                "📌 내 회사명 (선택) — 입력하면 헤더에 실시간 주가 표시",
                value="", placeholder="예: 삼성전자, SK하이닉스, 한국전력"
            )
            st.caption("코스피·코스닥 상장사명 입력 시 티커 자동 조회")
            uh1, uh2    = st.columns(2)
            with uh1:
                user_hour   = st.number_input("발송 시각 (0~23시)", min_value=0, max_value=23, value=6, step=1)
            with uh2:
                user_minute = st.selectbox("발송 분", [0, 10, 20, 30, 40, 50], index=3,
                                           format_func=lambda x: f"{x:02d}분")
            uc1, uc2 = st.columns(2)
            with uc1: sub_btn   = st.form_submit_button("구독 신청", use_container_width=True)
            with uc2: unsub_btn = st.form_submit_button("구독 해제", use_container_width=True)

        if sub_btn or unsub_btn:
            addr = user_email.strip().lower()
            if not addr or "@" not in addr:
                st.error("올바른 이메일 주소를 입력해 주세요.")
            else:
                # 회사명 → 티커 자동 조회
                resolved_name, resolved_ticker = "", ""
                if user_co_name.strip():
                    with st.spinner(f"'{user_co_name}' 티커 조회 중..."):
                        resolved_name, resolved_ticker = lookup_krx_ticker(user_co_name.strip())
                    if resolved_ticker:
                        st.info(f"📌 {resolved_name} ({resolved_ticker}) — 헤더 주가 등록")
                    else:
                        st.warning(f"'{user_co_name}' 을(를) KRX에서 찾지 못했습니다. 상장사명을 정확히 입력해 주세요.")

                fresh      = load_sub()
                fresh_subs = fresh.get("subscribers", [])
                emails     = [s["email"].lower() for s in fresh_subs]
                if sub_btn:
                    if addr in emails:
                        for s in fresh_subs:
                            if s["email"].lower() == addr:
                                s["keyword"]        = user_kw.strip() or "뉴스"
                                s["send_hour"]      = int(user_hour)
                                s["send_minute"]    = int(user_minute)
                                s["days"]           = int(user_days)
                                if resolved_ticker:
                                    s["company_name"]   = resolved_name
                                    s["company_ticker"] = resolved_ticker
                        fresh["subscribers"] = fresh_subs
                        save_sub(fresh); apply_scheduler(fresh)
                        # 헤더 즉시 반영
                        if resolved_ticker:
                            st.session_state["header_company"] = resolved_name
                            st.session_state["header_ticker"]  = resolved_ticker
                            get_market_data.clear()
                        st.success(f"✅ {addr} 설정 업데이트 완료!")
                    else:
                        fresh_subs.append({
                            "email":          addr,
                            "keyword":        user_kw.strip() or "뉴스",
                            "send_hour":      int(user_hour),
                            "send_minute":    int(user_minute),
                            "days":           int(user_days),
                            "company_name":   resolved_name,
                            "company_ticker": resolved_ticker,
                            "joined_at":      (datetime.utcnow()+timedelta(hours=9)).strftime("%Y-%m-%d %H:%M"),
                            "start_date":     (datetime.utcnow()+timedelta(hours=9)).strftime("%Y-%m-%d"),
                        })
                        fresh["subscribers"] = fresh_subs
                        save_sub(fresh); apply_scheduler(fresh)
                        # 헤더 즉시 반영
                        if resolved_ticker:
                            st.session_state["header_company"] = resolved_name
                            st.session_state["header_ticker"]  = resolved_ticker
                            get_market_data.clear()
                        co_txt = f" · 📌 {resolved_name} 주가 헤더 등록" if resolved_ticker else ""
                        st.success(
                            f"매일 {int(user_hour):02d}:{int(user_minute):02d}에 "
                            f"[{user_kw}] 뉴스 리포트를 보내드립니다.{co_txt}"
                        )
                else:
                    if addr not in emails:
                        st.warning(f"{addr} 은(는) 구독 목록에 없습니다.")
                    else:
                        fresh_subs = [s for s in fresh_subs if s["email"].lower() != addr]
                        fresh["subscribers"] = fresh_subs
                        save_sub(fresh); apply_scheduler(fresh)
                        st.success(f"✅ {addr} 구독 해제 완료.")

        st.markdown(
            f"<div style='font-size:10px;color:#aaa;font-family:{FONT_KR};margin-top:4px;'>"
            "이미 구독 중이라면 신청 시 키워드·시각이 업데이트됩니다.</div>",
            unsafe_allow_html=True
        )

        # ── [관리자] 발신 계정·구독자 관리 ──
        st.markdown("<div style='height:12px;'></div>", unsafe_allow_html=True)
        st.markdown(
            f"<div style='font-size:12px;font-weight:800;color:#555;margin-bottom:6px;"
            f"font-family:{FONT_KR};'>🔒 관리자 설정</div>",
            unsafe_allow_html=True
        )
        if "sub_admin_ok" not in st.session_state:
            st.session_state.sub_admin_ok = False

        if not st.session_state.sub_admin_ok:
            with st.form("admin_gate_form", clear_on_submit=True):
                entered_pw = st.text_input("관리자 비밀번호", type="password", placeholder="비밀번호 입력")
                gate_btn   = st.form_submit_button("잠금 해제", use_container_width=True)
            if gate_btn:
                if entered_pw == ADMIN_PW:
                    st.session_state.sub_admin_ok = True; st.rerun()
                else:
                    st.error("비밀번호가 올바르지 않습니다.")
        else:
            cl1, _ = st.columns([1, 2])
            with cl1:
                if st.button("🔒 잠금", key="sub_lock_btn", use_container_width=True):
                    st.session_state.sub_admin_ok = False; st.rerun()

            adm     = load_sub()
            adm_subs = adm.get("subscribers", [])
            if adm_subs:
                st.markdown(
                    f"<div style='font-size:11px;font-weight:700;color:#003366;margin:8px 0 4px;"
                    f"font-family:{FONT_KR};'>구독자 목록 ({len(adm_subs)}명)</div>",
                    unsafe_allow_html=True
                )
                # ── 구독자 테이블 (재디자인) ──
                if "del_checks" not in st.session_state:
                    st.session_state.del_checks = {}

                # 헤더
                _hc = st.columns([0.4, 2.2, 2.5, 0.8, 1.4, 1.2])
                for _hcol, _hlbl in zip(_hc, ["선택","이메일","키워드","발송시각","기준일/주기","발송"]):
                    _hcol.markdown(
                        f"<div style='font-size:9px;font-weight:700;color:white;background:#003366;"
                        f"padding:5px 4px;text-align:center;'>{ _hlbl}</div>",
                        unsafe_allow_html=True
                    )

                for idx_s, s in enumerate(adm_subs):
                    email_key = s["email"].lower()
                    row_bg = "#FAFAFA" if idx_s % 2 == 0 else "white"
                    kw  = s.get("keyword", "—")
                    co  = s.get("company_name", "")
                    co_badge = (
                        f" <span style='background:#E8EAF6;color:#3949AB;font-size:8px;"
                        f"padding:1px 5px;border-radius:10px;'>📌{co}</span>"
                    ) if co else ""
                    d_val    = s.get("days", 1)
                    freq_lbl = {1:"매일",2:"2일마다",3:"3일마다",7:"1주일마다",30:"1개월마다"}.get(d_val, f"{d_val}일마다")
                    sdate    = s.get("start_date") or (s.get("joined_at","")[:10] if s.get("joined_at") else "—")
                    sdate_fmt = sdate.replace("-",".") if sdate and sdate != "—" else "—"
                    hhmm     = f"{s.get('send_hour',6):02d}:{s.get('send_minute',30):02d}"

                    c1, c2, c3, c4, c5, c6 = st.columns([0.4, 2.2, 2.5, 0.8, 1.4, 1.2])
                    with c1:
                        checked = st.checkbox("", key=f"del_chk_{idx_s}",
                                              value=st.session_state.del_checks.get(email_key, False),
                                              label_visibility="collapsed")
                        st.session_state.del_checks[email_key] = checked
                    with c2:
                        st.markdown(
                            f"<div style='font-size:10px;padding:8px 4px;word-break:break-all;"
                            f"min-height:44px;border-bottom:1px solid #eee;'>{s['email']}</div>",
                            unsafe_allow_html=True)
                    with c3:
                        st.markdown(
                            f"<div style='font-size:10px;padding:8px 4px;color:#003366;font-weight:700;"
                            f"min-height:44px;border-bottom:1px solid #eee;'>{kw}{co_badge}</div>",
                            unsafe_allow_html=True)
                    with c4:
                        st.markdown(
                            f"<div style='font-size:11px;padding:10px 4px;text-align:center;font-weight:600;"
                            f"min-height:44px;border-bottom:1px solid #eee;'>{hhmm}</div>",
                            unsafe_allow_html=True)
                    with c5:
                        st.markdown(
                            f"<div style='font-size:10px;padding:6px 4px;text-align:center;"
                            f"min-height:44px;border-bottom:1px solid #eee;'>"
                            f"<span style='color:#1565C0;font-weight:700;'>{sdate_fmt}부터</span><br>"
                            f"<span style='color:#666;'>{freq_lbl}</span></div>",
                            unsafe_allow_html=True)
                    with c6:
                        if st.button("📤 발송", key=f"ind_send_{idx_s}", use_container_width=True):
                            adm_now  = load_sub()
                            test_msg = st.session_state.get("admin_test_msg", "")
                            with st.spinner(f"{s['email']} 발송 중..."):
                                ok_i, msg_i = send_email_report(adm_now, test_addr=s["email"],
                                                                custom_message=test_msg, is_broadcast_test=True)
                            if ok_i:
                                st.toast(f"✅ 발송 완료", icon="✅")
                            else:
                                st.session_state[f"send_err_{idx_s}"] = msg_i

                    # 발송 실패 메시지 — 전체 너비 별도 행
                    if st.session_state.get(f"send_err_{idx_s}"):
                        _e1, _e2 = st.columns([5, 1])
                        with _e1:
                            st.error(f"❌ 발송 실패 ({s['email']}): {st.session_state[f'send_err_{idx_s}']}")
                        with _e2:
                            if st.button("닫기", key=f"err_close_{idx_s}", use_container_width=True):
                                del st.session_state[f"send_err_{idx_s}"]; st.rerun()

                st.markdown("<div style='height:6px;background:#e8e8e8;border-radius:0 0 4px 4px;margin-bottom:8px;'></div>", unsafe_allow_html=True)


                # ── 일괄 삭제 버튼 ──
                selected_emails = [email for email, chk in st.session_state.del_checks.items() if chk]
                dc1, dc2 = st.columns([3, 1])
                with dc1:
                    if selected_emails:
                        st.markdown(
                            f"<div style='font-size:10px;color:#C62828;padding-top:6px;'>"
                            f"선택된 구독자 {len(selected_emails)}명 삭제 예정</div>",
                            unsafe_allow_html=True
                        )
                with dc2:
                    if st.button("🗑 선택 삭제", key="bulk_del_btn", use_container_width=True,
                                 disabled=len(selected_emails)==0):
                        fresh2     = load_sub()
                        fresh_subs2 = fresh2.get("subscribers", [])
                        fresh_subs2 = [s for s in fresh_subs2
                                       if s["email"].lower() not in selected_emails]
                        fresh2["subscribers"] = fresh_subs2
                        save_sub(fresh2); apply_scheduler(fresh2)
                        # 체크 상태 초기화
                        for em in selected_emails:
                            st.session_state.del_checks.pop(em, None)
                        st.success(f"✅ {len(selected_emails)}명 삭제 완료")
                        st.rerun()
            else:
                st.caption("구독자 없음")

            with st.form("sub_form", clear_on_submit=False):
                st.markdown(
                    f"<div style='font-size:11px;font-weight:700;color:#003366;margin:8px 0 4px;"
                    f"font-family:{FONT_KR};'>네이버 발신 계정</div>",
                    unsafe_allow_html=True
                )
                sub_sender  = st.text_input("발신 이메일 (네이버)", value=adm.get("sender_email",""), placeholder="yourname@naver.com")
                sub_pw      = st.text_input("네이버 앱 비밀번호", value=adm.get("sender_pw",""), type="password")
                sub_enabled = st.checkbox("구독 활성화", value=bool(adm.get("enabled", False)))
                as1, as2    = st.columns(2)
                with as1: save_btn = st.form_submit_button("저장", use_container_width=True)
                with as2: test_btn = st.form_submit_button("테스트 발송", use_container_width=True)

            if save_btn or test_btn:
                adm.update({
                    "enabled":      sub_enabled,
                    "sender_email": sub_sender.strip(),
                    "sender_pw":    sub_pw,
                })
                if save_sub(adm):
                    apply_scheduler(adm)
                    if test_btn:
                        with st.spinner("테스트 발송 중..."):
                            ok, msg2 = send_email_report(adm, test_addr=adm["sender_email"])
                        if ok: st.success(f"✅ {msg2}")
                        else:  st.error(f"❌ {msg2}")
                    else:
                        st.success(f"✅ 저장 완료 — {'구독 활성화' if sub_enabled else '비활성화 상태'}")
                else:
                    st.error("저장 실패")

            # ── 전체 즉시 발송 버튼 (폼 밖) ──
            st.markdown("<div style='margin-top:8px;'>", unsafe_allow_html=True)
            st.markdown(
                f"<div style='font-size:11px;font-weight:700;color:#003366;margin-bottom:4px;font-family:{FONT_KR};'>"
                f"📝 테스트 메시지 (선택 — 이메일 상단에 삽입됩니다)</div>",
                unsafe_allow_html=True
            )
            test_msg_input = st.text_area(
                "테스트 메시지", height=68,
                placeholder="예) 안녕하세요! 뉴스 모니터링 서비스 테스트 발송입니다. 잘 도착하는지 확인 부탁드립니다.",
                key="admin_test_msg", label_visibility="collapsed"
            )
            if st.button("📨 전체 즉시 발송 (모든 구독자에게 지금 발송)", use_container_width=True, key="broadcast_now"):
                adm2 = load_sub()
                if not adm2.get("sender_email") or not adm2.get("sender_pw"):
                    st.error("발신 계정을 먼저 저장하세요.")
                elif not adm2.get("subscribers"):
                    st.warning("구독자가 없습니다.")
                else:
                    with st.spinner(f"전체 {len(adm2['subscribers'])}명에게 발송 중..."):
                        ok, msg3 = send_email_report(adm2, is_broadcast_test=True,
                                                     custom_message=test_msg_input)
                    if ok: st.success(f"✅ {msg3}  |  제목 말머리: (테스트)")
                    else:  st.error(f"❌ {msg3}")
            st.markdown("</div>", unsafe_allow_html=True)

            st.markdown(
                f"<div style='background:#FFF8E1;border-left:3px solid #F9A825;padding:8px 10px;"
                f"border-radius:0 4px 4px 0;font-size:10px;color:#555;line-height:1.6;"
                f"font-family:{FONT_KR};margin-top:8px;'>"
                "네이버 메일 → 환경설정 → POP3/SMTP → SMTP 사용 선택<br>"
                "내 정보 → 보안설정 → 앱 비밀번호 발급</div>",
                unsafe_allow_html=True
            )


# ── URL 파라미터 자동 분석 실행 ──
if st.session_state.get("auto_run_kw") and not st.session_state.get("auto_run_rendered"):
    st.session_state["auto_run_rendered"] = True
    _kw   = st.session_state["auto_run_kw"]
    _days = st.session_state.get("auto_run_days", 1)
    _end  = (datetime.utcnow() + timedelta(hours=9)).date()
    _start = _end - timedelta(days=max(1, int(_days)))
    with st.spinner(f"'{_kw}' 자동 분석 중..."):
        _g = parse_kw(_kw)[0]
        _raw = get_news(" ".join(_g["keywords"]), 1000)
        _res = []
        for _a in _raw:
            _pub = _a.get("pubDate","")
            try:
                _ad = datetime.strptime(_pub[:16], "%a, %d %b %Y").date()
                if not (_start <= _ad <= _end): continue
                _ds = _ad.strftime("%Y-%m-%d"); _hs = _pub[17:19] if len(_pub)>18 else "00"
            except: _ds = _pub[:10]; _hs = "00"
            _title = clean(_a.get("title","")); _desc = clean(_a.get("description",""))
            _text  = _title + " " + _desc
            if not is_relevant(_text): continue
            _media = get_media(_a.get("originallink",""), _a.get("link",""))
            _gi    = MEDIA_GRADE.get(_media, {})
            _reporter = extract_reporter(_title, _desc)
            _orig = _a.get("originallink",""); _link = _a.get("link","")
            _res.append({"키워드그룹":_kw,"일자":_ds,"월":_ds[:7],"시간":_hs,
                         "매체":_media,"등급":_gi.get("grade","—"),"열독률":_gi.get("rate",0.05),
                         "헤드라인":_title,"요약":summarize(_desc,30),
                         "감성":get_sentiment(_text),"카테고리":"",
                         "기자":_reporter,"링크":_orig if _orig else _link})
    if _res:
        _res = auto_cat(_res)
        _df  = pd.DataFrame(_res)
        _arts = _df.to_dict("records"); _total = len(_df); _cv = _df["감성"].value_counts()
        _pos_n = int(_cv.get("긍정",0)); _neg_n = int(_cv.get("부정",0)); _neu_n = int(_cv.get("중립",0))
        _period_str = f"{_start.strftime('%Y.%m.%d')} ~ {_end.strftime('%m.%d')}"
        _tnc = _df[_df["감성"]=="부정"]["카테고리"].value_counts().index[0] if _neg_n>0 else "없음"
        _tpc = _df[_df["감성"]=="긍정"]["카테고리"].value_counts().index[0] if _pos_n>0 else "없음"
        _nk  = extract_kws(_arts,"부정"); _pk = extract_kws(_arts,"긍정")
        _tnkw = _nk[0][0] if _nk else None
        _daily = _df.groupby("일자").size()
        _tt = f"총 {_total}건"
        _it = f"'{_kw}' 분석 결과, '{_tnc}' 이슈 중심으로 부정 언론 환경이 형성되어 선제적 대응이 필요합니다."
        _crs = gen_criticisms(_arts, _kw)
        _neg_med = [m for m,_ in _df[_df["감성"]=="부정"]["매체"].value_counts().head(5).items()]
        _pr_s, _pr_l, _pr_c = calc_pr_risk(_neg_n, _total, _nk, False, _neg_med)
        _ck = f"{_kw}_{_period_str}"
        _cd = {"label":_kw,"period_str":_period_str,"df":_df,"articles":_arts,"total":_total,
               "pos_n":_pos_n,"neg_n":_neg_n,"neu_n":_neu_n,"neg_kws":_nk,"neu_kws":[],
               "pos_kws":_pk,"top_neg_kw":_tnkw,"criticisms":_crs,"insights_text":_it,
               "top_neg_cat":_tnc,"top_pos_cat":_tpc,"top3_media":"","trend_txt":_tt,
               "crisis_kws":[],"pr_score":_pr_s,"pr_lvl":_pr_l,"pr_color":_pr_c}
        st.session_state.analysis_cache[_ck] = _cd
        st.session_state.active_key = _ck
        render_report(_cd)
    else:
        st.warning(f"'{_kw}' 관련 기사가 없습니다.")

elif run:

    st.session_state.active_key = None
    kw_groups = parse_kw(keywords_input)
    crisis_kws = ["전기요금 폭탄","정전","파업","감사원","수사","비리","횡령","사망","폭발"]
    all_res = []
    for g in kw_groups:
        lbl = g["label"]
        with st.spinner(f"'{lbl}' 수집 중... (최대 {max_articles}건)"):
            raw = get_news(" ".join(g["keywords"]), max_articles)
            for a in raw:
                pub = a.get("pubDate","")
                try:
                    ad = datetime.strptime(pub[:16], "%a, %d %b %Y").date()
                    if not (start_date<=ad<=end_date): continue
                    ds = ad.strftime("%Y-%m-%d"); hs = pub[17:19] if len(pub)>18 else "00"
                except: ds=pub[:10]; hs="00"
                title=clean(a.get("title","")); desc=clean(a.get("description",""))
                text=title+" "+desc; orig=a.get("originallink",""); link=a.get("link","")
                if not is_relevant(text): continue
                if g["type"]=="AND" and not matches_and(text,g): continue
                media=get_media(orig,link); gi=MEDIA_GRADE.get(media,{})
                reporter=extract_reporter(title, desc)
                all_res.append({"키워드그룹":lbl,"일자":ds,"월":ds[:7],"시간":hs,"매체":media,"등급":gi.get("grade","—"),"열독률":gi.get("rate",0.05),"헤드라인":title,"요약":summarize(desc,30),"감성":get_sentiment(text),"카테고리":"","기자":reporter,"링크":orig if orig else link})
    if not all_res: st.error("수집된 기사가 없습니다."); st.stop()
    all_res = auto_cat(all_res)
    all_res = [a for item in [apply_disambig([a],a["키워드그룹"]) for a in all_res] for a in item]
    df_all = pd.DataFrame(all_res)
    for g in kw_groups:
        lbl=g["label"]; df=df_all[df_all["키워드그룹"]==lbl].copy()
        if df.empty: st.warning(f"'{lbl}' 기사 없음"); continue
        arts=df.to_dict("records"); total=len(df); cv=df["감성"].value_counts()
        pos_n=int(cv.get("긍정",0)); neg_n=int(cv.get("부정",0)); neu_n=int(cv.get("중립",0))
        period_str=f"{start_date.strftime('%Y.%m.%d')} ~ {end_date.strftime('%m.%d')}"
        top3m=", ".join(list(df["매체"].value_counts().index[:3]))
        neg_cats_vc = df[df["감성"]=="부정"]["카테고리"].value_counts()
        neg_cats_filtered = [c for c in neg_cats_vc.index if c != "기타"]
        tnc = neg_cats_filtered[0] if neg_cats_filtered else (neg_cats_vc.index[0] if neg_n>0 else "없음")
        pos_cats_vc = df[df["감성"]=="긍정"]["카테고리"].value_counts()
        pos_cats_filtered = [c for c in pos_cats_vc.index if c != "기타"]
        tpc = pos_cats_filtered[0] if pos_cats_filtered else (pos_cats_vc.index[0] if pos_n>0 else "없음")
        nk=extract_kws(arts,"부정"); uk=[]; pk=extract_kws(arts,"긍정")
        tnkw=nk[0][0] if nk else None
        daily=df.groupby("일자").size()
        if len(daily)>=2:
            fh=daily.iloc[:len(daily)//2].mean(); sh=daily.iloc[len(daily)//2:].mean()
            tt=(f"증가 추세({fh:.1f}→{sh:.1f}건/일)" if sh>fh*1.3 else f"감소 추세({fh:.1f}→{sh:.1f}건/일)" if sh<fh*0.7 else f"안정적(일평균 {daily.mean():.1f}건)")
        else: tt=f"총 {total}건"
        it=f"'{lbl}' 관련 {period_str} 분석 결과, '{tnc}' 이슈 중심으로 부정 언론 환경이 형성되어 선제적 대응이 필요합니다. '{tpc}' 관련 성과는 수치 중심으로 적극 홍보해야 합니다. {tt}."
        crs=gen_criticisms(arts,lbl)
        neg_med=[m for m,_ in df[df["감성"]=="부정"]["매체"].value_counts().head(5).items()]
        crisis_found=any(any(ck in a["헤드라인"] or ck in a.get("요약","") for a in arts) for ck in crisis_kws)
        pr_s,pr_l,pr_c=calc_pr_risk(neg_n,total,nk,crisis_found,neg_med)
        ck=f"{lbl}_{period_str}"
        cd={"label":lbl,"period_str":period_str,"df":df,"articles":arts,"total":total,"pos_n":pos_n,"neg_n":neg_n,"neu_n":neu_n,"neg_kws":nk,"neu_kws":uk,"pos_kws":pk,"top_neg_kw":tnkw,"criticisms":crs,"insights_text":it,"top_neg_cat":tnc,"top_pos_cat":tpc,"top3_media":top3m,"trend_txt":tt,"crisis_kws":crisis_kws,"pr_score":pr_s,"pr_lvl":pr_l,"pr_color":pr_c}
        st.session_state.analysis_cache[ck]=cd; st.session_state.active_key=ck
        st.session_state.history=[h for h in st.session_state.history if not (h["keyword"]==lbl and h["period"]==period_str)]
        st.session_state.history.insert(0,{"keyword":lbl,"period":period_str,"total":total,"pos":pos_n,"neg":neg_n,"neu":neu_n,"top_issue":tnc,"cache_key":ck})
        st.session_state.history=st.session_state.history[:10]
        render_report(cd)
elif st.session_state.active_key:
    cd2=st.session_state.analysis_cache.get(st.session_state.active_key)
    if cd2: render_report(cd2)
    else: st.warning("다시 분석해 주세요.")
else:
    if st.session_state.history:
        st.markdown(f"<div style='font-size:14px;font-weight:800;color:#003366;border-bottom:2px solid #003366;padding-bottom:5px;margin:10px 0;font-family:{FONT_KR};'>📋 분석 이력</div>", unsafe_allow_html=True)
        for i,h in enumerate(st.session_state.history[:10],1):
            nr=h['neg']/h['total']*100 if h['total']>0 else 0; pr=h['pos']/h['total']*100 if h['total']>0 else 0
            ca,cb=st.columns([5,1])
            with ca:
                st.markdown(f"""<div style='background:white;border:1px solid #e8e8e8;border-radius:4px;padding:8px 12px;margin-bottom:4px;font-family:{FONT_KR};'><span style='font-size:13px;font-weight:700;color:#003366;'>#{i} {h['keyword']}</span><span style='color:#aaa;font-size:10px;margin-left:8px;'>{h['period']}</span><br><span style='font-size:11px;'>총 {h['total']}건 | <span style='color:#C62828;'>부정 {h['neg']}건({nr:.0f}%)</span> | <span style='color:#1565C0;'>긍정 {h['pos']}건({pr:.0f}%)</span></span></div>""", unsafe_allow_html=True)
            with cb:
                if st.button("열람", key=f"v_{i}", use_container_width=True):
                    st.session_state.active_key=h['cache_key']; st.rerun()
    else:
        st.markdown(f"""<div style='text-align:center;padding:50px;color:#aaa;font-family:{FONT_KR};'><div style='font-size:32px;'>⚡</div><div style='font-size:15px;font-weight:600;color:#003366;margin-top:8px;'>홍보실에 꼭 필요한 뉴스 분석시스템 <span style="font-size:9px;opacity:.7;">by 글쓰는 여행자</span></div><div style='font-size:12px;margin-top:6px;'>좌측 키워드 입력 후 🚀 클릭</div></div>""", unsafe_allow_html=True)
